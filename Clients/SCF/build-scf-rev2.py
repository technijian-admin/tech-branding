# -*- coding: utf-8 -*-
"""
SC Fuels - Strategic Digital, SEO & AI Growth Report - REVISION 2
Built after the June 2, 2026 strategy session with Derek Bettencourt.

Framing (per Ravi, 2026-06-03): "Complement, then offer more."
 - Keep the strong technical audit (Derek's team praised it).
 - Keep the coordinated WebFX + Technijian model.
 - ADD the net-new AI growth engine (AI SEO/GEO, AI content/social, AI lead gen, AI consulting).
 - Cost vs WebFX in PUBLISHED RANGES only (no invented client spend, no invented totals).
 - Consolidation to Technijian framed as an OPTIONAL future step.

Pricing discipline (feedback_no_invented_pricing):
 - My SEO published tiers + add-ons (real).
 - My AI Lead Gen $1,499 / $3,499 / $6,999 (real).
 - My AI consulting + My Dev build = "scoped at discovery" (NO invented figure).
Logo discipline (feedback_logo_use_authentic_files): use assets/Technijian Logo 2.png.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------- brand
BLUE      = "006DB6"
ORANGE    = "F67D4B"
TEAL      = "1EAAC8"
CHARTREUSE= "8FB81E"   # darkened CBDB2D for readable text/fill
GREY      = "59595B"
DARK      = "1A1A2E"
NEARBLACK = "2D2D2D"
OFFWHITE  = "F4F6F8"
LIGHTGREY = "E9ECEF"
WHITE     = "FFFFFF"
CRIT      = "CC0000"
PASS      = "1E8E3E"
AMBER     = "C77700"
FONT      = "Arial"

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.abspath(os.path.join(HERE, "..", "..", "assets"))
LOGO_LIGHT = os.path.join(ASSETS, "Technijian Logo 2.png")
LOGO_WHITE = os.path.join(ASSETS, "Technijian Logo - white text.png")
OUT_DOCX = os.path.join(HERE, "SCFuels_Strategic_Report_Technijian_June2026_Rev2.docx")

doc = Document()

# Continuous flow: section headers (with the orange rule + space-before) delineate
# sections, so we do NOT force a page break per section (that strands short sections
# on half-empty pages). Only the cover and TOC get hard breaks.
def secbreak():
    pass

def row_cantsplit(row):
    trPr = row._tr.get_or_add_trPr()
    cs = OxmlElement("w:cantSplit"); cs.set(qn("w:val"), "true"); trPr.append(cs)

# ---------------------------------------------------------------- base styles
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(NEARBLACK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.22

for hi, (sz, col) in {1: (15, BLUE), 2: (12.5, DARK), 3: (11, BLUE)}.items():
    st = doc.styles["Heading %d" % hi]
    st.font.name = FONT
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(col)
    st.paragraph_format.space_before = Pt(10 if hi == 1 else 8)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.keep_with_next = True

# ---------------------------------------------------------------- page setup
sec = doc.sections[0]
sec.page_width  = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin    = Inches(0.85)
sec.bottom_margin = Inches(0.8)
sec.left_margin   = Inches(0.9)
sec.right_margin  = Inches(0.9)
CONTENT_W = Inches(8.5 - 1.8)   # 6.7"

# ---------------------------------------------------------------- low-level helpers
def _shade(el, hexcolor):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexcolor)
    el.append(shd)

def cell_bg(cell, hexcolor):
    _shade(cell._tc.get_or_add_tcPr(), hexcolor)

def _set_borders(tcPr, edges, color, sz):
    tb = OxmlElement("w:tcBorders")
    for edge in edges:
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        tb.append(e)
    tcPr.append(tb)

def cell_border(cell, edges=("top","bottom","left","right"), color=LIGHTGREY, sz=4):
    _set_borders(cell._tc.get_or_add_tcPr(), edges, color, sz)

def no_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        e = OxmlElement("w:" + edge); e.set(qn("w:val"), "none"); e.set(qn("w:sz"), "0")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "auto"); borders.append(e)
    tblPr.append(borders)

def cell_margins(cell, top=40, bottom=40, left=90, right=90):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for k, v in (("top",top),("bottom",bottom),("left",left),("right",right)):
        n = OxmlElement("w:"+k); n.set(qn("w:w"), str(v)); n.set(qn("w:type"), "dxa"); m.append(n)
    tcPr.append(m)

def set_run(r, size=10.5, color=NEARBLACK, bold=False, italic=False, font=FONT):
    r.font.name = font; r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)

def fill_cell(cell, text, size=9.5, color=NEARBLACK, bold=False, align=None,
              bg=None, italic=False, valign="center"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.line_spacing = 1.08
    if align: p.alignment = align
    # support simple bold-prefix "**x** rest"
    parts = text.split("**")
    if len(parts) > 1:
        for i, seg in enumerate(parts):
            if seg == "": continue
            r = p.add_run(seg); set_run(r, size, color, bold or (i % 2 == 1), italic)
    else:
        r = p.add_run(text); set_run(r, size, color, bold, italic)
    if bg: cell_bg(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER if valign == "center" else WD_ALIGN_VERTICAL.TOP
    cell_margins(cell)
    return cell

def body(text, size=10.5, color=NEARBLACK, italic=False, space_after=6, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if bold_lead:
        r = p.add_run(bold_lead); set_run(r, size, DARK, True)
    r = p.add_run(text); set_run(r, size, color, False, italic)
    return p

def bullet(text, bold_lead=None, color=NEARBLACK):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.18
    if bold_lead:
        r = p.add_run(bold_lead); set_run(r, 10, DARK, True)
    r = p.add_run(text); set_run(r, 10, color)
    return p

def h1(num, title):
    p = doc.add_heading(level=1)
    p.paragraph_format.space_before = Pt(14)
    r = p.add_run(("%s  %s" % (num, title)) if num else title)
    set_run(r, 15, BLUE, True)
    # orange bottom rule
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "3"); bottom.set(qn("w:color"), ORANGE)
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def h2(title):
    p = doc.add_heading(level=2)
    r = p.add_run(title); set_run(r, 12.5, DARK, True)
    return p

def spacer(pts=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pts)
    for r in p.runs: r.font.size = Pt(2)
    return p

def data_table(headers, rows, widths=None, header_bg=BLUE, font_sz=9.5,
               zebra=True, status_col=None):
    n = len(headers)
    t = doc.add_table(rows=1, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.allow_autofit = False
    if widths:
        total = sum(widths)
        for ci, w in enumerate(widths):
            wEmu = int(CONTENT_W * (w / total))
            for r in t.rows:
                r.cells[ci].width = wEmu
    # header
    for ci, htext in enumerate(headers):
        c = t.rows[0].cells[ci]
        fill_cell(c, htext, size=font_sz, color=WHITE, bold=True, bg=header_bg,
                  align=WD_ALIGN_PARAGRAPH.LEFT)
        cell_border(c, ("bottom",), color=header_bg, sz=4)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        bg = OFFWHITE if (zebra and ri % 2 == 0) else WHITE
        for ci, val in enumerate(row):
            c = cells[ci]
            col = NEARBLACK
            cellbg = bg
            if status_col is not None and ci == status_col:
                token = str(val).strip().upper()
                cmap = {"PASS": PASS, "A": PASS, "QUICK WIN": PASS,
                        "WARN": AMBER, "B": BLUE, "C": GREY, "MEDIUM": AMBER,
                        "FAIL": CRIT, "CRITICAL": CRIT, "F": CRIT, "D": AMBER,
                        "URGENT": CRIT, "HIGH VALUE": BLUE, "ONGOING": TEAL,
                        "STRATEGIC": BLUE, "INFO": GREY, "NEW": TEAL,
                        "VALIDATED": PASS}
                if token in cmap:
                    cellbg = cmap[token]; col = WHITE
            fill_cell(c, str(val), size=font_sz, color=col, bg=cellbg,
                      bold=(ci == 0 and status_col != 0),
                      align=WD_ALIGN_PARAGRAPH.LEFT)
            if widths:
                total = sum(widths); c.width = int(CONTENT_W * (widths[ci] / total))
            cell_border(c, ("bottom",), color=LIGHTGREY, sz=2)
    for r in t.rows:
        row_cantsplit(r)
    no_table_borders(t)
    spacer(6)
    return t

def callout(title, lines, accent=BLUE, fill=OFFWHITE, title_color=None):
    """Single-cell box with thick colored left border."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]; c.width = CONTENT_W
    cell_bg(c, fill)
    _set_borders(c._tc.get_or_add_tcPr(), ("top","bottom","right"), LIGHTGREY, 4)
    # thick left accent
    tcPr = c._tc.get_or_add_tcPr()
    tb = tcPr.find(qn("w:tcBorders"))
    if tb is None:
        tb = OxmlElement("w:tcBorders"); tcPr.append(tb)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "28")
    left.set(qn("w:space"), "0"); left.set(qn("w:color"), accent)
    tb.append(left)
    cell_margins(c, top=90, bottom=90, left=160, right=140)
    c.text = ""
    if title:
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(3)
        r = p.add_run(title); set_run(r, 10.5, title_color or accent, True)
    for ln in lines:
        p = c.add_paragraph(); p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.18
        if isinstance(ln, tuple):
            lead, rest = ln
            r = p.add_run(lead); set_run(r, 9.7, DARK, True)
            r = p.add_run(rest); set_run(r, 9.7, NEARBLACK)
        else:
            r = p.add_run(ln); set_run(r, 9.7, NEARBLACK)
    row_cantsplit(t.rows[0])
    spacer(7)
    return t

def kpi_cards(cards):
    """Row of stat cards: list of (big, label, color)."""
    n = len(cards)
    t = doc.add_table(rows=1, cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, (big, label, color) in enumerate(cards):
        c = t.rows[0].cells[ci]
        c.width = int(CONTENT_W / n)
        cell_bg(c, OFFWHITE)
        cell_border(c, ("top","bottom","left","right"), color=WHITE, sz=20)
        cell_margins(c, top=110, bottom=110, left=80, right=80)
        c.text = ""
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(big); set_run(r, 21, color, True)
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        r = p2.add_run(label); set_run(r, 8, GREY, True)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    no_table_borders(t)
    spacer(8)
    return t

def full_bar(color, height_pt=10):
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]; c.width = CONTENT_W
    cell_bg(c, color); no_table_borders(t)
    c.text = ""
    c.paragraphs[0].paragraph_format.space_after = Pt(0)
    c.paragraphs[0].paragraph_format.space_before = Pt(0)
    for r in c.paragraphs[0].runs: r.font.size = Pt(height_pt)
    rr = c.paragraphs[0].add_run(" "); rr.font.size = Pt(height_pt)
    return t

# ---------------------------------------------------------------- header / footer
def add_footer():
    footer = sec.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]; p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # top border rule (orange)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); top = OxmlElement("w:top")
    top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "6"); top.set(qn("w:space"), "5"); top.set(qn("w:color"), ORANGE)
    pbdr.append(top); pPr.append(pbdr)
    r = p.add_run("Confidential  |  Technijian Technology Solutions  |  technijian.com  |  Page ")
    set_run(r, 8, GREY)
    # PAGE field
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    rr = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    f = OxmlElement("w:rFonts"); f.set(qn("w:ascii"), FONT); f.set(qn("w:hAnsi"), FONT); rpr.append(f)
    szz = OxmlElement("w:sz"); szz.set(qn("w:val"), "16"); rpr.append(szz)
    colr = OxmlElement("w:color"); colr.set(qn("w:val"), GREY); rpr.append(colr)
    rr.append(rpr); t = OxmlElement("w:t"); t.text = "1"; rr.append(t); fld.append(rr)
    p._p.append(fld)

def add_running_header():
    header = sec.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]; p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("SC Fuels  —  AI Growth & Digital Strategy  —  Revision 2")
    set_run(r, 8, "9AA0A6", italic=True)

# ================================================================ COVER
def cover():
    full_bar(BLUE, 14); spacer(2)
    # logo
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(LOGO_LIGHT, width=Inches(2.7))
    # orange divider
    full_bar(ORANGE, 4); spacer(10)

    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(40); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("SC FUELS"); set_run(r, 40, DARK, True)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Strategic Digital, SEO & AI Growth Report"); set_run(r, 20, BLUE, True)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(16)
    r = p.add_run("How an AI-Driven Growth Engine Helps SC Fuels Add Business — Without Adding Headcount")
    set_run(r, 12.5, GREY, italic=True)

    # revision tag
    t = doc.add_table(rows=1, cols=1); c = t.rows[0].cells[0]; c.width = Inches(2.5)
    cell_bg(c, ORANGE); cell_margins(c, 50, 50, 120, 120); no_table_borders(t)
    c.text = ""; pp = c.paragraphs[0]
    r = pp.add_run("REVISION 2"); set_run(r, 11, WHITE, True)
    spacer(26)

    meta = [
        ("Prepared for", "SC Fuels (Southern Counties Oil Co., L.P.)  ·  scfuels.com"),
        ("Attention", "Derek Bettencourt  ·  and the SC Fuels leadership & sales/marketing team"),
        ("Prepared by", "Technijian Technology Solutions  ·  Ravi Jain, Founder & CEO"),
        ("Date", "June 2026"),
        ("Context", "Follow-up to the June 2, 2026 strategy session"),
        ("Sources", "Semrush Domain Overview  ·  The HOTH SEO Audit  ·  June 2, 2026 session notes"),
    ]
    t = doc.add_table(rows=0, cols=2)
    for label, val in meta:
        cells = t.add_row().cells
        cells[0].width = Inches(1.5); cells[1].width = Inches(5.2)
        fill_cell(cells[0], label, size=9.5, color=BLUE, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        fill_cell(cells[1], val, size=9.5, color=NEARBLACK, align=WD_ALIGN_PARAGRAPH.LEFT)
    no_table_borders(t)
    spacer(40)
    full_bar(ORANGE, 4); spacer(2)
    p = doc.add_paragraph()
    r = p.add_run("CONFIDENTIAL — Prepared exclusively for SC Fuels leadership. Not for external distribution.")
    set_run(r, 8.5, GREY, italic=True)
    doc.add_page_break()

# ================================================================ TOC
def toc():
    p = doc.add_paragraph()
    r = p.add_run("Contents"); set_run(r, 16, BLUE, True)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12"); bottom.set(qn("w:space"), "3"); bottom.set(qn("w:color"), ORANGE)
    pbdr.append(bottom); pPr.append(pbdr)
    fld = OxmlElement("w:p"); r2 = OxmlElement("w:r")
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin"); r2.append(fc)
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '; r2.append(instr)
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "separate"); r2.append(fc2)
    fld.append(r2)
    r3 = OxmlElement("w:r"); t3 = OxmlElement("w:t")
    t3.text = "Right-click and choose “Update Field” to build the table of contents."
    r3.append(t3); fld.append(r3)
    r4 = OxmlElement("w:r"); fc3 = OxmlElement("w:fldChar"); fc3.set(qn("w:fldCharType"), "end"); r4.append(fc3); fld.append(r4)
    doc.paragraphs[-1]._p.addnext(fld)
    doc.add_page_break()

# ================================================================ build
add_footer(); add_running_header()
cover()
toc()

# ---------- scorecard up top
h1("", "Digital & AI Readiness Scorecard")
body("This revision keeps the full June 2026 technical audit (the section your team reviewed) and "
     "adds the AI growth plan we discussed on June 2 — inbound visibility, outbound lead generation, "
     "and internal AI that lets each salesperson cover far more ground. Grades below are unchanged from "
     "the audit; the AI-readiness rows are new.")
data_table(
    ["Category", "Grade", "Status Summary"],
    [
        ["Overall SEO Health", "B+", "Strong foundation — targeted gaps identified"],
        ["On-Page SEO", "A", "Well-optimized title, meta, schema, SSL"],
        ["Generative Engine Optimization (GEO)", "F", "CRITICAL — largely invisible to AI search engines"],
        ["Usability / Mobile", "A", "Responsive, correct viewport, no Flash/iFrames"],
        ["Page Performance", "D", "6.1MB load, 7.9s content, HTTP/1 only"],
        ["Organic Search (23,500/mo)", "B", "Stable — but 60% dependent on branded terms"],
        ["Paid Search (215 visits/mo)", "C", "Only 2 keywords active — significant gap"],
        ["Backlink Authority (5,100 links)", "B", "986 referring domains — healthy base"],
        ["Local SEO", "C", "GBP verified, but address & schema missing on page"],
        ["AI Outbound / Lead Gen", "F", "NEW — no AI-driven outbound today; biggest growth lever"],
        ["Internal AI (sales / credit / CRM)", "F", "NEW — Salesforce, credit, FleetSeek not AI-orchestrated yet"],
    ],
    widths=[40, 12, 78], status_col=1)
body("This report is intended exclusively for SC Fuels leadership and authorized partners.",
     size=8.5, color=GREY, italic=True)
secbreak()

# ================================================================ 1 EXEC SUMMARY
h1("1.", "Executive Summary")
body("SC Fuels (scfuels.com) is one of the oldest and largest petroleum distributors in the United "
     "States — a legacy industrial brand, acquired by Pilot Company in 2021, headquartered in Orange, "
     "California, and serving 11,000+ commercial customers across 15 western states. This revision builds "
     "on the June 2026 technical audit and folds in the priorities Derek Bettencourt shared in our June 2 "
     "strategy session.")
body("", bold_lead="The thread that ties it together: ")
body("Derek named a clear goal — grow the book of business without simply adding salespeople, by giving "
     "each rep AI that does the heavy lifting so one person can cover two, three, or four times the ground. "
     "Everything below serves that goal. WebFX continues to own the SEO and content layer it does well; "
     "Technijian adds the parts WebFX does not provide — AI-driven outbound lead generation, the AI "
     "search (GEO) layer, and the internal AI that multiplies each salesperson — and stands ready to take "
     "on more if and when SC Fuels wants to consolidate.")

kpi_cards([
    ("23,500", "ORGANIC VISITS / MO", BLUE),
    ("60%", "TRAFFIC IS BRANDED", ORANGE),
    ("Grade F", "AI-SEARCH VISIBILITY", CRIT),
    ("$100K+", "VALUE / NEW FLEET ACCOUNT", PASS),
])

h2("Three findings from the audit (unchanged)")
body("", bold_lead="1.  GEO Grade F (Critical).  ")
body("About 89% of the homepage is JavaScript-rendered, which makes it largely invisible to LLMs and AI "
     "search (Google AI Overviews, ChatGPT, Perplexity, Gemini). For a company whose buyers increasingly "
     "ask an AI assistant “who are the best bulk diesel suppliers in the west?” before they ever call, "
     "this is a real and growing blind spot.")
body("", bold_lead="2.  Paid search severely underutilized.  ")
body("SC Fuels runs exactly 2 paid keywords (~215 visits/mo) while a single competitor runs 195. Given the "
     "lifetime value of a commercial fleet contract, targeted paid search is one of the fastest paths to "
     "measurable new-account ROI available today.")
body("", bold_lead="3.  60% of organic traffic is branded.  ")
body("More than half of organic traffic comes from people who already know SC Fuels by name. Non-branded "
     "search — the genuine acquisition channel — is only ~40% and needs a deliberate content and keyword "
     "program.")

h2("Three things we added after June 2")
callout("What changed in Revision 2", [
    ("Inbound is necessary but not sufficient.  ",
     "SEO brings warmer leads in. To grow without adding headcount you also need outbound that runs "
     "itself — so we added an AI Lead Generation engine (Section 12, Pillar 3)."),
    ("Reach the buyer who is replacing the phone call.  ",
     "Derek flagged a generational shift: younger buyers at construction, trucking and tree-service "
     "firms research on social, video and AI — not cold calls. The content + social engine (Pillar 2) "
     "is built for them."),
    ("Multiply the salesperson, don’t replace them.  ",
     "Pilot has already cut new-customer onboarding from 27 days to 3 with AI. The internal-AI plan "
     "(Pillar 4) targets the same wins for SC Fuels — Salesforce opportunity progression and the credit "
     "process — so reps spend time on relationships, not paperwork."),
], accent=ORANGE)

body("The rest of this report walks the audit findings first (Sections 2–11), then lays out the four-pillar "
     "AI growth engine and how it works alongside WebFX (Section 12), the coordinated operating model and the "
     "optional path to consolidate (Section 13), the investment picture against what agencies like WebFX "
     "typically charge (Section 14), a sequenced roadmap (Section 15), and no-commitment quick wins (Section 16).")
secbreak()

# ================================================================ 2 WHAT WE HEARD
h1("2.", "What We Heard — June 2 Strategy Session")
body("Before the analysis, here is what Derek told us on June 2, and how each point maps to the plan. "
     "Every row below is something SC Fuels said out loud — not a cold assumption — which is why we treat "
     "these as validated priorities rather than guesses.")
data_table(
    ["What SC Fuels said", "What it means for the plan", "Status"],
    [
        ["“Add business without [just] adding salespeople” — Derek’s stated goal",
         "Reframe as: AI lets one rep cover 2–4× the accounts. The organizing idea of the whole engine.", "VALIDATED"],
        ["“We had a budget approved to redo the entire website… to incorporate AI tools”",
         "Build the new site AI-ready and GEO-ready from day one (Pillars 1 & 4). Timing is now.", "VALIDATED"],
        ["Wants “warmer leads”, less dependence on cold calls / door knocks",
         "Inbound content + AI outbound that lands warm, ICP-scored leads in the rep’s inbox.", "VALIDATED"],
        ["“Glaring hole… municipalities, government sector” (RFPs favor the incumbent)",
         "AI RFP discovery + auto-drafting to compete on more bids with the same team (Pillar 3).", "VALIDATED"],
        ["Younger buyers research on social, text, video, AI — not phone calls",
         "Blog → video → social → podcast factory + AI search visibility (Pillar 2).", "VALIDATED"],
        ["“Not thrilled with how we’re utilizing Salesforce” — opportunities stall",
         "AI that pushes opportunities through the sales path; ties InsightView + FleetSeek together (Pillar 4).", "VALIDATED"],
        ["“I’d love a tool that got the credit process refined or improved”",
         "AI-assisted credit intake & decision acceleration — faster yes/no, less manual review.", "VALIDATED"],
        ["Pilot cut onboarding 27 days → 3 days with AI; “AI is a huge topic” at Pilot",
         "A concrete internal benchmark SC Fuels can match — and show Pilot it is moving.", "VALIDATED"],
        ["Competes with card NETWORKS (WEX, Voyager), not just other distributors",
         "Widen the competitive lens beyond fuel distributors (Section 8).", "VALIDATED"],
    ],
    widths=[40, 48, 14], status_col=2, font_sz=9)
callout("Why this matters", [
    "An AI-built plan is only as good as its accuracy. Derek’s feedback confirmed the audit got the "
    "business right (“I didn’t see anything in here that did not fit who it is we are”) and sharpened two "
    "things our cold research under-weighted: the card-network competition (WEX/Voyager) and the "
    "government/RFP opportunity. Both are now reflected in the plan.",
], accent=BLUE)
secbreak()

# ================================================================ 3 COMPANY & SITE
h1("3.", "SC Fuels — Company & Site Architecture")
h2("3.1  Company Profile")
data_table(
    ["Field", "Detail"],
    [
        ["Website", "www.scfuels.com"],
        ["Headquarters", "1800 W Katella Ave #400, Orange, CA 92867"],
        ["Parent Company", "Pilot Company (acquired 2021). Note: the website still reads “family owned” — outdated."],
        ["Service Territory", "15 western states"],
        ["Customers", "11,000+ commercial and enterprise customers"],
        ["Tagline", "Your Single Choice for Petroleum Products"],
        ["Google Business Profile", "Verified — 4.5 stars / 20 reviews — Orange, CA"],
        ["Current marketing partner", "WebFX (SEO/analytics) + InsightView + FleetSeek → Salesforce CRM"],
    ],
    widths=[28, 72])
body("Derek confirmed the live site is dated (“the website’s really outdated… it still says family owned "
     "when it’s actually Pilot owned now”) and that a full rebuild is already budgeted with AI tooling in "
     "scope. That makes this the right moment to build the next site AI-ready and GEO-ready from day one, "
     "rather than retrofitting later.", )

h2("3.2  Core Services")
data_table(
    ["Service Line", "Description", "B2B Value"],
    [
        ["Unbranded Wholesale Fuel", "Bulk diesel & gasoline delivery to commercial sites", "Price-competitive supply"],
        ["Onsite Fleet Fueling", "Fuel brought directly to fleet vehicles at depot", "Eliminates driver downtime"],
        ["Cardlock Fleet Fueling", "Proprietary fleet card network, secure access", "Purchase controls, IFTA reporting"],
        ["Lubricants & DEF", "Engine oils, hydraulic fluids, DEF", "Single-supplier consolidation"],
        ["Alternative Fuels", "CNG, LNG, hydrogen, biodiesel, renewable diesel", "Clean-fuel mandate compliance"],
    ],
    widths=[26, 44, 30])

h2("3.3  Site Architecture & the Retention Moat")
body("scfuels.com is a pure B2B acquisition-and-retention engine, and understanding it is essential to "
     "every SEO and AI decision:")
bullet("New commercial accounts: primary CTAs are credit applications and quote requests — a multi-step, "
       "high-touch sales process that can span weeks. Marketing’s job is to get the right fleet manager or "
       "procurement officer into that funnel.", bold_lead="Dual-purpose funnel.  ")
bullet("Fleet Card, Lubricant and Wholesale portals embed SC Fuels into the customer’s daily workflow — "
       "MPG tracking, per-driver controls, IFTA exports. A fleet manager who runs compliance through your "
       "portal does not switch over a marginally cheaper gallon. Every portal feature deepens the switching "
       "cost.", bold_lead="The portal is a moat.  ")
bullet("Across 15 states, the site must capture both national intent (“bulk diesel delivery”) and local "
       "intent (“onsite fleet fueling Los Angeles”). Dedicated service-area pages are a clear opportunity.",
       bold_lead="Localized SEO need.  ")
secbreak()

# ================================================================ 4 WEBFX
h1("4.", "The Marketing Landscape Today — WebFX, and Where AI Extends It")
body("Derek asked us to look honestly at WebFX — what they do well, where the gaps are, and how Technijian "
     "is different. This section does exactly that. The short version: WebFX is a capable SEO and analytics "
     "partner, and nothing here suggests ripping that out. What WebFX does not do is the AI-driven outbound, "
     "AI search, and internal-AI work that Derek’s goals actually require — and that is precisely where "
     "Technijian adds value.")

h2("4.1  WebFX — Agency Profile")
data_table(
    ["Field", "Detail"],
    [
        ["Headquarters", "Harrisburg, Pennsylvania (multiple US offices)"],
        ["Founded / Size", "1996 · 500+ digital marketing specialists"],
        ["Reputation", "Clutch 4.9/5 · Google Premier Partner"],
        ["Core Services", "SEO, PPC/paid media, content, web design, social, email, CRO"],
        ["Proprietary Tech", "MarketingCloudFX (lead-lifecycle attribution) + CRM integrations"],
        ["Published Pricing", "SEO from ~$2,500/mo · core services $3,000+/mo minimum · enterprise $10K–$50K+/mo"],
        ["Contract Terms", "Typically 6–12 month initial terms + one-time setup fees (per public pricing guides)"],
    ],
    widths=[26, 74])

h2("4.2  What WebFX does well (keep it)")
bullet("Dominating organic rankings for specific, high-intent commercial keywords — their core competency, "
       "and exactly what converts a 60% branded dependency into real acquisition.", bold_lead="Search & content authority.  ")
bullet("MarketingCloudFX traces a credit application back to the keyword and content that produced it, and "
       "ties campaigns to signed contracts — critical for justifying spend in a 3–6 month B2B sales cycle.",
       bold_lead="Closed-loop attribution.  ")
bullet("Their model fits high-LTV B2B: you don’t need cheap clicks, you need to own the handful of terms a "
       "procurement director searches before a six-figure contract.", bold_lead="High-LTV specialization.  ")

h2("4.3  Where the gaps are")
data_table(
    ["Gap", "What it means for SC Fuels"],
    [
        ["No AI-driven outbound", "WebFX is an inbound/SEO shop. It does not run AI prospecting, ICP scoring, or RFP auto-drafting — the outbound Derek wants."],
        ["GEO / AI-search is nascent", "Like most traditional SEO agencies, AI-search optimization (llms.txt, AI-citation content) is still emerging for them. SC Fuels’ GEO grade is F today."],
        ["Niche-B2B content", "Public reviews note WebFX writers can lack deep expertise in specialized industries — fuel/energy content often needs heavy client editing."],
        ["Social effectiveness", "Derek’s own read: “WebFX also does some of that [social]. I don’t know that they’re doing it very effectively.”"],
        ["No internal-AI / CRM automation", "Salesforce opportunity progression, the credit process, and FleetSeek/InsightView orchestration are simply outside an SEO agency’s scope."],
        ["Cost + lock-in", "Enterprise engagements run $10K–$50K+/mo on 6–12 month terms; reporting can overwhelm. Worth confirming what you pay today vs. the value you receive."],
    ],
    widths=[30, 70], font_sz=9.3)
callout("The honest framing", [
    "This is not a case for firing WebFX. It is a case for adding the AI layer they do not provide — outbound "
    "lead gen, AI search, and internal automation — and, separately, for confirming what you pay WebFX today "
    "so you can decide later whether to consolidate. We model that economics in Section 14 using published "
    "rates only; we have not assumed your actual WebFX spend.",
], accent=TEAL)
secbreak()

# ================================================================ 5 ORGANIC
h1("5.", "Organic Search Performance")
h2("5.1  Traffic Summary")
data_table(
    ["Metric", "Value"],
    [
        ["Monthly Organic Traffic", "~23,500 estimated visits (Semrush)"],
        ["Semrush Domain Rank", "80,300 (lower is better)"],
        ["Total Ranking Keywords", "13,100+"],
        ["Est. Monthly Traffic Value", "$36,900"],
        ["Branded Traffic Share", "60.06% — strong recognition, but fragile dependency"],
        ["Non-Branded Share", "39.94% — the genuine acquisition channel"],
        ["Primary Market", "United States (76% of keywords); Canada ~4.4%"],
    ],
    widths=[34, 66])
h2("5.2  Top Organic Keywords")
data_table(
    ["Keyword", "Position", "Monthly Volume"],
    [
        ["sc fuels", "#1", "9,900"],
        ["sc fuels near me", "Top 3", "3,600"],
        ["diesel fuel suppliers near me", "#2", "18,100"],
        ["diesel fuel distributors near me", "#3", "18,100"],
        ["sci def", "#4", "33,100"],
        ["sc gas", "#10", "135,000"],
    ],
    widths=[52, 22, 26])
body("Ranking #2–#3 for “diesel fuel suppliers near me” and related terms (18,100 searches/mo) is a real "
     "commercial asset. Pushing these to #1 — and widening the footprint of similar commercial terms — is the "
     "highest-priority organic growth lever, and a natural place for AI-accelerated content to help.")
h2("5.3  Position Distribution — the biggest single lever")
data_table(
    ["Position Range", "Keywords", "Implication"],
    [
        ["1–3 (top SERP)", "333", "Protect with freshness + links"],
        ["4–10 (page 1)", "774", "Quick wins — small gains drive real traffic"],
        ["11–20 (page 2)", "1,100", "Active optimization target"],
        ["21–50", "4,700", "Content depth + internal linking"],
        ["51–100 (invisible)", "5,400", "Largest untapped opportunity"],
    ],
    widths=[34, 18, 48])
body("The 5,400 keywords sitting in positions 51–100 are the single biggest lever. Moving even 500 of them "
     "into the top 20 with targeted content could add 5,000–10,000 monthly visits — the kind of volume that "
     "an AI content factory (Section 12, Pillar 1) is purpose-built to produce at scale.")
secbreak()

# ================================================================ 6 PAID
h1("6.", "Paid Search (PPC) — A Critical Underinvestment")
body("This is the most immediately actionable gap in the audit. SC Fuels runs ~215 paid visits/mo on just "
     "2 keywords (both brand variants) — defensively appropriate, but absent from the commercial non-branded "
     "terms where new-customer acquisition actually happens.")
data_table(
    ["Metric", "Value"],
    [
        ["Monthly Paid Traffic", "215 visits"],
        ["Active Paid Keywords", "2 (both variants of “sc fuels near me”)"],
        ["Est. Monthly Paid Spend", "~$31 (per Semrush — minimal)"],
        ["Active Competitor (PPC)", "fuellogic.net — 195 paid keywords"],
        ["Ad Quality", "Both ads Top 3 — good execution, wrong scale"],
    ],
    widths=[36, 64])
body("Given the 15-state footprint and the LTV of a fleet account, a disciplined paid program targeting "
     "commercial terms (“commercial diesel delivery [city]”, “fleet fueling company”, “wholesale diesel "
     "supplier”) could realistically produce qualified enterprise leads at a cost-per-lead well inside an "
     "acceptable CAC for accounts worth $100K+ annually. Technijian’s My SEO program includes paid-ads "
     "management (Google + Meta) — this can run alongside or in place of WebFX’s PPC.")
secbreak()

# ================================================================ 7 BACKLINKS
h1("7.", "Backlink Profile & Domain Authority")
data_table(
    ["Metric", "Value"],
    [
        ["Total Backlinks", "5,100"],
        ["Referring Domains", "986 unique domains"],
        ["Follow Links", "2,900 (57.5%) — pass ranking authority"],
        ["Empty-Anchor Links", "31% (1,452 links) — missed keyword signal"],
        ["Top Referring Domain", "downsenergy.com (1,352) — acquired-brand link equity"],
    ],
    widths=[36, 64])
body("A healthy base. Two opportunities: (1) a link-building outreach program that earns descriptive, "
     "keyword-rich anchors (“commercial fuel delivery California”) rather than bare URLs; and (2) confirm "
     "301 redirects consolidate all http:// and non-www variants into https://www.scfuels.com/ so link "
     "equity isn’t split. Technijian’s My SEO add-ons (Content Syndication, PR Releases) build authority "
     "links systematically.")
secbreak()

# ================================================================ 8 COMPETITIVE
h1("8.", "Competitive Landscape")
h2("8.1  Organic Competitors")
data_table(
    ["Competitor", "Shared / SE Keywords", "Level"],
    [
        ["gosenergy.com", "494 shared / 8,300 SE", "Strongest organic rival"],
        ["ricochetfuel.com", "326 shared / 5,800 SE", "Strong in fuel delivery"],
        ["dieseldirect.com", "266 shared / 2,800 SE", "Diesel specialist"],
        ["boosterusa.com", "229 shared / 2,500 SE", "Mobile-fueling, digital-native"],
        ["walthall-oil.com", "205 shared / 2,200 SE", "Regional operator"],
    ],
    widths=[30, 38, 32])
body("SC Fuels holds the dominant keyword footprint and the highest organic traffic in this group. That lead "
     "should be defended and extended with consistent content and link acquisition — the area an AI content "
     "engine accelerates most.")

h2("8.2  The competitive lens Derek added — card networks")
callout("WEX and Voyager are competitors too", [
    "Derek’s point on June 2: “We compete against other SC Fuels-type companies, and we compete against the "
    "networks — WEX or Voyager. Even though we offer co-branded Voyager cards, we still compete against them "
    "for larger-volume customers.”",
    ("Implication.  ", "The real battleground is not only other distributors — it is the fuel-card networks "
     "courting your highest-volume fleets. Content and outbound should speak to why a managed SC Fuels "
     "relationship (local delivery, cardlock + onsite + lubricants + DEF under one roof, IFTA tooling) beats "
     "a pure card network for a serious fleet."),
], accent=ORANGE)

h2("8.3  AI Overview Presence — already real, under-optimized")
data_table(
    ["Keyword (AI Overview)", "Position", "Est. Visits/mo"],
    [
        ["diesel supplier", "Position 2", "1,604"],
        ["delivering fuel", "Position 1", "1,094"],
        ["diesel fuel distributors near me", "Position 5", "849"],
        ["commercial fuel", "Position 1", "64"],
    ],
    widths=[50, 24, 26])
body("SC Fuels already receives ~4,600 monthly visits from Google AI Overviews — with a GEO grade of F. "
     "Proper AI-search optimization could multiply that several times over. Whoever fixes GEO first locks in "
     "citation positions that are hard to displace later.")
secbreak()

# ================================================================ 9 TECHNICAL
h1("9.", "Technical SEO & On-Page Audit")
h2("9.1  On-Page SEO — Grade A")
body("Core on-page elements are well implemented (Yoast configured, good title/meta, valid canonical, SSL). "
     "The notable exceptions below are quick, high-value fixes.")
data_table(
    ["Check", "Status", "Finding / Recommendation"],
    [
        ["Title / Meta / Canonical", "PASS", "Within optimal ranges; canonical set correctly"],
        ["Single H1", "WARN", "Two H1 tags on homepage — standardize to one"],
        ["Image Alt Text", "WARN", "4 of 19 images missing alt attributes — simple fix"],
        ["Local Business Schema", "CRITICAL", "No Local Business schema on homepage — major local-SEO gap"],
        ["Address in static HTML", "CRITICAL", "Phone shown, but address not in static HTML — NAP failure"],
        ["Plain-text email exposed", "FAIL", "info@scfuels.com in raw HTML — replace with form"],
        ["SPF email record", "FAIL", "No SPF DNS record — spoofing & deliverability risk"],
        ["HTTP/2", "FAIL", "Serving on HTTP/1 — Nginx supports HTTP/2; config change"],
        ["GBP", "PASS", "Verified, 4.5 stars, 20 reviews — Orange CA"],
        ["YouTube channel", "WARN", "Linked but 4 subscribers / 526 views — underutilized"],
    ],
    widths=[28, 16, 56], status_col=1, font_sz=9.2)
secbreak()

# ================================================================ 10 GEO
h1("10.", "Generative Engine Optimization (GEO) — Grade F")
body("GEO is the practice of making sure LLMs and AI search (Google AI Overviews, ChatGPT, Perplexity, "
     "Copilot, Gemini) can read, understand and cite your content. In 2026 this is an active traffic channel, "
     "not a future one — and it is the single area where Technijian’s AI work is most differentiated.")
h2("10.1  Why this is urgent for SC Fuels specifically")
body("Your ideal buyer — a fleet manager, operations director, or procurement officer — increasingly asks "
     "an AI assistant “what should I look for in a commercial fuel supplier?” or “who are the best bulk "
     "diesel companies in the western US?” before engaging anyone. If your content isn’t readable by the "
     "model, you are invisible in that answer. Derek named exactly this — the younger buyer who starts with "
     "AI and social, not a phone call.")
h2("10.2  GEO Failure Analysis")
data_table(
    ["Check", "Status", "Finding"],
    [
        ["Dynamic rendering (89%)", "CRITICAL", "89% of page content is JS-rendered — LLMs read raw HTML and miss it"],
        ["llms.txt file", "CRITICAL", "No /llms.txt — the standard guidance file for AI crawlers is absent"],
        ["Static service content", "CRITICAL", "Core service copy & CTAs are JS-injected — invisible to AI crawlers"],
        ["FAQ schema", "FAIL", "No FAQ-structured content — heavily cited by AI Overviews"],
        ["Entity clarity", "FAIL", "Company entity not explicitly structured for AI comprehension"],
        ["Authority content", "WARN", "Limited long-form on fleet/IFTA/DEF — the topics LLMs cite"],
    ],
    widths=[26, 16, 58], status_col=1, font_sz=9.2)
body("The remediation — llms.txt, static-HTML migration of business-critical copy, FAQ schema, and an "
     "authority-content program — is detailed in Section 12, Pillar 1, and is core to the budgeted website "
     "rebuild Derek described.")
secbreak()

# ================================================================ 11 PERFORMANCE
h1("11.", "Site Performance — Grade D")
data_table(
    ["Metric", "Value", "Target"],
    [
        ["Content loaded", "7.9 s", "< 3 s"],
        ["Total page size", "6.10 MB", "< 5 MB"],
        ["JavaScript", "2.4 MB / 62 files", "Bundle + defer"],
        ["CSS compression", "7% compressed", "70%+"],
        ["Images", "1.4 MB / 27% compressed", "WebP + lazy-load"],
        ["HTTP protocol", "HTTP/1", "HTTP/2 (Nginx supports it)"],
    ],
    widths=[34, 38, 28])
body("Server response is excellent (0.1s) — the weight is front-end. The fixes are mostly configuration, not "
     "redevelopment: enable HTTP/2 (15–30% faster), turn on Brotli compression (CSS 1.86MB → ~0.5MB), and "
     "run images through a WebP/CDN pipeline (1.4MB → under 400KB). These pair naturally with the website "
     "rebuild and are owned by Technijian in the operating model (Section 13).")
secbreak()

# ================================================================ 12 THE ENGINE (4 pillars)
h1("12.", "How Technijian’s AI-Driven Engine Gets SC Fuels There")
body("This is the heart of Revision 2. The engine has four pillars. Pillars 1–2 strengthen what brings buyers "
     "in (and overlap with what WebFX does — we can support or run it). Pillars 3–4 are net-new: the outbound "
     "and internal AI that no SEO agency provides, and that Derek’s goals actually require. Each pillar lists "
     "what Technijian has already built, so this is proven capability — not theory.")

# capability proof up front
callout("Proven build — Multi-Agent SEO & Content Platform", [
    ("What we built.  ", "A multi-agent platform orchestrating Claude, GPT-4o and Gemini with MCP servers for "
     "SEMrush, GA4, Perplexity, Firecrawl and social distribution — ~70% less content-production time."),
    ("How it applies.  ", "This is the factory behind Pillars 1–2: blogs, AI-search content, videos and social "
     "produced at scale while the team manages quality, not keystrokes."),
], accent=BLUE)
callout("Proven build — AI Document Intelligence (RFPs: days → minutes)", [
    ("What we built.  ", "AI that auto-populates complex vendor questionnaires for FINRA broker-dealers — RFP "
     "responses cut from days to minutes, 60–80% less manual review."),
    ("How it applies.  ", "This is the engine behind the government/RFP play (Pillar 3) and the credit-process "
     "acceleration (Pillar 4) — the same document-AI, pointed at fuel bids and credit packets."),
], accent=ORANGE)

# ---- Pillar 1
h2("Pillar 1 — AI-Driven SEO & GEO (get found, by people and by AI)")
body("Fix the audit gaps and make SC Fuels the cited answer in both Google and AI search.")
bullet("Author and maintain /llms.txt and migrate business-critical copy to static HTML so AI crawlers can "
       "read it — directly fixing the Grade-F GEO.", bold_lead="AI-readability.  ")
bullet("FAQ-schema pages and Q&A-style content for the questions buyers actually ask (“how does cardlock "
       "fueling work?”, “what states does SC Fuels serve?”) — the format AI Overviews cite.", bold_lead="AI-search content.  ")
bullet("Target the 5,400 near-miss keywords and the service-area pages for 15 states with AI-accelerated, "
       "human-reviewed content — grow non-branded traffic past 50%.", bold_lead="Scale content.  ")
bullet("Maps to My SEO (Tiers 2–3) + the AI Search Optimization add-on.", bold_lead="Technijian service.  ", color=GREY)

# ---- Pillar 2
h2("Pillar 2 — AI-Driven Content & Social (reach the younger buyer)")
body("Answer Derek’s generational-shift question directly — meet buyers on social, video and audio.")
bullet("One source blog becomes a video (AI presenter), short social clips, a two-person podcast, and "
       "syndicated posts — across YouTube, LinkedIn, TikTok, Instagram, Facebook, Spotify and Apple.",
       bold_lead="Blog → video → social → podcast factory.  ")
bullet("PR + syndication to authority outlets builds backlink reputation; quarterly announcements get "
       "distributed for real reach.", bold_lead="Authority & syndication.  ")
bullet("Paid Google + Meta management runs the commercial-keyword PPC the audit calls for (Section 6).",
       bold_lead="Paid amplification.  ")
bullet("Maps to My SEO (Tiers 4–5: videos + viral clips) + Content Syndication + PR Releases + DMA paid-ads.",
       bold_lead="Technijian service.  ", color=GREY)

# ---- Pillar 3
h2("Pillar 3 — AI Lead Generation / Outbound (warmer leads, no new headcount)")
body("This is net-new — WebFX does not do it. It is how you “add business without adding salespeople.”")
bullet("AI scans public sources, scores prospects against your ICP, and sorts them into hot / warm / watch "
       "lists — the same pattern that produced 50 hot leads for another Technijian client from city plan-check "
       "data. For SC Fuels it augments FleetSeek and InsightView rather than replacing them.",
       bold_lead="ICP scoring.  ")
bullet("Each rep gets customized, one-to-one outreach drafts (not blast email) — the engine personalizes per "
       "company and reason-to-buy, at a volume a person can’t match by hand.", bold_lead="Personalized outreach.  ")
bullet("AI discovers relevant municipal/government RFPs and auto-drafts responses — letting SC Fuels compete "
       "on more bids with the same team, and chip at the incumbent-favored RFP problem Derek named.",
       bold_lead="Government / RFP automation.  ")
bullet("Maps to My AI Lead Gen (Starter / Professional / Enterprise). The CRM-sync tiers push scored leads "
       "straight into Salesforce.", bold_lead="Technijian service.  ", color=GREY)

# ---- Pillar 4
h2("Pillar 4 — AI Consulting & Internal AI (multiply the salesperson)")
body("Give each rep “ten assistants” so one person covers 2–4× the accounts — and match Pilot’s 27 → 3 day "
     "onboarding win.")
bullet("AI that watches Salesforce and pushes opportunities through the path — next-best-action, stalled-deal "
       "alerts, auto-summaries — tying InsightView + FleetSeek + Salesforce together (Derek’s explicit ask).",
       bold_lead="Salesforce opportunity progression.  ")
bullet("AI-assisted credit intake and decision support — faster, more consistent yes/no, less manual review "
       "(the tool Derek said he’d love).", bold_lead="Credit-process acceleration.  ")
bullet("An AI strategy roadmap, executive workshop, and a fractional AI advisor so the wins compound and the "
       "team is enabled, not bypassed.", bold_lead="Advisory.  ")
bullet("Maps to My AI (Strategy, Executive Workshop, Fractional Advisor) + My Dev for custom builds (new "
       "AI-ready website, credit & CRM automation). Scoped at discovery.", bold_lead="Technijian service.  ", color=GREY)

spacer(4)
body("Together, these four pillars are the engine. The next section shows how it runs alongside WebFX, and "
     "Section 14 shows what it costs.", italic=True, color=GREY)
secbreak()

# ================================================================ 13 OPERATING MODEL
h1("13.", "How It Works With WebFX — and the Optional Path to Consolidate")
body("The recommended model keeps WebFX where it is strong and gives Technijian the AI layers it is built "
     "for. Clear lanes, no overlap, no turf war.")
data_table(
    ["Layer", "WebFX (today)", "Technijian (adds)"],
    [
        ["SEO strategy & content", "Keyword targeting, content calendar, on-page", "AI content factory at scale; AI-search (GEO) content"],
        ["Paid search", "Current PPC", "AI-managed Google + Meta (or take it over)"],
        ["AI search / GEO", "Emerging", "llms.txt, static-HTML migration, FAQ schema, AI-crawler testing"],
        ["Outbound lead gen", "Not in scope", "AI prospecting, ICP scoring, RFP auto-draft (Pillar 3)"],
        ["Internal AI", "Not in scope", "Salesforce progression, credit AI, knowledge systems (Pillar 4)"],
        ["Site performance & security", "Recommendations", "HTTP/2, compression, CDN, SPF/DKIM/DMARC, monitoring"],
        ["Website rebuild", "Design input", "Build it AI-ready & GEO-ready from day one (My Dev)"],
        ["Reporting", "MarketingCloudFX attribution", "AI/infra dashboards; lead-engine + automation metrics"],
    ],
    widths=[22, 36, 42], font_sz=9)
callout("Complement now — consolidate later, if you choose", [
    "Start by adding the AI layers above on top of your current setup — nothing to unwind, fast to show value. "
    "Then, once the engine is proving out, SC Fuels has the option to bring the SEO and site work under "
    "Technijian’s fixed-fee My SEO model (12-month, unlimited hours) at a published cost that typically sits "
    "below enterprise-agency rates. That is a choice to make later with real data — not a switch we are asking "
    "for today. The first step is confirming what WebFX costs you now versus the value you get.",
], accent=BLUE)
secbreak()

# ================================================================ 14 INVESTMENT
h1("14.", "Investment — and How It Compares")
body("Pricing discipline: every Technijian figure below is a published rate. Where a number depends on scope "
     "(custom builds, the website rebuild, internal-AI automation), we say “scoped at discovery” rather than "
     "guess. We have not assumed your current WebFX spend — confirming it is a discovery item.")

h2("14.1  Technijian published rates")
data_table(
    ["Service", "What it covers", "Published rate"],
    [
        ["My SEO (Tiers 1–5)", "Site/hosting → SEO → blog content → video → viral video (cumulative, 12-mo, unlimited hours)", "$500 – $1,500 / mo"],
        ["My SEO add-ons", "AI Search Optimization (GEO) $200 · Content Syndication $200 · PR Releases $150", "+$150 – $200 / mo ea."],
        ["My SEO — paid ads (DMA)", "Google + Meta management (media spend separate, paid to the platforms)", "$100 – $300 / mo"],
        ["My AI Lead Gen — Starter", "1 pipeline, 500 enriched leads/mo, CSV + email drafts", "$1,499 / mo"],
        ["My AI Lead Gen — Professional", "3 pipelines, 2,500 leads/mo, outreach automation, CRM sync", "$3,499 / mo"],
        ["My AI Lead Gen — Enterprise", "Unlimited pipelines, 10,000+ leads/mo, dedicated CRM sync", "$6,999 / mo"],
        ["My AI (consulting)", "AI roadmap, executive workshop, fractional advisor", "Scoped at discovery"],
        ["My Dev (build)", "AI-ready website rebuild, Salesforce/credit automation", "Scoped at discovery"],
    ],
    widths=[24, 52, 24], font_sz=9)

h2("14.2  How that compares to a full-service agency")
data_table(
    ["", "WebFX (public ranges)", "Technijian (published)"],
    [
        ["SEO / content", "from ~$2,500/mo; core $3,000+/mo min", "My SEO $500–$1,500/mo, unlimited hours"],
        ["Paid search", "Included in larger retainers", "DMA $100–$300/mo + media spend"],
        ["AI search / GEO", "Emerging add-on", "AI Search Opt $200/mo"],
        ["AI outbound lead gen", "Not offered", "My AI Lead Gen $1,499–$6,999/mo"],
        ["Internal AI / CRM / credit", "Not offered", "My AI + My Dev (scoped)"],
        ["Enterprise engagement", "$10K–$50K+/mo, 6–12 mo terms + setup", "Fixed-fee, 12-mo; outbound is net-new scope"],
    ],
    widths=[26, 37, 37], font_sz=9)
callout("Reading the comparison fairly", [
    ("The headline.  ", "At published rates, a full inbound + paid + AI-search program via My SEO and its "
     "add-ons runs in the low thousands per month — and you can add the AI outbound engine WebFX does not "
     "offer (My AI Lead Gen) and still sit at or below the enterprise-agency range, while getting more scope."),
    ("The honest caveat.  ", "Enterprise WebFX engagements include strategy depth and MarketingCloudFX "
     "attribution that have real value. The right comparison is total scope for total dollars — which is why "
     "the first move is confirming what you pay WebFX today against what you receive."),
], accent=PASS, fill="EEF7EE")

h2("14.3  Land-and-expand — sequence the investment")
data_table(
    ["Phase", "What", "Approx. economics"],
    [
        ["Entry (now)", "AI Lead Gen + AI Search/GEO + content factory layered on the current setup — no big build", "Published recurring rates"],
        ["Expansion", "AI-ready website rebuild (already budgeted) + Salesforce/credit AI", "Scoped at discovery (My Dev/My AI)"],
        ["Consolidate (optional)", "Move SEO/site to Technijian fixed-fee My SEO once the engine proves out", "Typically below enterprise-agency rates"],
    ],
    widths=[22, 56, 22], font_sz=9)
h2("14.4  The ROI logic")
body("We model ROI as method, not a promised multiple. A single new enterprise fleet account is worth $100K+ "
     "in annual revenue. The entire entry-phase program runs in the low thousands per month. The engine "
     "therefore only needs to help land a small number of incremental accounts per year — or retain a few "
     "at-risk ones, or free enough rep time to cover more territory — to return many times its cost. Closed-loop "
     "attribution (WebFX’s MarketingCloudFX today, plus Technijian’s lead-engine metrics) keeps that honest.")
secbreak()

# ================================================================ 15 ROADMAP
h1("15.", "Prioritized Roadmap")
h2("15.1  Immediate — within 30 days (quick technical wins + engine kickoff)")
data_table(
    ["Action", "Owner", "Type"],
    [
        ["Enable HTTP/2 on Nginx", "Technijian", "QUICK WIN"],
        ["Implement SPF DNS record", "Technijian", "URGENT"],
        ["Create /llms.txt", "Technijian", "QUICK WIN"],
        ["Fix duplicate H1 + missing alt text", "WebFX/Dev", "QUICK WIN"],
        ["Replace plain-text email with form", "Technijian", "QUICK WIN"],
        ["Stand up My AI Lead Gen pilot (1 vertical)", "Technijian", "NEW"],
    ],
    widths=[52, 24, 24], status_col=2, font_sz=9.3)
h2("15.2  Build — 60–90 days (the engine + the rebuild)")
data_table(
    ["Action", "Owner", "Type"],
    [
        ["JS-render audit + static-HTML migration (GEO core)", "Technijian", "HIGH VALUE"],
        ["Brotli compression + WebP/CDN image pipeline", "Technijian", "HIGH VALUE"],
        ["Local Business schema + address in static HTML", "WebFX/Dev", "HIGH VALUE"],
        ["Expand PPC to 20–30 commercial keywords", "Technijian/WebFX", "HIGH VALUE"],
        ["Content + social factory live (blog→video→social)", "Technijian", "NEW"],
        ["Begin AI-ready website rebuild", "Technijian/Dev", "STRATEGIC"],
    ],
    widths=[52, 24, 24], status_col=2, font_sz=9.3)
h2("15.3  Scale — 3–6 months (internal AI + authority)")
data_table(
    ["Action", "Owner", "Type"],
    [
        ["Salesforce opportunity-progression AI", "Technijian", "STRATEGIC"],
        ["Credit-process AI acceleration", "Technijian", "STRATEGIC"],
        ["Government/RFP discovery + auto-draft", "Technijian", "STRATEGIC"],
        ["Service-area pages for 15 states", "Technijian/WebFX", "STRATEGIC"],
        ["Grow non-branded traffic to 50%+", "Technijian/WebFX", "STRATEGIC"],
        ["Fractional AI advisor cadence", "Technijian", "ONGOING"],
    ],
    widths=[52, 24, 24], status_col=2, font_sz=9.3)
secbreak()

# ================================================================ 16 QUICK WINS
h1("16.", "Quick Wins — No Commitment Required")
body("Five things SC Fuels can do this week, with or without an engagement:")
bullet("Run a free Nexus Assess. Technijian’s no-cost IT & security assessment (internal + external "
       "vulnerability, dark-web credential check, and a Microsoft 365 review) returns a prioritized "
       "remediation roadmap — a concrete starting point with no commitment.", bold_lead="1.  ")
bullet("Add /llms.txt. Even a simple version starts making the brand legible to AI search today.", bold_lead="2.  ")
bullet("Ask WebFX three questions: what is our total monthly spend, what does it cover, and what is your "
       "plan for AI search (GEO) and AI-driven outbound? The answers frame the consolidation decision.", bold_lead="3.  ")
bullet("Claim the YouTube channel for real. Four subscribers today; it is the natural home for the video "
       "content that reaches younger buyers.", bold_lead="4.  ")
bullet("List your 10 best closed accounts. They become the ICP seed for the AI Lead Gen pilot and the basis "
       "for anonymized case-study content.", bold_lead="5.  ")
secbreak()

# ================================================================ 17 NEXT STEPS
h1("17.", "Conclusion & Next Steps")
body("SC Fuels enters mid-2026 with a genuinely strong digital foundation — dominant brand rankings, a real "
     "retention moat in the portals, and a capable SEO partner in WebFX. The audit also shows a clear pattern: "
     "the site reflects 2020-era practices in a 2026, AI-first search world, and there is no AI-driven outbound "
     "or internal automation yet. Both are squarely fixable, and both map directly to the goal Derek named — "
     "grow without simply adding salespeople.")
body("As promised on June 2, this revision delivers the second-opinion site analysis, the honest WebFX read, "
     "and the AI growth plan in one place.", bold_lead="What we committed to.  ")
callout("Recommended next step", [
    "Bring this to the working session with your number-two and your Director of Sales & Marketing, with their "
    "questions in hand. We will walk the team through the engine, answer everything, and agree a small entry "
    "phase (the AI Lead Gen pilot + the quick technical wins) that proves the lift before any larger build.",
    ("To schedule.  ", "Use “Book a meeting” in Ravi’s email signature to grab a time — and we can cover this "
     "plan plus the broader set of AI strategies Technijian is putting into place for itself and its clients."),
], accent=ORANGE)

spacer(6)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Technijian Technology Solutions"); set_run(r, 11, BLUE, True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("technology as a solution"); set_run(r, 9.5, GREY, italic=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("18 Technology Dr., Ste 141, Irvine, CA 92618  ·  949.379.8499  ·  technijian.com")
set_run(r, 9, GREY)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sources: Semrush Domain Overview · The HOTH SEO Audit · June 2, 2026 strategy session · published Technijian & WebFX rates")
set_run(r, 8, GREY, italic=True)

doc.save(OUT_DOCX)
print("SAVED:", OUT_DOCX)
print("Paragraphs:", len(doc.paragraphs), "Tables:", len(doc.tables))
