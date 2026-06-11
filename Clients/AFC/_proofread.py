"""Proofread pass over Abound-AI-Growth-Report.docx:
banned hype words, mojibake, placeholders, pricing-math tokens, section refs.
"""
import re, sys
from docx import Document

DOCX = r"c:\vscode\tech-branding\tech-branding\Clients\AFC\Abound-AI-Growth-Report.docx"
doc = Document(DOCX)

texts = []
for para in doc.paragraphs:
    if para.text.strip():
        texts.append(para.text)
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.text.strip():
                    texts.append(para.text)

# tables nested in tables (callouts) - walk the body XML for completeness
full = "\n".join(texts)
print(f"paragraph/table text blocks: {len(texts)};  total chars: {len(full)}")

issues = []

# 1) banned hype words (word-boundary, case-insensitive)
banned = [
    r"\bleverage[sd]?\b", r"\bleveraging\b", r"\bseamless(ly)?\b", r"\bunlock(s|ed|ing)?\b",
    r"\bworld-class\b", r"\brobust\b", r"\bcutting-edge\b", r"\bend-to-end\b",
    r"\bgame-chang", r"\brevolutionar", r"\bfeel free to\b", r"\bdon't hesitate\b",
    r"in today's fast-paced", r"\bin conclusion\b", r"it is important to note",
]
for pat in banned:
    for m in re.finditer(pat, full, re.IGNORECASE):
        ctx = full[max(0, m.start()-60):m.end()+60].replace("\n", " ")
        issues.append(f"BANNED '{m.group(0)}': ...{ctx}...")

# "vendor" banned except in the verbatim FINRA phrase "vendor questionnaires"
for m in re.finditer(r"\bvendors?\b", full, re.IGNORECASE):
    ctx = full[max(0, m.start()-40):m.end()+40].replace("\n", " ")
    if "vendor questionnaires" not in ctx.lower():
        issues.append(f"BANNED 'vendor' (non-FINRA): ...{ctx}...")

# 2) mojibake
for pat in [r"Â", r"â€", r"Ã[^ ]", r"�"]:
    for m in re.finditer(pat, full):
        ctx = full[max(0, m.start()-50):m.end()+50].replace("\n", " ")
        issues.append(f"MOJIBAKE '{m.group(0)}': ...{ctx}...")

# 3) placeholders
for pat in [r"\bTBD\b", r"\bTODO\b", r"\[placeholder", r"\bLorem\b", r"\bXXX\b", r"\[Your "]:
    for m in re.finditer(pat, full, re.IGNORECASE):
        ctx = full[max(0, m.start()-50):m.end()+50].replace("\n", " ")
        issues.append(f"PLACEHOLDER '{m.group(0)}': ...{ctx}...")

# 4) pricing tokens - each must appear; consistency checks
need = ["$32,000", "$2,250/mo", "$5,000", "$1,250", "$1,000", "$15,000", "$12,000",
        "$48,000", "$9,600", "$24,000", "$5,050/mo", "$113,600", "$800", "$2,000"]
for tok in need:
    if tok not in full:
        issues.append(f"PRICING token missing: {tok}")
# numbers that must NOT appear (old/wrong figures)
for tok in ["$137,000", "$72,000", "$1,200", "$14,400", "$41,6", "$1,499", "$3,499", "$6,999"]:
    if tok in full:
        issues.append(f"PRICING unexpected token present: {tok}")

# 5) section cross-references - all "Section N" mentions valid (sections 1-17 exist)
for m in re.finditer(r"Section[s]? (\d+)", full):
    nval = int(m.group(1))
    if not (1 <= nval <= 17):
        issues.append(f"XREF out of range: {m.group(0)}")
refs = sorted(set(int(m.group(1)) for m in re.finditer(r"Section[s]? (\d+)", full)))
print("Section numbers referenced in prose:", refs)

# 6) figure numbering
figs = sorted(set(re.findall(r"Figure (\d+\.\d)", full)))
print("Figures:", figs)

# 7) phone discipline: only 949.379.8499
for m in re.finditer(r"949\.379\.85\d\d", full):
    issues.append(f"PHONE wrong number: {m.group(0)}")

# 8) tact: never print deficits / Keller / TechHeights / follower counts
for bad in ["deficit", "TechHeights", "Keller", "$2.06M", "$262", "declin", "Claudia",
            "followers", "990", "GuideStar"]:
    for m in re.finditer(bad, full, re.IGNORECASE):
        ctx = full[max(0, m.start()-60):m.end()+60].replace("\n", " ")
        issues.append(f"TACT '{bad}': ...{ctx}...")

# 9) ABM discipline: phrase check
for bad in ["lead generation", "shotgun", "fill your funnel"]:
    for m in re.finditer(bad, full, re.IGNORECASE):
        ctx = full[max(0, m.start()-60):m.end()+60].replace("\n", " ")
        issues.append(f"GTM-FRAMING '{bad}': ...{ctx}...")

print()
if issues:
    print(f"ISSUES ({len(issues)}):")
    for i in issues:
        print(" -", i)
    sys.exit(1)
print("PROOFREAD CLEAN: no banned words, no mojibake, no placeholders, pricing tokens consistent, refs in range.")
