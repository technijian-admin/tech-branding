"""
Build OXPLive MCP Server — Client Approval Report (DOCX).
Run: py -3.12 build-approval-docx.py
Output: OXP-Live-MCP-Server-Approval-Report.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT_DIR  = Path(__file__).parent
OUT_DOCX = OUT_DIR / "OXP-Live-MCP-Server-Approval-Report.docx"

LOGO_PATH = Path(r"C:\VSCode\tech-branding\tech-branding\assets\logos\Technijian Logo - blue text.png")

# Brand colours
BLUE   = RGBColor(0x00, 0x6D, 0xB6)
ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
TEAL   = RGBColor(0x1E, 0xAA, 0xC8)
DARK   = RGBColor(0x1A, 0x2B, 0x3C)
LIGHT  = RGBColor(0xF5, 0xF7, 0xFA)
MUTED  = RGBColor(0x71, 0x80, 0x96)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
TEXT   = RGBColor(0x2D, 0x37, 0x48)
BORDER = RGBColor(0xD0, 0xD7, 0xE2)

DOMAINS = [
    (1,  "Auth & Session",                5,   4,   2,  "User authentication, token lifecycle, session management"),
    (2,  "Patient / Account",             18,  16,  16, "Core patient record — demographics, insurance, job creation, medical history, status, cancellation"),
    (3,  "Insurance Authorization",        5,   5,   5,  "Auth requests, line-item management, status tracking, approval workflow"),
    (4,  "Queue & Workflow",               7,   7,   6,  "13-stage workflow queue (Intake → Billing → Collections); job transitions, queue counts"),
    (5,  "Billing & Claims",               5,   5,   5,  "Fee schedules, billing search, claims processing, remittance management"),
    (6,  "Scheduling",                    10,   9,   9,  "Appointment creation, calendar search, therapist scheduling, 48-hour delivery alerts"),
    (7,  "Patient Experience",             7,   6,   5,  "PSAT surveys, post-delivery call-back tracking, satisfaction scores, follow-up"),
    (8,  "Collector Actions",              5,   5,   4,  "Daily collector call logs, AR action tracking, daily-total summaries"),
    (9,  "Products & Inventory",          10,  10,   9,  "Equipment catalog, SKU management, inventory levels, assignment to patients at delivery"),
    (10, "Orders & Purchasing",            3,   3,   3,  "Purchase-order listing, vendor ordering, reorder tracking"),
    (11, "Documents / Files",              4,   4,   3,  "Document metadata, HIPAA-compliant signed download URLs for CMN, EDF, PHI files"),
    (12, "Directory, Users & Reps",        9,   8,   6,  "Doctor/facility directory, sales rep management, referring physician lookup, staff directory"),
    (13, "Time Records (HR)",              9,   9,   6,  "Timecard entry, manager approval workflow, payroll-ready export"),
    (14, "Lookups (Generic)",              4,   4,   2,  "Shared reference data — insurance payer names, territories, status codes, dropdown values"),
    (15, "Config",                         3,   3,   2,  "System-wide configuration and company-settings retrieval"),
    (16, "Reporting",                      4,   4,   4,  "Billing collections, PSR, variance analysis, KPI dashboard data"),
    (17, "Audit & Activity Logs",          4,   4,   3,  "HIPAA audit trail — user access log, data-change history, compliance events"),
    (18, "Compliance / Scheduled",         2,   2,   2,  "Mass job-move operations and compliance batch processing"),
    (19, "Yes/No Module (PA/CMN/EDF)",     1,   1,   1,  "Three compliance gates: PA status, CMN receipt, EDF sign-off — block queue progression"),
    (20, "Dashboard",                      1,   1,   1,  "Real-time territory and user activity dashboard — session-aware, role-filtered"),
]
TOTAL_STPS      = sum(d[2] for d in DOMAINS)
TOTAL_ENDPOINTS = sum(d[3] for d in DOMAINS)
TOTAL_TOOLS     = sum(d[4] for d in DOMAINS)


# ─── XML helpers ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color: str = "D0D7E2", sz: int = 4) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(sz))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type())


def docx_break_type():
    from docx.enum.text import WD_BREAK
    return WD_BREAK.PAGE


def set_para_spacing(para, before: int = 0, after: int = 6, line: int = 276) -> None:
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def add_colored_rule(doc: Document, color: str = "006DB6", sz: int = 12) -> None:
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_para_spacing(para, before=0, after=0)


def heading1(doc: Document, text: str, num: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(f"{num}   {text}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = BLUE
    set_para_spacing(para, before=240, after=80)
    # bottom border
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "006DB6")
    pBdr.append(bottom)
    pPr.append(pBdr)


def heading2(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    set_para_spacing(para, before=160, after=60)


def body_para(doc: Document, text: str, indent: bool = False) -> None:
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if indent:
        para.paragraph_format.left_indent = Cm(0.5)
    para.runs[0].font.size = Pt(9.5)
    para.runs[0].font.color.rgb = TEXT
    set_para_spacing(para, before=40, after=60)


def bullet_para(doc: Document, text: str, bold_prefix: str = "") -> None:
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = para.add_run(bold_prefix + " ")
        r1.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = TEXT
        r2 = para.add_run(text)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = TEXT
    else:
        run = para.add_run(text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = TEXT
    set_para_spacing(para, before=20, after=20)


def make_table_header_row(table, cells: list[str], col_widths: list[float]) -> None:
    row = table.rows[0]
    for i, (cell_text, width) in enumerate(zip(cells, col_widths)):
        cell = row.cells[i]
        cell.width = Inches(width)
        set_cell_bg(cell, "006DB6")
        set_cell_borders(cell, color="006DB6")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(cell_text)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = WHITE
        set_para_spacing(para, before=40, after=40)


def add_table_data_row(table, values: list[str], row_idx: int,
                       col_widths: list[float], bold_cols: list[int] = None) -> None:
    row = table.add_row()
    bg = "F5F7FA" if row_idx % 2 == 0 else "FFFFFF"
    for i, (val, width) in enumerate(zip(values, col_widths)):
        cell = row.cells[i]
        cell.width = Inches(width)
        set_cell_bg(cell, bg)
        set_cell_borders(cell, color="D0D7E2", sz=4)
        para = cell.paragraphs[0]
        run = para.add_run(val)
        run.font.size = Pt(8.5)
        run.font.color.rgb = TEXT
        if bold_cols and i in bold_cols:
            run.bold = True
        set_para_spacing(para, before=30, after=30)


# ─── Document sections ────────────────────────────────────────────────────────

def build_cover(doc: Document) -> None:
    """Cover page — logo, title, metadata, KPI summary."""
    # Logo
    if LOGO_PATH.exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run()
        run.add_picture(str(LOGO_PATH), height=Pt(28))
        set_para_spacing(para, before=0, after=120)
    else:
        para = doc.add_paragraph("TECHNIJIAN")
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = para.runs[0]
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = ORANGE
        set_para_spacing(para, before=0, after=120)

    # Eyebrow
    para = doc.add_paragraph("CLIENT APPROVAL REPORT")
    para.runs[0].font.size = Pt(8.5)
    para.runs[0].font.color.rgb = TEAL
    para.runs[0].bold = True
    set_para_spacing(para, before=0, after=60)

    # Title
    para = doc.add_paragraph()
    run = para.add_run("OXPLive MCP Server")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = DARK
    set_para_spacing(para, before=0, after=60)

    # Subtitle
    para = doc.add_paragraph("Secure AI Access to the OXPLive Platform")
    para.runs[0].font.size = Pt(13)
    para.runs[0].font.color.rgb = BLUE
    para.runs[0].bold = True
    set_para_spacing(para, before=0, after=200)

    # Horizontal rule
    add_colored_rule(doc)

    # Metadata table (2-col)
    meta = [
        ("Prepared for",     "OrthoKinetix (ORX)"),
        ("Prepared by",      "Technijian"),
        ("Project",          "OXPLive Modernization — Phase B Deliverable"),
        ("Date",             "June 2026"),
        ("Status",           "Pending Client Approval — Phase F (Build)"),
        ("Classification",   "Confidential — OrthoKinetix / Technijian Use Only"),
    ]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row repurposed as first data row
    row = tbl.rows[0]
    row.cells[0].width = Inches(1.5)
    row.cells[1].width = Inches(5.0)
    set_cell_bg(row.cells[0], "1A2B3C")
    set_cell_bg(row.cells[1], "1A2B3C")
    set_cell_borders(row.cells[0], "FFFFFF", sz=2)
    set_cell_borders(row.cells[1], "FFFFFF", sz=2)
    r0 = row.cells[0].paragraphs[0].add_run(meta[0][0])
    r0.bold = True; r0.font.color.rgb = WHITE; r0.font.size = Pt(9)
    r1 = row.cells[1].paragraphs[0].add_run(meta[0][1])
    r1.font.color.rgb = WHITE; r1.font.size = Pt(9)
    set_para_spacing(row.cells[0].paragraphs[0], before=30, after=30)
    set_para_spacing(row.cells[1].paragraphs[0], before=30, after=30)

    for key, val in meta[1:]:
        row = tbl.add_row()
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(5.0)
        set_cell_bg(row.cells[0], "1A2B3C")
        set_cell_bg(row.cells[1], "1A2B3C")
        set_cell_borders(row.cells[0], "FFFFFF", sz=2)
        set_cell_borders(row.cells[1], "FFFFFF", sz=2)
        rk = row.cells[0].paragraphs[0].add_run(key)
        rk.bold = True; rk.font.color.rgb = WHITE; rk.font.size = Pt(9)
        rv = row.cells[1].paragraphs[0].add_run(val)
        rv.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC); rv.font.size = Pt(9)
        set_para_spacing(row.cells[0].paragraphs[0], before=30, after=30)
        set_para_spacing(row.cells[1].paragraphs[0], before=30, after=30)

    # Spacer
    sp = doc.add_paragraph()
    set_para_spacing(sp, before=0, after=120)

    # KPI summary row
    kpis = [
        (str(TOTAL_STPS),      "Stored Procedures"),
        (str(TOTAL_ENDPOINTS), "API Endpoints"),
        (str(TOTAL_TOOLS),     "MCP Tools"),
        ("20",                  "Functional Domains"),
        ("100%",                "DB Call Coverage"),
    ]
    kpi_tbl = doc.add_table(rows=2, cols=len(kpis))
    kpi_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (num, lbl) in enumerate(kpis):
        # num row
        cell_num = kpi_tbl.rows[0].cells[i]
        set_cell_bg(cell_num, "1A2B3C")
        set_cell_borders(cell_num, "FFFFFF", sz=4)
        p = cell_num.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn = p.add_run(num)
        rn.bold = True; rn.font.size = Pt(20); rn.font.color.rgb = TEAL
        set_para_spacing(p, before=60, after=20)
        # label row
        cell_lbl = kpi_tbl.rows[1].cells[i]
        set_cell_bg(cell_lbl, "1A2B3C")
        set_cell_borders(cell_lbl, "FFFFFF", sz=4)
        pl = cell_lbl.paragraphs[0]
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = pl.add_run(lbl)
        rl.font.size = Pt(7.5); rl.font.color.rgb = RGBColor(0xAA, 0xBB, 0xCC)
        set_para_spacing(pl, before=0, after=60)

    # Cover footer line
    sp2 = doc.add_paragraph()
    set_para_spacing(sp2, before=160, after=0)
    add_colored_rule(doc, color="F67D4B", sz=8)

    para_foot = doc.add_paragraph()
    r1 = para_foot.add_run("Technijian · technology as a solution     ")
    r1.font.size = Pt(7.5); r1.font.color.rgb = MUTED
    r2 = para_foot.add_run("949.379.8499 · technijian.com     ")
    r2.font.size = Pt(7.5); r2.font.color.rgb = MUTED
    r3 = para_foot.add_run("Confidential — OrthoKinetix / Technijian Only")
    r3.font.size = Pt(7.5); r3.font.color.rgb = MUTED
    set_para_spacing(para_foot, before=40, after=0)

    doc.add_page_break()


def section_exec_summary(doc: Document) -> None:
    heading1(doc, "Executive Summary", "01")
    body_para(doc,
        "OrthoKinetix operates the OXPLive practice management platform — a mature SQL Server system "
        "running the full DME order lifecycle: intake, benefits verification, insurance authorization, "
        "shipping, billing, re-billing, and collections across 13 workflow queues.")
    body_para(doc,
        f"Technijian proposes to make that data safely available to AI assistants (Claude) through a "
        f"purpose-built Model Context Protocol (MCP) server. The server connects Claude directly to the "
        f"OXPLive database via a set of {TOTAL_STPS} curated stored procedures across 20 functional "
        f"domains — without granting the AI any direct table access.")
    body_para(doc,
        "The governing rule: every AI request follows a single, auditable path — Claude calls a "
        "documented MCP tool → the tool executes an approved stored procedure → the stored procedure "
        "reads or writes the database. No exceptions. No inline SQL. No table-level permissions for "
        "the AI service account.")
    body_para(doc,
        f"This document presents the complete scope of what will be built, the {TOTAL_STPS}-procedure "
        f"coverage matrix, the security and HIPAA controls in place, the delivery plan, and the approval "
        f"request to begin Phase F (Build). The full engineering specification — including every stored "
        f"procedure, its parameters, return schema, API endpoint, and MCP tool definition — is provided "
        f"in the companion STP-API-MCP Coverage Matrix.")
    doc.add_page_break()


def section_current_state(doc: Document) -> None:
    heading1(doc, "Why Now — The Current-State Gap", "02")
    body_para(doc,
        "Technijian has been the active IT and development partner for OrthoKinetix since approximately "
        "2015. As part of the ongoing OXPLive Modernization project (migrating from ColdFusion to "
        ".NET 8 / React 18 under Technijian SDLC v6.0), a complete database-call inventory was "
        "conducted in early 2026.")
    heading2(doc, "What Was Found")
    body_para(doc,
        "The legacy OXPLive ColdFusion frontend makes hundreds of direct database calls across "
        "13 workflow queues. These calls are not governed — they bypass stored procedures in many places, "
        "embed business logic in the UI layer, and contain hardcoded credentials. This creates three "
        "gaps that block safe AI access today:")
    bullet_para(doc,
        "The application reaches the database through a mix of stored procedures, inline SQL, and "
        "direct table reads. An AI assistant granted the same access would have uncontrolled write "
        "access to production data.",
        bold_prefix="No single governed access path.")
    bullet_para(doc,
        "Direct database calls leave no application-level record of who accessed what, when, or why "
        "— a HIPAA audit requirement.",
        bold_prefix="No audit trail by design.")
    bullet_para(doc,
        "Without a procedure layer, there is no clean way to expose only what the AI should see "
        "while hiding raw PHI fields, financial data, and operational tables.",
        bold_prefix="No semantic boundary.")
    heading2(doc, "The Solution")
    body_para(doc,
        "The OXPLive Modernization addresses all three gaps by establishing a complete, SP-Only "
        "database access pattern. Every frontend interaction — and every AI tool call — must go "
        "through a named, documented, version-controlled stored procedure. This is both a "
        "modernization best practice and the prerequisite for safe AI integration.")
    doc.add_page_break()


def section_architecture(doc: Document) -> None:
    heading1(doc, "Architecture — DB-Direct Design", "03")
    body_para(doc,
        "The OXPLive MCP server uses a DB-Direct architecture: Claude connects to the MCP server "
        "via Streamable HTTP, and the server executes stored procedures directly against SQL Server "
        "via a secure connection pool. There is no intermediate REST API layer for MCP access.")

    # Architecture diagram as a table
    heading2(doc, "Request Path")
    arch_rows = [
        ("Claude (Desktop / Code / API)",     "AI assistant — your staff or automated workflows",       "F67D4B"),
        ("Streamable HTTP  POST /mcp",        "MCP 2025-03-26 spec · session-keyed · optional API-key", "006DB6"),
        (f"OXPLive MCP Server  Node.js + mssql", f"{TOTAL_TOOLS} MCP tools · dispatches to named stored procedures", "1EAAC8"),
        ("SQL Server — oxadmin schema",       "Service account EXECUTE-only on oxadmin schema · no table permissions", "1A2B3C"),
        ("OXPLive Tables · Views · Functions","Production OrthoXpressDB — never directly accessible to the service account", "718096"),
    ]
    arch_tbl = doc.add_table(rows=len(arch_rows), cols=2)
    arch_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (layer, note, color) in enumerate(arch_rows):
        c0 = arch_tbl.rows[i].cells[0]
        c1 = arch_tbl.rows[i].cells[1]
        c0.width = Inches(2.8)
        c1.width = Inches(3.8)
        set_cell_bg(c0, color)
        set_cell_borders(c0, "FFFFFF", sz=4)
        set_cell_borders(c1, "D0D7E2", sz=4)
        r = c0.paragraphs[0].add_run(layer)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE
        set_para_spacing(c0.paragraphs[0], before=50, after=50)
        rn = c1.paragraphs[0].add_run(note)
        rn.font.size = Pt(8.5); rn.font.color.rgb = MUTED; rn.italic = True
        set_para_spacing(c1.paragraphs[0], before=50, after=50)

    sp = doc.add_paragraph()
    set_para_spacing(sp, before=0, after=80)

    heading2(doc, "Why DB-Direct")
    body_para(doc,
        "The OXPLive Modernization Phase F will build a .NET 8 REST API, but the stored procedures "
        "are being designed now as part of Phase B. DB-Direct lets the MCP server go live the moment "
        "the stored procedures are tested — before the full API layer is complete — giving OrthoKinetix "
        "early AI access while the broader modernization continues.")
    body_para(doc,
        "Security posture is unchanged from an API-first design. The SQL Server service account "
        "(oxplive_mcp_svc) holds only EXECUTE on the oxadmin schema. It cannot read tables directly, "
        "run dynamic SQL, or access any object outside the curated procedure set.")

    heading2(doc, "Transport — Streamable HTTP")
    body_para(doc,
        "The server implements the MCP 2025-03-26 Streamable HTTP transport specification:")
    for method, path, desc in [
        ("POST",   "/mcp",    "Creates or resumes a session; handles all JSON-RPC tool calls"),
        ("GET",    "/mcp",    "SSE stream for server-initiated messages (session-keyed)"),
        ("DELETE", "/mcp",    "Gracefully closes a session"),
        ("GET",    "/health", "Returns { status: \"ok\", sessions: N }"),
    ]:
        para = doc.add_paragraph(style="List Bullet")
        r1 = para.add_run(f"{method} {path}"); r1.bold = True; r1.font.size = Pt(9.5)
        r1.font.color.rgb = BLUE
        r2 = para.add_run(f" — {desc}"); r2.font.size = Pt(9.5); r2.font.color.rgb = TEXT
        set_para_spacing(para, before=20, after=20)

    doc.add_page_break()


def section_domain_matrix(doc: Document) -> None:
    heading1(doc, "Domain Coverage Matrix", "04")
    body_para(doc,
        f"The table below summarises the complete stored-procedure coverage across all 20 functional "
        f"domains of the OXPLive system. Every domain corresponds to a distinct area of the OXPLive "
        f"workflow UI. The full {TOTAL_STPS}-row matrix — with individual SP parameters, return schemas, "
        f"API endpoint definitions, and MCP tool specifications — is provided in the companion "
        f"STP-API-MCP Coverage Matrix document.")

    headers = ["#", "Domain", "STPs", "API\nEndpoints", "MCP\nTools", "Business Purpose"]
    widths  = [0.25, 1.5, 0.5, 0.6, 0.55, 3.7]

    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    make_table_header_row(tbl, headers, widths)

    for i, d in enumerate(DOMAINS):
        num, name, stps, eps, tools, biz = d
        add_table_data_row(
            tbl,
            [str(num), name, str(stps), str(eps), str(tools), biz],
            i,
            widths,
            bold_cols=[1],
        )

    # Total row
    row = tbl.add_row()
    vals = ["", "TOTAL", str(TOTAL_STPS), str(TOTAL_ENDPOINTS), str(TOTAL_TOOLS),
            "100% of OXPLive database calls covered"]
    for i, (val, width) in enumerate(zip(vals, widths)):
        cell = row.cells[i]
        cell.width = Inches(width)
        set_cell_bg(cell, "1A2B3C")
        set_cell_borders(cell, "FFFFFF", sz=4)
        r = cell.paragraphs[0].add_run(val)
        r.bold = True; r.font.size = Pt(8.5)
        r.font.color.rgb = WHITE if i < 5 else RGBColor(0xAA, 0xBB, 0xCC)
        set_para_spacing(cell.paragraphs[0], before=40, after=40)

    sp = doc.add_paragraph()
    p = sp.add_run("STPs = Stored Procedures.  API Endpoints = planned .NET 8 REST endpoints (Phase F).  "
                   "MCP Tools = tools exposed to Claude via DB-Direct MCP server.  "
                   "Some STPs are INTERNAL and are not exposed as MCP tools.")
    p.font.size = Pt(8); p.font.color.rgb = MUTED; p.italic = True
    set_para_spacing(sp, before=40, after=0)

    doc.add_page_break()


def section_detailed_capability(doc: Document) -> None:
    """Section 05 — per-SP capability list grouped by domain."""
    from itertools import groupby

    STPS = [
        (1,"1.1","usp_Auth_EnsureUser","INT","","","INTERNAL — confirms/creates the Azure AD user on first login; called by API middleware only"),
        (1,"1.2","usp_Auth_GetCurrentUser","READ","GET","/api/me","Returns the logged-in user's name, email, roles, division, and territory to populate the app shell"),
        (1,"1.3","usp_Auth_GetIdpConfig","INT","","","INTERNAL — returns MSAL configuration for the SPA; agents use client_credentials instead"),
        (1,"1.4","usp_Auth_LogSessionHistory","INT","","","INTERNAL — records every page navigation for the Admin activity log; called by request pipeline"),
        (1,"1.5","usp_Auth_RefreshToken","INT","","","INTERNAL — silently rotates access/refresh token pair when the SPA token expires; MSAL-managed"),
        (2,"2.1","usp_Patient_Search","READ","GET","/api/patients/search","Find patients by name, account number, DOB, SSN-last-4, or phone"),
        (2,"2.2","usp_Patient_GetDetail","READ","GET","/api/patients/{jobId}","Retrieve complete patient record: demographics, insurance, medical info, referring doctor, status"),
        (2,"2.3","usp_Patient_Create","WRITE","POST","/api/patients","Create a new patient account with demographics, insurance, and medical intake data"),
        (2,"2.4","usp_Patient_Update","WRITE","PUT","/api/patients/{jobId}","Save edits to an existing patient record; changes are audited before/after"),
        (2,"2.5","usp_Patient_GetSnapshot","READ","GET","/api/patients/{jobId}/snapshot","Get a complete printable snapshot: all products, billing codes, insurance, clinical notes in one call"),
        (2,"2.6","usp_Patient_GetCoverage","READ","GET","/api/patients/{jobId}/coverage","Get payer name, policy number, group number, coverage type, and current auth numbers"),
        (2,"2.7","usp_Patient_UpdateCoverage","WRITE","PUT","/api/patients/{jobId}/coverage","Update primary or secondary insurance coverage without touching demographics"),
        (2,"2.8","usp_Patient_GetOrigRep","READ","GET","/api/patients/{jobId}/orig-rep","Get the original rep assigned at intake — preserves commission origination credit"),
        (2,"2.9","usp_Patient_SetOrigRep","WRITE","PUT","/api/patients/{jobId}/orig-rep","Set or update the originating rep when first assigned during intake"),
        (2,"2.10","usp_Patient_ListProducts","READ","GET","/api/patients/{jobId}/products","List all product/equipment line items with HCPCS codes, billing status, and auth notes"),
        (2,"2.11","usp_Patient_AddProduct","WRITE","POST","/api/patients/{jobId}/products","Add a new product line item to the account; validates HCPCS code mapping"),
        (2,"2.12","usp_Patient_UpdateProduct","WRITE","PUT","/api/patients/{jobId}/products/{id}","Update quantity or notes on a product line item"),
        (2,"2.13","usp_Patient_RemoveProduct","WRITE","DELETE","/api/patients/{jobId}/products/{id}","Remove a product line item; blocked if a claim has already been submitted"),
        (2,"2.14","usp_Patient_ListNotes","READ","GET","/api/patients/{jobId}/notes","Read all free-text notes on the account in chronological order"),
        (2,"2.15","usp_Patient_AddNote","WRITE","POST","/api/patients/{jobId}/notes","Append a note to the account; notes are immutable once created"),
        (2,"2.16","usp_Patient_GetAgreement","READ","GET","/api/patients/{jobId}/agreement","Check whether the patient has accepted the service agreement and on what date"),
        (2,"2.17","usp_Medical_UpdateDiagnosis","WRITE","PUT","/api/patients/{jobId}/medical/diagnosis","Update ICD-10 diagnosis codes, referring doctor, body part, and procedure code"),
        (2,"2.18","usp_Patient_Cancel","WRITE","POST","/api/patients/{jobId}/cancel","Cancel the account with a required reason code and notes; writes a final status history entry"),
        (3,"3.1","usp_Auth_ListByPatient","READ","GET","/api/patients/{jobId}/authorizations","List all insurance authorizations: auth numbers, dates, payer, status, linked products"),
        (3,"3.2","usp_Auth_Create","WRITE","POST","/api/patients/{jobId}/authorizations","Create a new insurance authorization record with auth number, dates, payer, and notes"),
        (3,"3.3","usp_Auth_Update","WRITE","PUT","/api/authorizations/{authId}","Update an existing authorization — correct dates, extend, or change status"),
        (3,"3.4","usp_Auth_AddItem","WRITE","POST","/api/authorizations/{authId}/items","Link a product line item to an authorization; required for billing to confirm each line is covered"),
        (3,"3.5","usp_Auth_RemoveItem","WRITE","DELETE","/api/authorizations/{authId}/items/{id}","Remove a product from an authorization; blocked if the claim has been submitted"),
        (4,"4.1","usp_Queue_List","READ","GET","/api/queues","List all 13 workflow queues visible to the current user with current job counts"),
        (4,"4.2","usp_Queue_GetJobs","READ","GET","/api/queues/{queueId}/jobs","Get all patient accounts in a specific queue — paged, with name, status age, and rep"),
        (4,"4.3","usp_Queue_TransitionJob","WRITE","POST","/api/queues/transition","Move a patient account from one queue to another; validates the transition, writes status history"),
        (4,"4.4","usp_Queue_GetStatusHistory","READ","GET","/api/patients/{jobId}/status-history","Get the complete queue-transition history for an account — every move with user and timestamp"),
        (4,"4.5","usp_Queue_GetActions","READ","GET","/api/queues/{queueId}/actions","List the permitted moves out of a queue — used to render the correct action buttons"),
        (4,"4.6","usp_Queue_GetBillingStatus","READ","GET","/api/patients/{jobId}/billing-status","Get billing status badge: Unbilled / Partially Billed / Fully Billed / Cash"),
        (4,"4.7","usp_Queue_GetCounts","READ","GET","/api/queues/counts","Get count of accounts in each of the 13 queues in one call — powers the sidebar badges and dashboard"),
        (5,"5.1","usp_Billing_ListEntities","READ","GET","/api/billing/entities","List all billing entities (insurance companies, payers)"),
        (5,"5.2","usp_Billing_Search","READ","GET","/api/billing/search","Search billable records by account, payer, or status — returns line items with charge and collection data"),
        (5,"5.3","usp_Billing_GetFeeSchedules","READ","GET","/api/billing/fee-schedules","Get contracted fee schedule amounts for a HCPCS code and billing entity"),
        (5,"5.4","usp_Billing_GetProductValuation","READ","GET","/api/billing/valuations/{productId}","Get pricing and rental term options (capped rental, rent-to-own, purchase)"),
        (5,"5.5","usp_Billing_GetReBillData","READ","GET","/api/patients/{jobId}/rebill","Get billed line items flagged for re-billing due to denial, partial payment, or correction"),
        (6,"6.1","usp_Schedule_GetByRange","READ","GET","/api/schedules","List appointments filtered by date range, territory, or rep — powers the calendar view"),
        (6,"6.2","usp_Schedule_Create","WRITE","POST","/api/schedules","Create a new appointment for a patient — type, date/time, assigned rep, and notes"),
        (6,"6.3","usp_Schedule_Update","WRITE","PUT","/api/schedules/{scheduleId}","Reschedule an appointment or update its type or notes"),
        (6,"6.4","usp_Schedule_Delete","WRITE","DELETE","/api/schedules/{scheduleId}","Cancel and remove a scheduled appointment; also removes Exchange calendar sync"),
        (6,"6.5","usp_Schedule_GetUnscheduled","READ","GET","/api/schedules/unscheduled","Find accounts in the scheduling queue with no upcoming appointment"),
        (6,"6.6","usp_Schedule_GetTypes","READ","GET","/api/lookups/schedule-types","Get appointment type dropdown (delivery, follow-up, fitting, re-evaluation, etc.)"),
        (6,"6.7","usp_Schedule_GetExchangeSync","READ","GET","/api/schedules/{scheduleId}/exchange-sync","Get the Outlook calendar sync status for an appointment"),
        (6,"6.8","usp_Schedule_SetExchangeSync","WRITE","PUT","/api/schedules/{scheduleId}/exchange-sync","Link an appointment to an Outlook calendar event after pushing to Exchange"),
        (6,"6.9","usp_Schedule_DeleteExchangeSync","WRITE","DELETE","/api/schedules/{scheduleId}/exchange-sync","Remove the Exchange calendar link when an appointment is cancelled"),
        (6,"6.10","usp_Scheduler_Get48HourAlerts","READ","GET","/api/schedules/alerts","Get deliveries and follow-ups due within 48 hours for the user's territory — same-day alert panel"),
        (7,"7.1","usp_PatientExperience_List","READ","GET","/api/patients/{jobId}/experiences","List all patient contacts on the account — calls, inquiries, complaints — chronologically"),
        (7,"7.2","usp_PatientExperience_Create","WRITE","POST","/api/patients/{jobId}/experiences","Log a patient contact: call attempt, inquiry, complaint, or scheduling contact with outcome"),
        (7,"7.3","usp_PatientExperience_Update","WRITE","PUT","/api/patient-experiences/{peId}","Update status, follow-up date, or notes on a patient contact record"),
        (7,"7.4","usp_PatientExperience_GetTypes","READ","GET","/api/lookups/patient-experience-types","Get contact type dropdown (inquiry, scheduling, complaint, satisfaction, etc.)"),
        (7,"7.5","usp_PatientExperience_GetStatuses","READ","GET","/api/lookups/patient-experience-statuses","Get contact status dropdown (open, resolved, escalated, etc.)"),
        (7,"7.6","usp_PatientExperience_GetComplaintTypes","READ","GET","/api/lookups/complaint-types","Get complaint type dropdown — used when logging a formal patient complaint"),
        (7,"7.7","usp_PatientExperience_GetTemplates","READ","GET","/api/lookups/patient-experience-templates","Get message templates that pre-populate contact notes for common contact types"),
        (8,"8.1","usp_CollectorAction_List","READ","GET","/api/collector-actions","List collector actions for an account or by rep — calls made, voicemails, payment arrangements"),
        (8,"8.2","usp_CollectorAction_Create","WRITE","POST","/api/collector-actions","Log a new collection activity: call, voicemail, payment plan, referral to agency"),
        (8,"8.3","usp_CollectorAction_Update","WRITE","PUT","/api/collector-actions/{actionId}","Update a collector action — revise follow-up date or add resolution notes"),
        (8,"8.4","usp_CollectorAction_GetReasonCodes","READ","GET","/api/lookups/collector-action-reasons","Get reason code dropdown (called/no answer, payment plan set, referred to agency, etc.)"),
        (8,"8.5","usp_CollectorAction_GetDailyTotals","READ","GET","/api/collector-actions/totals","Get aggregated collector performance stats by user and date range"),
        (9,"9.1","usp_Product_List","READ","GET","/api/products","Browse the product catalog filtered by category or active status"),
        (9,"9.2","usp_Product_Create","WRITE","POST","/api/products","Add a new product to the catalog with pricing and vendor info (admin only)"),
        (9,"9.3","usp_Product_Update","WRITE","PUT","/api/products/{productId}","Update product name, part number, pricing, or active status (admin only)"),
        (9,"9.4","usp_Product_GetBillingCodes","READ","GET","/api/products/{productId}/billing-codes","Get HCPCS billing code mappings for a product — used in billing validation"),
        (9,"9.5","usp_Product_GetBodyParts","READ","GET","/api/lookups/body-parts","Get body part dropdown (Left Knee, Right Ankle, etc.) for medical intake"),
        (9,"9.6","usp_Product_GetCategories","READ","GET","/api/lookups/product-categories","Get product category dropdown (AFO, KAFO, TLSO, etc.)"),
        (9,"9.7","usp_Product_GetManufacturers","READ","GET","/api/lookups/manufacturers","Get vendor/manufacturer dropdown for product assignment"),
        (9,"9.8","usp_Inventory_ListLocations","READ","GET","/api/inventory/locations","List warehouse/storage locations by site"),
        (9,"9.9","usp_Inventory_ListSites","READ","GET","/api/inventory/sites","List inventory sites (warehouses and offices)"),
        (9,"9.10","usp_Inventory_Assign","WRITE","PUT","/api/inventory/{equipmentId}/assign","Assign a specific piece of equipment to a patient — records serial number and date for stock-and-bill compliance"),
        (10,"10.1","usp_Order_ListByPatient","READ","GET","/api/patients/{jobId}/orders","List all orders and purchase orders on the account"),
        (10,"10.2","usp_Order_GetDetail","READ","GET","/api/orders/{orderId}","Get full order detail including all line items with SKU, description, quantity, and price"),
        (10,"10.3","usp_Purchasing_GetByPatient","READ","GET","/api/patients/{jobId}/purchasing","Get sales order number, PO number, and invoice number for the account"),
        (11,"11.1","usp_Document_ListByPatient","READ","GET","/api/patients/{jobId}/documents","List all documents attached to the account: type, filename, upload date, uploader"),
        (11,"11.2","usp_Document_GetMeta","INT","","","INTERNAL — retrieves blob path for signed URL generation; raw path never returned to client (HIPAA)"),
        (11,"11.3","usp_Document_CreateRecord","WRITE","POST","/api/patients/{jobId}/documents","Attach an uploaded file to the patient account; blob stored in Azure, metadata in database"),
        (11,"11.4","usp_Document_GetViewUrl","READ","GET","/api/documents/{documentId}/view","Get a 5-minute signed URL to view a patient document; access logged to HIPAA audit trail"),
        (12,"12.1","usp_Directory_GetList","READ","GET","/api/directory","Browse the staff directory filtered by division or territory"),
        (12,"12.2","usp_Directory_GetReps","READ","GET","/api/lookups/reps","Get the sales rep list, optionally filtered by territory"),
        (12,"12.3","usp_Directory_GetTerritories","READ","GET","/api/lookups/territories","Get territory lookup — used in scheduling, patient intake, and queue filtering"),
        (12,"12.4","usp_Directory_GetDivisions","READ","GET","/api/lookups/divisions","Get division/business-unit lookup"),
        (12,"12.5","usp_Directory_GetDoctors","READ","GET","/api/doctors","Browse the physician directory filtered by specialty"),
        (12,"12.6","usp_Directory_GetPractitioners","READ","GET","/api/practitioners","List clinical practitioners (PTs, OTs, RNs) for clinical assignment"),
        (12,"12.7","usp_User_ListAll","READ","GET","/api/admin/users","List all user accounts with roles, division, territory (admin only)"),
        (12,"12.8","usp_User_GetSecurity","READ","GET","/api/admin/users/{userId}/security","Get role assignments and territory access for a user (admin only)"),
        (12,"12.9","usp_User_GetRoleTypes","READ","GET","/api/admin/roles","Get all defined role types (Admin, Billing, Clinical, etc.)"),
        (13,"13.1","usp_TimeRecord_List","READ","GET","/api/time-records","List time record submissions by employee, pay period, or approval status"),
        (13,"13.2","usp_TimeRecord_GetDetail","READ","GET","/api/time-records/{timeRecordId}","Get a specific time record with all hourly or salaried line items"),
        (13,"13.3","usp_TimeRecord_Create","WRITE","POST","/api/time-records","Create a new time record header for a pay period"),
        (13,"13.4","usp_TimeRecord_Update","WRITE","PUT","/api/time-records/{timeRecordId}","Submit, approve, or reject a time record"),
        (13,"13.5","usp_TimeRecord_GetSalariedData","READ","GET","/api/time-records/{id}/salaried-data","Get daily salaried entries and PTO allocations for a salaried time record"),
        (13,"13.6","usp_TimeRecord_GetHourlyData","READ","GET","/api/time-records/{id}/hourly-data","Get hourly shift entries with start/end times and OT hours"),
        (13,"13.7","usp_TimeRecord_GetPTOTypes","READ","GET","/api/lookups/pto-types","Get PTO type dropdown (Vacation, Sick, Personal, Holiday, Bereavement)"),
        (13,"13.8","usp_TimeRecord_GetPaymentCalendar","READ","GET","/api/time-records/payment-calendar","Get pay period dates for a year — valid week ranges for timecard entry"),
        (13,"13.9","usp_TimeRecord_GetHolidayCalendar","READ","GET","/api/time-records/holiday-calendar","Get company holidays for a year — used in timecard calculations"),
        (14,"14.1","usp_Lookup_GetValues","READ","GET","/api/lookups/{tableName}","Get dropdown values for any lookup table: cancellation reasons, document types, regions, territories, etc."),
        (14,"14.2","usp_Lookup_Create","WRITE","POST","/api/admin/lookups/{tableName}","Add a new value to a lookup table (admin only); change is audit-logged"),
        (14,"14.3","usp_Lookup_Update","WRITE","PUT","/api/admin/lookups/{tableName}/{id}","Update a lookup table value (admin only); change is audit-logged"),
        (14,"14.4","usp_Lookup_Delete","INT","","","INTERNAL — soft-deletes a lookup value; admin UI action only, not an MCP tool (prevents agents from deleting reference data)"),
        (15,"15.1","usp_Config_GetIdp","INT","","","INTERNAL — MSAL configuration for the SPA; agents use client_credentials, not IDP config"),
        (15,"15.2","usp_Config_GetAll","READ","GET","/api/config","Get all system configuration values: session timeout, file size limits, feature flags"),
        (15,"15.3","usp_Config_GetCompany","READ","GET","/api/config/company","Get company name, address, and branding for display in reports and print views"),
        (16,"16.1","usp_Report_BillingCollections","READ","GET","/api/reports/billing-collections","Billing and collections report: outstanding balances, payer performance, aging buckets"),
        (16,"16.2","usp_Report_DepartmentDeliverables","READ","GET","/api/reports/department-deliverables","Department productivity: deliverables completed, time logged, productivity score by division"),
        (16,"16.3","usp_Report_PSR","READ","GET","/api/reports/psr","Patient Satisfaction report: contact outcomes, complaint resolution, rep performance"),
        (16,"16.4","usp_Report_Variance","READ","GET","/api/reports/variance","Pricing variance: billed amounts vs contracted rates — surfaces underbilling and margin issues"),
        (17,"17.1","usp_AuditLog_Insert","INT","","","INTERNAL — writes HIPAA audit events; called by API middleware only"),
        (17,"17.2","usp_AuditLog_GetList","READ","GET","/api/admin/audit-logs","Search the HIPAA audit log: PHI access, data changes, security events (admin/auditor only)"),
        (17,"17.3","usp_ActivityLog_Insert","INT","","","INTERNAL — records page views, logins, logouts; called by request pipeline only"),
        (17,"17.4","usp_ActivityLog_GetList","READ","GET","/api/admin/activity-logs","Search user activity log: login, logout, and page-view events (admin only)"),
        (18,"18.1","usp_Compliance_GetPendingJobs","READ","GET","/api/admin/compliance/pending","List accounts awaiting compliance review that are overdue — daily compliance sweep"),
        (18,"18.2","usp_Compliance_MassMoveJobs","WRITE","POST","/api/admin/compliance/mass-move","Bulk-move accounts past the compliance hold period to the next queue; returns count moved"),
        (19,"19.1","usp_YesNo_Update","WRITE","PUT","/api/patients/{jobId}/yesno","Update all three compliance gates at once: PA authorization, CMN receipt, EDF sign-off. PA Denied blocks the account from advancing to Billing"),
        (20,"20.1","usp_Dashboard_GetActivity","READ","GET","/api/dashboard/activity","Get six headline KPIs: new intakes today, queue transitions, deliveries, documents uploaded, open collector follow-ups, overdue compliance — territory-filtered"),
    ]

    DNAMES = {
        1:"Auth & Session", 2:"Patient / Account", 3:"Insurance Authorization",
        4:"Queue & Workflow", 5:"Billing & Claims", 6:"Scheduling",
        7:"Patient Experience (PSR)", 8:"Collector Actions", 9:"Products & Inventory",
        10:"Orders & Purchasing", 11:"Documents / Files", 12:"Directory, Users & Reps",
        13:"Time Records (HR)", 14:"Lookups (Generic)", 15:"Config",
        16:"Reporting", 17:"Audit & Activity Logs", 18:"Compliance / Scheduled",
        19:"Yes/No Module (PA/CMN/EDF)", 20:"Dashboard",
    }

    heading1(doc, "Detailed Capability Reference", "05")
    read_count  = sum(1 for s in STPS if s[3]=="READ")
    write_count = sum(1 for s in STPS if s[3]=="WRITE")
    int_count   = len(STPS) - read_count - write_count
    body_para(doc,
        f"Every stored procedure in the OXPLive MCP server is listed below, grouped by domain. "
        f"Each row shows whether the operation reads data or writes/updates it, the API endpoint it "
        f"maps to in Phase F, and a plain-English description of what OrthoKinetix staff and AI "
        f"assistants can do with it. In total: {read_count} read operations, {write_count} write "
        f"operations, and {int_count} internal procedures not exposed to clients.")

    from itertools import groupby
    for domain_num, rows in groupby(STPS, key=lambda s: s[0]):
        rows = list(rows)
        heading2(doc, f"Domain {domain_num} — {DNAMES[domain_num]}")

        headers = ["#", "Stored Procedure", "R/W", "API Endpoint", "What It Does"]
        widths  = [0.3, 1.6, 0.5, 2.2, 3.0]
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        make_table_header_row(tbl, headers, widths)

        for i, (_, ref, sp, rw, method, endpoint, desc) in enumerate(rows):
            row = tbl.add_row()
            bg = "F5F7FA" if i % 2 == 0 else "FFFFFF"
            vals = [ref, sp, rw, f"{method} {endpoint}".strip(), desc]
            for j, (val, width) in enumerate(zip(vals, widths)):
                cell = row.cells[j]
                cell.width = Inches(width)
                set_cell_bg(cell, bg)
                set_cell_borders(cell, "D0D7E2", sz=4)
                para = cell.paragraphs[0]
                r = para.add_run(val)
                r.font.size = Pt(8)
                r.font.color.rgb = TEXT
                if j == 1:  # sp name — monospace-ish
                    r.font.color.rgb = BLUE
                if j == 2:  # R/W badge
                    if rw == "READ":
                        r.font.color.rgb = RGBColor(0x27, 0x67, 0x49)
                    elif rw == "WRITE":
                        r.font.color.rgb = RGBColor(0xB3, 0x52, 0x1E)
                    else:
                        r.font.color.rgb = MUTED
                    r.bold = True
                if j == 3:  # endpoint
                    r.font.color.rgb = BLUE
                set_para_spacing(para, before=20, after=20)

        sp2 = doc.add_paragraph()
        set_para_spacing(sp2, before=0, after=80)

    doc.add_page_break()


def section_security(doc: Document) -> None:
    heading1(doc, "Security, HIPAA & Governance", "06")
    body_para(doc,
        "OXPLive contains protected health information — patient demographics, insurance details, "
        "DME order data, and EDI 835/837 claims. Security controls are designed into every layer "
        "of this architecture, not added afterward.")

    sec_rows = [
        ("Service Account Isolation",
         "oxplive_mcp_svc holds only GRANT EXECUTE ON SCHEMA::oxadmin. No table-level SELECT, INSERT, "
         "UPDATE, or DELETE. No dynamic SQL is possible."),
        ("SP-Only Access Pattern",
         "Every AI request resolves to a named, version-controlled stored procedure. Inline SQL from "
         "AI-generated queries is architecturally impossible."),
        ("Audit Logging",
         "Domain 17 exposes audit-log procedures. HIPAA-sensitive SPs (Document_GetViewUrl, all "
         "patient-record writers) trigger server-side audit entries recording user, action, timestamp, "
         "and affected record ID."),
        ("PHI Field Masking",
         "SSN and other PHI identifiers are masked at the stored-procedure level before they reach the "
         "MCP response payload. Patient-data fields require explicit role clearance."),
        ("Document URLs",
         "File access (usp_Document_GetViewUrl) returns HIPAA-compliant signed URLs with short-lived "
         "expiry. No raw file paths or permanent links are returned."),
        ("Transport Security",
         "Streamable HTTP runs over HTTPS in production. Optional MCP_API_KEY header provides HTTP-level "
         "bearer auth. Network access restricted by firewall to authorised IP ranges."),
        ("Session Isolation",
         "Each Claude session receives a unique mcp-session-id. Sessions are server-memory-scoped; "
         "one session cannot access another session's state or results."),
        ("Compliance Domains",
         "Domain 19 (Yes/No Module) enforces the three compliance gates — PA Status, CMN receipt, "
         "EDF sign-off — that must be satisfied before a job can progress. Domain 18 handles mass-move "
         "and scheduled compliance batch operations."),
    ]

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = tbl.rows[0]
    for cell_text, width in zip(["Control", "Implementation"], [1.5, 5.1]):
        c = hrow.cells[sec_rows.index((cell_text, "")) if (cell_text, "") in sec_rows else list(zip(["Control","Implementation"], [1.5,5.1])).index((cell_text, width))]
        # simplify: just set by index
    for i, cell in enumerate(hrow.cells):
        w = [1.5, 5.1][i]
        cell.width = Inches(w)
        set_cell_bg(cell, "006DB6")
        set_cell_borders(cell, "006DB6")
        r = cell.paragraphs[0].add_run(["Control", "Implementation"][i])
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE
        set_para_spacing(cell.paragraphs[0], before=40, after=40)

    for i, (control, impl) in enumerate(sec_rows):
        row = tbl.add_row()
        bg = "F5F7FA" if i % 2 == 0 else "FFFFFF"
        for j, (val, width) in enumerate(zip([control, impl], [1.5, 5.1])):
            cell = row.cells[j]
            cell.width = Inches(width)
            set_cell_bg(cell, bg)
            set_cell_borders(cell, "D0D7E2", sz=4)
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(8.5)
            r.font.color.rgb = TEXT
            if j == 0:
                r.bold = True
            set_para_spacing(cell.paragraphs[0], before=30, after=30)

    body_para(doc,
        "This design supports OrthoKinetix's HIPAA compliance posture. It does not replace the "
        "organisation's existing policies, Business Associate Agreements, or staff training, "
        "which remain in place.")
    doc.add_page_break()


def section_what_gets_built(doc: Document) -> None:
    heading1(doc, "What Gets Built", "07")
    heading2(doc, "MCP Server — Node.js Application")
    body_para(doc,
        "The OXPLive MCP server is a Node.js / TypeScript application built on the official "
        "@modelcontextprotocol/sdk package. It uses the mssql library for SQL Server connectivity "
        "and Express for the Streamable HTTP transport.")
    for item in [
        ("DB Client", "Singleton mssql connection pool; executeSP(), querySP(), and querySPSingle() "
         "helpers that auto-prefix oxadmin. to every procedure name"),
        ("20 Domain Tool Files", "One TypeScript file per domain, each exporting its MCP tool "
         "definitions and request handlers"),
        (f"Server Core", f"createOxpliveMcpServer() registers all {TOTAL_TOOLS} tools and dispatches "
         "calls to the matching domain handler"),
        ("Streamable HTTP Transport", "Express app with session map; supports multiple concurrent "
         "Claude sessions"),
    ]:
        bullet_para(doc, item[1], bold_prefix=item[0] + " —")

    heading2(doc, "SQL Server Service Account")
    body_para(doc,
        "A dedicated, least-privilege SQL Server login is created specifically for the MCP server. "
        "No further permissions are required. The service account cannot read, write, or alter any "
        "table, view, or function directly.")
    body_para(doc,
        "Provisioning script:  CREATE LOGIN oxplive_mcp_svc WITH PASSWORD = '<vault-managed>';  "
        "CREATE USER oxplive_mcp_svc FOR LOGIN oxplive_mcp_svc;  "
        "GRANT EXECUTE ON SCHEMA::oxadmin TO oxplive_mcp_svc;",
        indent=True)

    heading2(doc, "Configuration & Deployment")
    body_para(doc,
        "All connection strings and secrets are managed via environment variables (never committed to "
        "source control). The server ships with a .env.example file and a config/claude-code.mcp.json "
        "snippet for immediate Claude Code integration.")

    heading2(doc, "What Is NOT in Scope")
    for item in [
        "The .NET 8 REST API layer (Phase F — planned, not yet approved)",
        "An admin portal or management dashboard",
        "Changes to the existing OXPLive ColdFusion application",
        "Changes to the OrthoXpressDB schema or existing stored procedures",
    ]:
        bullet_para(doc, item)
    doc.add_page_break()


def section_delivery(doc: Document) -> None:
    heading1(doc, "Delivery Plan", "08")
    body_para(doc,
        "The MCP server delivery is structured in three phases. Each phase ends with a clear, "
        "demonstrable result that OrthoKinetix can test before the next phase begins.")

    headers = ["Phase", "Deliverable", "Scope", "Exit Criteria"]
    widths  = [0.7, 1.4, 3.2, 1.6]
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    make_table_header_row(tbl, headers, widths)

    phase_data = [
        ("Phase 1",
         "Stored Procedure Foundation",
         f"All {TOTAL_STPS} stored procedures authored, reviewed, and deployed to the OXPLive "
         "test database. Includes read-only procedures for all 20 domains plus write procedures "
         "for core entities.",
         "All SPs pass unit tests against test DB; naming convention compliant; HIPAA-sensitive "
         "fields verified"),
        ("Phase 2",
         "MCP Server — Build & Test",
         f"Node.js MCP server with all {TOTAL_TOOLS} tools, DB-Direct wiring, Streamable HTTP "
         "transport, service account provisioning, and Claude Code integration.",
         f"Claude can invoke all {TOTAL_TOOLS} tools against test database; audit log confirms "
         "all calls recorded"),
        ("Phase 3",
         "Production Deployment",
         "Production environment setup, credential vault integration, firewall rules, HTTPS, "
         "staff training session, and final sign-off from OrthoKinetix IT.",
         "Server running in production; OrthoKinetix team can connect Claude and execute "
         "representative queries"),
    ]
    for i, row_vals in enumerate(phase_data):
        add_table_data_row(tbl, list(row_vals), i, widths)

    body_para(doc,
        "Detailed timeline and effort estimates are confirmed during the kick-off session and "
        "provided in a companion Statement of Work upon approval of this document.")
    doc.add_page_break()


def section_next_steps(doc: Document) -> None:
    heading1(doc, "What We Need From You", "09")
    body_para(doc,
        "Approval of this document authorises Technijian to begin Phase 1. A short kick-off session "
        "is scheduled immediately afterward to confirm the items below. None of these decisions "
        "block approval — they shape the detailed plan.")
    for bp, text in [
        ("Domain priority order.",
         "Which of the 20 domains should be built first? Suggested default: Domains 1–5 — Auth, "
         "Patient, Insurance Auth, Queue, Billing."),
        ("Test database access.",
         "Confirmation that Technijian may continue working against the existing test/UAT "
         "environment for Phase 1 and Phase 2."),
        ("Service account provisioning.",
         "Authorisation to create the oxplive_mcp_svc login and grant it EXECUTE on the oxadmin schema."),
        ("Credential vault.",
         "Preferred secret-management approach (Azure Key Vault, Windows Credential Manager, or "
         "Technijian-managed vault) for the service account password and API key."),
        ("Hosting location.",
         "Where the MCP server will run — on-premises IIS, Windows service, Docker container, "
         "or Technijian-hosted."),
        ("Internal users.",
         "Who at OrthoKinetix will use Claude with the MCP tools, and at what initial permission "
         "level (read-only vs. read/write access to write-capable domains)."),
    ]:
        bullet_para(doc, text, bold_prefix=bp)
    doc.add_page_break()


def section_approval(doc: Document) -> None:
    heading1(doc, "Approval & Sign-Off", "10")
    body_para(doc,
        "By signing below, OrthoKinetix authorises Technijian to proceed with the OXPLive MCP "
        "Server project as described in this document, beginning with Phase 1 (Stored Procedure "
        "Foundation). This approval covers the technical approach, domain scope, and phased "
        "delivery plan.")
    body_para(doc,
        "Specific timeline, effort hours, and fees are confirmed in a companion Statement of Work, "
        "which will be issued within five (5) business days of receiving this signed approval.")

    sp = doc.add_paragraph()
    set_para_spacing(sp, before=0, after=120)

    # Signature table (2 columns)
    sig_tbl = doc.add_table(rows=1, cols=2)
    sig_tbl.style = "Table Grid"
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    parties = [
        ("OrthoKinetix", "Authorised on behalf of OrthoKinetix / OrthoXpress", None, None),
        ("Technijian",   "Prepared and submitted by Technijian Inc.",
         "Rohit Jain, Technijian", "June 2026"),
    ]

    row = sig_tbl.rows[0]
    for i, (name, subtitle, signer, date) in enumerate(parties):
        cell = row.cells[i]
        cell.width = Inches(3.2)
        set_cell_borders(cell, "D0D7E2", sz=6)

        p = cell.paragraphs[0]
        rn = p.add_run(name)
        rn.bold = True; rn.font.size = Pt(11); rn.font.color.rgb = DARK
        set_para_spacing(p, before=60, after=40)

        p2 = cell.add_paragraph(subtitle)
        p2.runs[0].font.size = Pt(8.5); p2.runs[0].font.color.rgb = MUTED
        set_para_spacing(p2, before=0, after=120)

        # Signature line
        p3 = cell.add_paragraph()
        set_para_spacing(p3, before=0, after=0)
        add_sig_line_in_cell(cell, signer or "_________________________", date or "_____")

        p5 = cell.add_paragraph()
        set_para_spacing(p5, before=120, after=0)
        add_sig_line_in_cell(cell, "Name (print)" if not signer else "", "Title" if not date else "")

    sp2 = doc.add_paragraph()
    set_para_spacing(sp2, before=0, after=120)

    body_para(doc,
        "Questions? Contact Technijian at 949.379.8499 (main) or rjain@technijian.com. "
        "The full engineering specification, STP-API-MCP Coverage Matrix, and Statement of Work "
        "are available on request.")
    doc.add_page_break()


def add_sig_line_in_cell(cell, left_text: str, right_text: str) -> None:
    p = cell.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "2D3748")
    pBdr.append(top)
    pPr.append(pBdr)
    r1 = p.add_run(left_text)
    r1.font.size = Pt(8.5); r1.font.color.rgb = MUTED
    if right_text:
        tab = p.add_run("\t")
        r2 = p.add_run(right_text)
        r2.font.size = Pt(8.5); r2.font.color.rgb = MUTED
    set_para_spacing(p, before=40, after=60)


def section_appendix(doc: Document) -> None:
    heading1(doc, "Appendix — Full Domain Reference", "A")

    heading2(doc, "A.1   Domain Summary (authoritative counts)")
    headers = ["#", "Domain", "STPs", "API Endpoints", "MCP Tools"]
    widths  = [0.3, 2.5, 0.8, 1.0, 0.9]
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    make_table_header_row(tbl, headers, widths)
    for i, d in enumerate(DOMAINS):
        add_table_data_row(tbl, [str(d[0]), d[1], str(d[2]), str(d[3]), str(d[4])], i, widths)
    row = tbl.add_row()
    for j, (val, width) in enumerate(zip(["", "TOTAL", str(TOTAL_STPS), str(TOTAL_ENDPOINTS), str(TOTAL_TOOLS)], widths)):
        cell = row.cells[j]
        cell.width = Inches(width)
        set_cell_bg(cell, "1A2B3C")
        set_cell_borders(cell, "FFFFFF", sz=4)
        r = cell.paragraphs[0].add_run(val)
        r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = WHITE
        set_para_spacing(cell.paragraphs[0], before=40, after=40)

    sp = doc.add_paragraph()
    set_para_spacing(sp, before=0, after=120)

    heading2(doc, "A.2   The 13 Workflow Queues")
    body_para(doc,
        "The OXPLive platform routes DME orders through 13 named queues. Domain 4 (Queue & Workflow) "
        "provides the stored procedures that govern job transitions between queues.")

    queue_headers = ["#", "Queue", "Purpose"]
    queue_widths  = [0.3, 1.2, 4.3]
    queues = [
        ("1",  "Intake",            "New DME orders received; initial eligibility check"),
        ("2",  "Benefits",          "Insurance benefit verification in progress"),
        ("3",  "Auth",              "Prior authorisation requested or under review"),
        ("4",  "Pending",           "Awaiting patient or clinical information"),
        ("5",  "Shipping",          "Order approved; equipment dispatch in progress"),
        ("6",  "Billing",           "Delivered; claim submitted to payer"),
        ("7",  "Re-Bill",           "Claim denied or adjusted; re-submission in progress"),
        ("8",  "EDF",               "Equipment Delivery Form outstanding — compliance hold"),
        ("9",  "Collections",       "Overdue balance; active collector follow-up"),
        ("10", "Yes/No",            "PA Status / CMN / EDF compliance gates pending"),
        ("11", "Inactive",          "Order on hold; no active work"),
        ("12", "Cancelled",         "Order cancelled; record retained for audit"),
        ("13", "Collector Actions", "Collector daily-action tracking and totals"),
    ]
    q_tbl = doc.add_table(rows=1, cols=3)
    q_tbl.style = "Table Grid"
    q_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    make_table_header_row(q_tbl, queue_headers, queue_widths)
    for i, vals in enumerate(queues):
        add_table_data_row(q_tbl, list(vals), i, queue_widths)

    sp2 = doc.add_paragraph()
    set_para_spacing(sp2, before=0, after=120)

    heading2(doc, "A.3   Companion Documents")
    bullet_para(doc,
        f"All {TOTAL_STPS} stored procedures with full parameter lists, return schemas, API endpoint "
        "mappings, and MCP tool definitions. Provided as a separate A3-landscape PDF deliverable.",
        bold_prefix="STP-API-MCP Coverage Matrix —")
    bullet_para(doc,
        "Issued within 5 business days of approval. Covers phased timeline, effort hours, and fees.",
        bold_prefix="Statement of Work —")


# ─── Main ─────────────────────────────────────────────────────────────────────

def main() -> None:
    doc = Document()

    # Page setup: A4 portrait
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width  = Cm(21.0)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.2)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(9.5)

    build_cover(doc)
    section_exec_summary(doc)
    section_current_state(doc)
    section_architecture(doc)
    section_domain_matrix(doc)
    section_detailed_capability(doc)
    section_security(doc)
    section_what_gets_built(doc)
    section_delivery(doc)
    section_next_steps(doc)
    section_approval(doc)
    section_appendix(doc)

    doc.save(str(OUT_DOCX))
    size_kb = OUT_DOCX.stat().st_size // 1024
    print(f"DOCX: {OUT_DOCX}  ({size_kb} KB)")


if __name__ == "__main__":
    main()
