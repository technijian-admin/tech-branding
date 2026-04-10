"""
Nexus Assess Marketing Datasheet — python-docx Generator
Produces a branded 2-4 page DOCX with Word COM PDF conversion.
"""
import os
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUT_DIR, "Nexus Assess Marketing Datasheet.docx")
PDF_PATH = os.path.join(OUT_DIR, "Nexus Assess Marketing Datasheet.pdf")

# Brand colors
BLUE = RGBColor(0x00, 0x6D, 0xB6)
ORANGE = RGBColor(0xF6, 0x7D, 0x4B)
TEAL = RGBColor(0x1E, 0xAA, 0xC8)
CHARTREUSE = RGBColor(0xCB, 0xDB, 0x2D)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GREY = RGBColor(0x59, 0x59, 0x5B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xE9, 0xEC, 0xEF)

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_colored_heading(doc, text, level=1, color=BLUE):
    """Add a heading with brand color."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Open Sans'
    return h

def add_colored_divider(doc, color_hex="006DB6"):
    """Add a thin colored bar."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def build_doc():
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ── HERO ──
    h = doc.add_heading('', level=0)
    run = h.add_run('Nexus ')
    run.font.size = Pt(32)
    run.font.color.rgb = DARK
    run.font.name = 'Open Sans'
    run.bold = True
    run = h.add_run('Assess')
    run.font.size = Pt(32)
    run.font.color.rgb = ORANGE
    run.font.name = 'Open Sans'
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run('IT RISK & COMPLIANCE PLATFORM')
    run.font.size = Pt(10)
    run.font.color.rgb = TEAL
    run.font.name = 'Open Sans'
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run('Automated IT environment scanning, risk scoring, and compliance reporting across your entire infrastructure.')
    run.font.size = Pt(12)
    run.font.color.rgb = GREY
    run.font.name = 'Open Sans'
    run.italic = True

    add_colored_divider(doc, "F67D4B")

    # ── HERO STATS TABLE ──
    stats_table = doc.add_table(rows=1, cols=4)
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats = [("94+", "Report Types"), ("3", "OS Platforms"), ("24/7", "Monitoring"), ("10", "Frameworks")]
    for i, (num, label) in enumerate(stats):
        cell = stats_table.cell(0, i)
        set_cell_shading(cell, "1A1A2E")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(num + "\n")
        run.font.size = Pt(20)
        run.font.color.rgb = ORANGE
        run.font.name = 'Open Sans'
        run.bold = True
        run = p.add_run(label)
        run.font.size = Pt(8)
        run.font.color.rgb = WHITE
        run.font.name = 'Open Sans'

    doc.add_paragraph()
    add_colored_divider(doc, "1EAAC8")

    # ── THE PROBLEM ──
    add_colored_heading(doc, "The Problem You're Facing", level=2)

    problems = [
        ("Blind Spots in Hybrid Environments", "Your network spans on-premise, cloud, and remote endpoints \u2014 but no single tool sees it all."),
        ("Manual Audits Drain Resources", "Security assessments take weeks across multiple tools. By the time reports compile, data is stale."),
        ("Compliance Complexity Multiplies", "HIPAA, SOC 2, PCI-DSS, NIST, CMMC \u2014 each demands unique evidence and control mapping."),
    ]
    for title, desc in problems:
        p = doc.add_paragraph()
        run = p.add_run("\u25A0 " + title + "  ")
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
        run.font.name = 'Open Sans'
        run.bold = True
        run = p.add_run(desc)
        run.font.size = Pt(10)
        run.font.color.rgb = GREY
        run.font.name = 'Open Sans'

    add_colored_divider(doc, "006DB6")

    # ── CORE CAPABILITIES ──
    add_colored_heading(doc, "Core Capabilities", level=2)

    capabilities = [
        ("\U0001F50D Discovery & Scanning", BLUE, [
            "Non-intrusive scans on Windows, macOS, and Linux",
            "Remote data collectors & lightweight ISAA agents",
            "Automated scheduling with continuous monitoring",
            "SQL Server database security assessments",
        ]),
        ("\u2601 Cloud & Microsoft 365", ORANGE, [
            "M365, Teams, SharePoint, OneDrive & Exchange audits",
            "Azure AD & Active Directory assessments",
            "AWS infrastructure evaluation & config review",
            "Conditional Access Policy gap analysis",
        ]),
        ("\U0001F6E1 Risk Intelligence", TEAL, [
            "Dual scoring: proprietary + CVSS + MS Secure Score",
            "Dark web monitoring & credential exposure detection",
            "Vulnerability scanning with CVE correlation",
            "AI-powered executive summaries",
        ]),
        ("\U0001F4C4 94+ Report Types", DARK, [
            "Network (32), Security (39), Cloud (13), Compliance (10)",
            "Executive infographics & QBR presentations",
            "Multi-format: DOCX, PDF, XLSX, PPTX, VSDX",
            "Branded client-facing reports with your logo",
        ]),
    ]

    for title, color, items in capabilities:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(12)
        run.font.color.rgb = color
        run.font.name = 'Open Sans'
        run.bold = True
        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run("  \u2713 " + item)
            run.font.size = Pt(9.5)
            run.font.color.rgb = GREY
            run.font.name = 'Open Sans'

    add_colored_divider(doc, "F67D4B")

    # ── PRICING ──
    add_colored_heading(doc, "Plans & Pricing", level=2)

    pricing_table = doc.add_table(rows=5, cols=4)
    pricing_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Plan", "Monthly", "Annual", "Endpoints"]
    for i, h_text in enumerate(headers):
        cell = pricing_table.cell(0, i)
        set_cell_shading(cell, "1A1A2E")
        p = cell.paragraphs[0]
        run = p.add_run(h_text)
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.font.name = 'Open Sans'
        run.bold = True

    tiers = [
        ("Essentials", "$499/mo", "$4,990/yr", "Up to 50"),
        ("Professional", "$999/mo", "$9,990/yr", "Up to 250"),
        ("Enterprise", "$1,999/mo", "$19,990/yr", "Unlimited"),
        ("Managed", "Custom", "Custom", "Full service"),
    ]
    for ri, (name, mo, yr, ep) in enumerate(tiers):
        for ci, val in enumerate([(name, DARK), (mo, GREY), (yr, GREY), (ep, GREY)]):
            cell = pricing_table.cell(ri + 1, ci)
            if ri == 1:  # Popular
                set_cell_shading(cell, "EBF5FB")
            p = cell.paragraphs[0]
            run = p.add_run(val[0])
            run.font.size = Pt(10)
            run.font.color.rgb = BLUE if (ri == 1 and ci == 0) else val[1]
            run.font.name = 'Open Sans'
            if ci == 0:
                run.bold = True

    p = doc.add_paragraph()
    run = p.add_run("Annual plans save 17% \u2022 Professional includes dark web monitoring \u2022 30-day free pilot available")
    run.font.size = Pt(8.5)
    run.font.color.rgb = TEAL
    run.font.name = 'Open Sans'
    run.bold = True

    add_colored_divider(doc, "1EAAC8")

    # ── COMPETITIVE COMPARISON ──
    add_colored_heading(doc, "Competitive Comparison", level=2)

    cmp_table = doc.add_table(rows=11, cols=5)
    cmp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cmp_headers = ["Capability", "Qualys", "Nessus", "Auvik", "Nexus Assess"]
    for i, h_text in enumerate(cmp_headers):
        cell = cmp_table.cell(0, i)
        set_cell_shading(cell, "006DB6" if i == 4 else "1A1A2E")
        p = cell.paragraphs[0]
        run = p.add_run(h_text)
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.font.name = 'Open Sans'
        run.bold = True

    cmp_rows = [
        ("Agentless scans", "\u2713", "\u2713", "\u2014", "\u2713"),
        ("M365 auditing", "\u2014", "\u2014", "\u2014", "\u2713"),
        ("94+ report types", "\u2014", "\u2014", "\u2014", "\u2713"),
        ("Dark web monitor", "\u2014", "\u2014", "\u2014", "\u2713"),
        ("Compliance GRC", "Add-on", "\u2014", "\u2014", "\u2713 Built-in"),
        ("AI exec summaries", "\u2014", "\u2014", "\u2014", "\u2713"),
        ("PSA integration", "\u2014", "\u2014", "Limited", "\u2713 Full"),
        ("Multi-format export", "PDF only", "PDF/CSV", "\u2014", "\u2713 5 formats"),
        ("Branded reports", "\u2014", "\u2014", "\u2014", "\u2713"),
        ("10 frameworks", "3", "\u2014", "\u2014", "\u2713"),
    ]
    for ri, row in enumerate(cmp_rows):
        for ci, val in enumerate(row):
            cell = cmp_table.cell(ri + 1, ci)
            if ci == 4:
                set_cell_shading(cell, "EBF5FB")
            p = cell.paragraphs[0]
            is_check = val.startswith("\u2713")
            run = p.add_run(val)
            run.font.size = Pt(9)
            run.font.color.rgb = TEAL if (is_check and ci == 4) else (GREY if val == "\u2014" else DARK)
            run.font.name = 'Open Sans'
            if is_check and ci == 4:
                run.bold = True

    add_colored_divider(doc, "CBDB2D")

    # ── ROI ──
    add_colored_heading(doc, "Risk Reduction & ROI", level=2)

    roi_table = doc.add_table(rows=1, cols=3)
    roi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    roi_data = [("85%", "Faster assessments\n(50 endpoints)"), ("$4.45M", "Avg breach cost\navoided (250 ep)"), ("70%", "Less audit\nprep time")]
    for i, (val, desc) in enumerate(roi_data):
        cell = roi_table.cell(0, i)
        set_cell_shading(cell, "1A1A2E")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val + "\n")
        run.font.size = Pt(22)
        run.font.color.rgb = CHARTREUSE
        run.font.name = 'Open Sans'
        run.bold = True
        run = p.add_run(desc)
        run.font.size = Pt(8)
        run.font.color.rgb = WHITE
        run.font.name = 'Open Sans'

    doc.add_paragraph()

    # ── SECURITY BADGES ──
    add_colored_heading(doc, "Security & Compliance Frameworks", level=2)

    badges = ["HIPAA", "SOC 2", "PCI-DSS", "NIST CSF", "CIS Controls", "GDPR", "CMMC", "CJIS", "FFIEC", "HPH CPGs"]
    p = doc.add_paragraph()
    for badge in badges:
        run = p.add_run("\u2713 ")
        run.font.size = Pt(10)
        run.font.color.rgb = TEAL
        run.font.name = 'Open Sans'
        run.bold = True
        run = p.add_run(badge + "    ")
        run.font.size = Pt(10)
        run.font.color.rgb = DARK
        run.font.name = 'Open Sans'
        run.bold = True

    add_colored_divider(doc, "006DB6")

    # ── CTA ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("See Your Risk Posture in Minutes")
    run.font.size = Pt(18)
    run.font.color.rgb = BLUE
    run.font.name = 'Open Sans'
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Schedule a Demo: ")
    run.font.size = Pt(11)
    run.font.color.rgb = GREY
    run.font.name = 'Open Sans'
    run = p.add_run("RJain@technijian.com")
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    run.font.name = 'Open Sans'
    run.bold = True
    run = p.add_run("  |  ")
    run.font.size = Pt(11)
    run.font.color.rgb = GREY
    run.font.name = 'Open Sans'
    run = p.add_run("949.379.8500")
    run.font.size = Pt(11)
    run.font.color.rgb = BLUE
    run.font.name = 'Open Sans'
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u00a9 2026 Technijian \u2022 18 Technology Dr. Ste 141, Irvine CA 92618 \u2022 technijian.com")
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    run.font.name = 'Open Sans'

    doc.save(DOCX_PATH)
    print(f"DOCX saved: {DOCX_PATH}")

    # Convert to PDF via Word COM
    try:
        import comtypes.client
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        doc_obj = word.Documents.Open(DOCX_PATH)
        doc_obj.SaveAs(PDF_PATH, FileFormat=17)
        doc_obj.Close()
        word.Quit()
        print(f"PDF saved: {PDF_PATH}")
    except Exception as e:
        print(f"PDF conversion skipped (Word COM not available): {e}")

if __name__ == "__main__":
    build_doc()
