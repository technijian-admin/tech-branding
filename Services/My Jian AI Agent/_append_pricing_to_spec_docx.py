"""
Append a Pricing & Packaging section to My-Jian-Specification.docx.

Idempotent: scans existing Heading 1s and skips if "Pricing & Packaging" is
already present. Otherwise inserts the pricing block immediately before the
trailing "Founded in 2000" / "Contact the Technijian team" closing paragraphs.

Run:
    py docs/_append_pricing_to_spec_docx.py
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor

DOCX_PATH = Path(r"C:\vscode\tech-branding\tech-branding\Services\My Jian AI Agent\My-Jian-Specification.docx")

DARK_BLUE_HEX = "1F4E79"
ALT_ROW_HEX = "F2F7FB"
WHITE_HEX = "FFFFFF"
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_shading(cell, color_hex: str) -> None:
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    )


def _set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor = BLACK,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT, size: Pt = Pt(9)) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(str(text))
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.name = "Calibri"


def _apply_cell_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:color="BFBFBF"/>'
        '<w:left w:val="single" w:sz="4" w:color="BFBFBF"/>'
        '<w:bottom w:val="single" w:sz="4" w:color="BFBFBF"/>'
        '<w:right w:val="single" w:sz="4" w:color="BFBFBF"/>'
        '</w:tcBorders>'
    ))


def _add_table_at(doc, headers: list[str], rows: list[list[str]], anchor_paragraph) -> None:
    """Add a table immediately after `anchor_paragraph`, return it."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        _set_cell_text(cell, h, bold=True, color=WHITE, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_shading(cell, DARK_BLUE_HEX)
        _apply_cell_borders(cell)

    for r, row in enumerate(rows, start=1):
        bg = ALT_ROW_HEX if r % 2 == 0 else WHITE_HEX
        for c, val in enumerate(row):
            cell = tbl.rows[r].cells[c]
            _set_cell_text(cell, val)
            _set_cell_shading(cell, bg)
            _apply_cell_borders(cell)

    # Move table to sit right after the anchor paragraph
    anchor_paragraph._p.addnext(tbl._tbl)


def _insert_paragraph_after(paragraph, text: str = ""):
    """Insert a new paragraph immediately after `paragraph` and return it."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_p, paragraph._parent)
    if text:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
    return p


def already_inserted(doc) -> bool:
    for p in doc.paragraphs:
        if p.style and p.style.name == "Heading 1" and "Pricing" in p.text and "Packaging" in p.text:
            return True
    return False


def find_anchor(doc):
    """Find the closing paragraph block — return the paragraph BEFORE which
    we will insert the new section. Prefer 'Founded in 2000' line; fall back
    to the last paragraph."""
    target = None
    for p in doc.paragraphs:
        if "Founded in 2000" in p.text or "founded in 2000" in p.text:
            target = p
            break
    return target or doc.paragraphs[-1]


def remove_existing_pricing_section(doc) -> bool:
    """Find the 'Pricing & Packaging' Heading-1 block and delete every
    sibling element (paragraphs + tables) after it up until the next
    Heading-1 OR the 'Founded in 2000' / 'Contact the Technijian' closing
    paragraphs. Returns True if a section was found and removed."""
    body = doc.element.body
    children = list(body.iterchildren())
    start_idx = None

    # Find the Pricing & Packaging H1
    for i, el in enumerate(children):
        if el.tag.endswith("}p"):
            from docx.text.paragraph import Paragraph
            p = Paragraph(el, doc)
            if p.style and p.style.name == "Heading 1" and "Pricing" in p.text and "Packaging" in p.text:
                start_idx = i
                break

    if start_idx is None:
        return False

    # Find the end of the section — next H1 OR closing paragraph
    end_idx = None
    for j in range(start_idx + 1, len(children)):
        el = children[j]
        if el.tag.endswith("}p"):
            from docx.text.paragraph import Paragraph
            p = Paragraph(el, doc)
            if p.style and p.style.name == "Heading 1":
                end_idx = j
                break
            if "Founded in 2000" in p.text or "founded in 2000" in p.text:
                end_idx = j
                break
    if end_idx is None:
        end_idx = len(children)

    # Delete from start_idx to end_idx - 1 (exclusive of end_idx)
    for el in children[start_idx:end_idx]:
        el.getparent().remove(el)
    return True


def _find_style(doc, name: str):
    """Look up a style by display name OR style_id, since python-docx 1.x
    refuses direct dict lookup on some DOCX files."""
    for s in doc.styles:
        try:
            if s.name == name or s.style_id == name.replace(" ", ""):
                return s
        except Exception:
            continue
    return None


def _make_helpers(doc):
    """Return closures bound to this document for inserting headings/paragraphs."""
    h_styles = {
        1: _find_style(doc, "Heading 1"),
        2: _find_style(doc, "Heading 2"),
    }

    def add_heading_after(anchor, text: str, *, level: int = 1):
        p = _insert_paragraph_after(anchor, "")
        st = h_styles.get(level)
        if st is not None:
            p.style = st
        run = p.add_run(text)
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        else:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        return p

    def add_para_after(anchor, text: str):
        p = _insert_paragraph_after(anchor, "")
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        return p

    return add_heading_after, add_para_after


def main() -> int:
    if not DOCX_PATH.exists():
        print(f"ERROR: not found: {DOCX_PATH}", file=sys.stderr)
        return 2

    doc = Document(str(DOCX_PATH))

    if already_inserted(doc):
        if remove_existing_pricing_section(doc):
            print("Existing Pricing & Packaging section removed — re-inserting fresh.")
        else:
            print("Pricing & Packaging section already present and could not be removed — no changes.")
            return 0

    cursor = find_anchor(doc)
    add_heading_after, add_para_after = _make_helpers(doc)

    # Header
    cursor = add_heading_after(cursor, "Pricing & Packaging (V3.1 — Hybrid, 40%-of-Name-Brand, Confusion-Reduced)", level=1)
    cursor = add_para_after(
        cursor,
        "Brand promise: 40% of name-brand price. Same protection. One platform. This mirrors "
        "the Technijian private-cloud pricing playbook against AWS and Azure. Every Jian line "
        "item is anchored at approximately 40% (or less) of the equivalent name-brand competitor "
        "price, so every comparison is winnable on the brochure. Pricing covers the V1 capability "
        "set described earlier in this document. Auto-Remediation, Jian Voice, and Jian Remote "
        "Agent are V2 and priced separately when they ship.",
    )

    # 1. Pricing model
    cursor = add_heading_after(cursor, "Pricing model — hybrid (per-user + per-resource)", level=2)
    cursor = add_para_after(
        cursor,
        "Two pricing axes. Per-user prices the people layer (M365 monitoring, @Jian Teams "
        "Self-Service bot, compliance evidence). Per-resource prices the infrastructure layer "
        "(per firewall, per Meraki organization, per backup environment, per vCenter, etc.). "
        "This mirrors how clients already buy IT services (Sophos per-firewall plus per-endpoint, "
        "Microsoft per-user, Veeam per-VM) — familiar buyer mental model means lower sales friction.",
    )

    # 2. Per-user pricing — two parallel tiers
    cursor = add_heading_after(cursor, "Per-user pricing — Jian for M365 (two parallel tiers)", level=2)
    user_headers = ["Tier", "Price", "Annual (2 mo free)", "Best for"]
    user_rows = [
        ["Jian for M365 — Standard",       "$9/user/mo",  "$90/user/yr",  "General SMB; covers M365 monitoring, @Jian Teams bot, HIPAA/SOC 2/PCI compliance, branded monthly M365 report, SIEM correlation, MITRE tagging"],
        ["Jian for M365 — Compliance Pro", "$24/user/mo", "$240/user/yr", "Regulated / DoD / multi-site; everything in Standard PLUS CMMC + custom frameworks + dedicated CISSP-certified CSM + cross-client correlator + Service Level Agreement (terms in MSA) + custom @Jian skills deployed to your Microsoft Teams workspace + CISSP-led audit-firm liaison"],
    ]
    _add_table_at(doc, user_headers, user_rows, cursor)
    cursor = add_para_after(cursor, "")
    cursor = add_para_after(
        cursor,
        "Minimum: 25 seats on Standard ($225/mo floor). Why parallel tiers (not 'upgrade'): "
        "two prices on the page is cleaner than '+$15 add-on.' Standard is fully featured; "
        "Compliance Pro is for clients who need additional regulatory scope and white-glove service.",
    )

    # 2b. Bot-only SKU
    cursor = add_heading_after(cursor, "Bot-only SKU — Jian Teams Bot Only", level=2)
    bot_headers = ["SKU", "Price", "Cap", "Best for"]
    bot_rows = [
        ["Jian Teams Bot Only", "$99/mo flat", "Up to 25 users", "5–25 seat client who just wants @Jian in Teams. No platform monitoring, no compliance, no reports. Above 25 users, must upgrade to Standard."],
    ]
    _add_table_at(doc, bot_headers, bot_rows, cursor)
    cursor = add_para_after(cursor, "")
    cursor = add_para_after(
        cursor,
        "Replaces the V1 $49/mo IT Self-Service Agent for new buyers (existing $49 customers "
        "are grandfathered for a defined grandfathering period — see Grandfathering section).",
    )

    # 3. Per-resource pricing
    cursor = add_heading_after(cursor, "Per-resource pricing — Jian for [Infrastructure]", level=2)
    res_headers = ["Service", "Jian Price", "Annual (2 mo free)", "Name-Brand Comp", "% of Comp"]
    res_rows = [
        ["Jian for Firewall (per Sophos/Meraki firewall)",                    "$99/firewall/mo",     "$990/yr",     "Arctic Wolf $300–$500/site; Sophos My Sophos $99–$299",     "20–33%"],
        ["Jian for Meraki Network (per Meraki org)",                          "$79/org/mo",          "$790/yr",     "Auvik $250/site",                                          "32%"],
        ["Jian for Backup (per Veeam VBR / 365 / ONE / NAKIVO env)",          "$59/env/mo",          "$590/yr",     "Liongard $50–$100 (no NAKIVO)",                            "59–118%"],
        ["Jian for DNS (per Cisco Umbrella tenant)",                          "$39/tenant/mo",       "$390/yr",     "Cisco Umbrella reporting add-ons ~$80+",                   "49%"],
        ["Jian for Endpoint EDR — Workstation (overlay your EDR vendor)",     "$1.50/workstation/mo","$15/yr",      "Huntress $6/desktop",                                      "25%"],
        ["Jian for Endpoint EDR — Server (overlay your EDR vendor)",          "$3.50/server/mo",     "$35/yr",      "Huntress $6/server, CrowdStrike $10–$15",                  "23–58%"],
        ["Jian for vCenter (per vCenter Server)",                             "$99/vCenter/mo",      "$990/yr",     "VMware vROps Premium ~$150+",                              "66%"],
        ["Jian for Web / Cloudflare (per zone)",                              "$29/zone/mo",         "$290/yr",     "Cloudflare paid + manual labor",                           "~40%"],
    ]
    _add_table_at(doc, res_headers, res_rows, cursor)
    cursor = add_para_after(cursor, "")
    cursor = add_para_after(
        cursor,
        "Backup vendor agnosticism is a distinct selling point. Jian for Backup covers Veeam "
        "(VBR + 365 + ONE) AND NAKIVO at the same per-environment price. Liongard does not "
        "cover NAKIVO. Veeam ONE Reporter only covers Veeam. Jian bridges any mix at one rate "
        "and one branded monthly report format.",
    )

    # 3b. Starter Bundles
    cursor = add_heading_after(cursor, "Starter Bundles (flat-price alternative to itemized)", level=2)
    bundle_headers = ["Bundle", "Includes", "Flat Price", "vs. Itemized"]
    bundle_rows = [
        ["Jian Starter",             "Standard 25 seats + 1 firewall + 1 backup",                                    "$349/mo",   "itemized $383 (saves $34)"],
        ["Jian Pro Bundle",          "Compliance Pro 50 seats + 2 firewalls + 1 backup + 1 Meraki",                  "$1,495/mo", "itemized $1,615 (saves $120)"],
        ["Jian Enterprise Bundle",   "Compliance Pro 100 seats + 3 firewalls + 1 backup + 2 Meraki + 1 DNS",          "$2,695/mo", "itemized $2,953 (saves $258)"],
    ]
    _add_table_at(doc, bundle_headers, bundle_rows, cursor)
    cursor = add_para_after(cursor, "")
    cursor = add_para_after(
        cursor,
        "Bundles are an alternative path for the buyer who wants one number, not a replacement "
        "of line-item pricing. Sales rep offers either: 'You can quote line-by-line, or we have "
        "these three bundles.' Most SMB buyers will choose the bundle that fits.",
    )

    # 4. Discounts — simplified to TWO mechanics
    cursor = add_heading_after(cursor, "Discounts — simplified to two mechanics (stack multiplicatively)", level=2)
    disc_headers = ["Mechanic", "Discount", "Trigger"]
    disc_rows = [
        ["Annual prepay",          "17% off (2 months free)",                                       "Annual billing on any service, tier, or bundle"],
        ["Volume tier",            "0% under $1k/mo; 10% $1k–$5k; 15% $5k–$15k; 20% $15k–$50k; 25% $50k+", "Auto-applied based on total monthly invoice"],
    ]
    _add_table_at(doc, disc_headers, disc_rows, cursor)
    cursor = add_para_after(cursor, "")
    cursor = add_para_after(
        cursor,
        "Worked example: a $5,000/mo client annually prepaid pays $5,000 × 0.85 (15% volume) × "
        "0.83 (17% annual) = $3,528/mo effective. CFO can verify in 10 seconds. "
        "V3.1 dropped the 4+ service bundle (value folded into list prices) and collapsed "
        "MSP-volume + seat-volume into one volume tier based on monthly invoice — cleaner math.",
    )

    # 4b. Glossary
    cursor = add_heading_after(cursor, "Resource definitions (glossary — pre-empts billing disputes)", level=2)
    gloss_headers = ["Term", "Definition"]
    gloss_rows = [
        ["Firewall",            "One managed Sophos or Meraki firewall appliance. HA pair = 1 unit. Multi-WAN on one appliance = 1 unit."],
        ["Meraki organization", "One Meraki dashboard organization, regardless of network or device count."],
        ["Backup environment",  "One Veeam VBR Server OR one Veeam 365 tenant OR one Veeam ONE Server OR one NAKIVO Director instance. Mix and match at $59 each."],
        ["Cloudflare zone",     "One apex domain (e.g., example.com). All subdomains included free."],
        ["vCenter",             "One vCenter Server instance. Enhanced Linked Mode counts as 1 per PSC."],
        ["Umbrella tenant",     "One Cisco Umbrella organization (child org under parent partner relationship)."],
        ["Workstation",         "One desktop, laptop, or VDI session running an EDR agent."],
        ["Server",              "One server endpoint (physical or virtual) with an EDR agent."],
        ["Licensed user",       "One Microsoft 365 paid license (excludes shared mailboxes, room accounts, guests)."],
    ]
    _add_table_at(doc, gloss_headers, gloss_rows, cursor)
    cursor = add_para_after(cursor, "")

    # 4c. Backup exception
    cursor = add_heading_after(cursor, "Backup line exception (transparent)", level=2)
    cursor = add_para_after(
        cursor,
        "Jian for Backup at $59 vs. Liongard at $50–$100 = 59–118 percent of comp — the one "
        "line that breaks our 40 percent promise. We hold $59 because: (1) Liongard is already "
        "a discount product; pricing Jian at 40 percent of Liongard would be below cost-to-serve, "
        "(2) NAKIVO support is unique — Liongard doesn't cover NAKIVO, Veeam ONE Reporter only "
        "covers Veeam, Jian bridges both at one rate, (3) we win on integration with the rest "
        "of Jian, not on raw price. This exception is stated explicitly so a procurement audit "
        "doesn't catch us being inconsistent.",
    )

    # 5. Grandfathering + legacy SKUs
    cursor = add_heading_after(cursor, "Grandfathering and legacy SKUs", level=2)
    cursor = add_para_after(
        cursor,
        "Existing $49/month IT Self-Service Agent customers stay at $49 for the original "
        "bot-only feature set for a defined grandfathering period from V3 launch. Any upgrade to platform monitoring "
        "(Jian for M365 Core or any per-resource service) moves them to V3 pricing for that scope.",
    )
    cursor = add_para_after(
        cursor,
        "Standalone report SKUs (My M365 Report $15–$129/tenant, My Meraki Report $150/org, "
        "My Sophos Report $99/$199/$299/tenant) remain available for buyers who want a single "
        "report and never the platform. Platform subscribers get equivalent reports auto-included.",
    )

    # 6. Worked examples
    cursor = add_heading_after(cursor, "Worked examples — three client profiles", level=2)
    ex_headers = ["Profile", "Setup", "Monthly", "Annual prepay", "Comp stack", "% of Comp"]
    ex_rows = [
        ["Small SMB (20-user law firm, 1 site, 1 firewall, 1 backup)",                          "Core (25-seat min) + 1 firewall + 1 backup",                                                    "$383",   "$3,830",  "~$960",    "40%"],
        ["Medium regulated (50-user dental group, HIPAA, 2 sites, 2 firewalls, 2 Meraki, 1 backup)", "Enterprise (50 × $24) + 2 firewalls + 2 Meraki + 1 backup, –10% bundle",                  "$1,453", "$14,530", "~$5,550",  "26%"],
        ["Mid-market (250 users multi-site, 5 firewalls, 3 Meraki, 2 backups, regulated)",      "Enterprise (250 × $24) + 5 firewalls + 3 Meraki + 2 backups, –10% bundle, –15% seat-volume",   "$5,265", "$52,650", "~$28,350", "19%"],
    ]
    _add_table_at(doc, ex_headers, ex_rows, cursor)
    cursor = add_para_after(cursor, "")

    # 7. Labor displacement
    cursor = add_heading_after(cursor, "Labor displacement — IT support hours reduced (V1 proactive only)", level=2)
    cursor = add_para_after(
        cursor,
        "V1 monitors, classifies, correlates, opens tickets, routes them, and produces branded "
        "monthly reports. It does NOT auto-fix (remediation is V2). The labor savings below come "
        "entirely from proactive detection and triage — not machine-executed repair.",
    )
    labor_headers = ["Client profile", "Manual baseline (hrs/mo)", "With Jian (hrs/mo)", "Hours saved (hrs/mo)", "Value at $75–$100/hr"]
    labor_rows = [
        ["Small (10–25 users)",  "18–30",  "4–8",   "14–22", "$1,050–$2,200"],
        ["Medium (26–100 users)",       "35–60",  "8–15",  "27–45", "$2,025–$4,500"],
        ["Large (101+ users)",        "60–120", "15–30", "45–90", "$3,375–$9,000"],
    ]
    _add_table_at(doc, labor_headers, labor_rows, cursor)
    cursor = add_para_after(cursor, "")
    cursor = add_para_after(
        cursor,
        "Net reduction in proactive IT support hours: roughly 45 percent at the Enterprise "
        "upgrade tier (the headline number). Driven by seven labor sinks Jian removes: Tier-1 "
        "alert triage across 5+ monitoring tools (5–15 hours), cross-platform incident "
        "correlation (1.5–6 hours), compliance evidence collection (2–6 hours), monthly client "
        "report assembly (6–18 hours), M365 service-health user triage (0.5–3 hours), "
        "Teams/CP channel coordination (0.5–1.5 hours), shift-handover friction (0.5–1.5 hours). "
        "Total typical: 16–51 hours per client per month.",
    )

    # 8. Cost-to-serve
    cursor = add_heading_after(cursor, "Cost-to-serve and margin discipline", level=2)
    cost_headers = ["Line", "Per client per month"]
    cost_rows = [
        ["LLM (Gemini Flash primary; includes safety multiple)",  "$5"],
        ["Infrastructure (SQL/n8n/IIS/LiteLLM/Azure App Service)", "$7"],
        ["Engineering maintenance (amortized; drops with scale)", "$150–$300"],
        ["Cost to serve (one client)",                             "$165–$315/mo"],
    ]
    _add_table_at(doc, cost_headers, cost_rows, cursor)
    cursor = add_para_after(cursor, "")
    cursor = add_para_after(
        cursor,
        "No single Jian line item sells below cost-to-serve at the 25-seat minimum. At year-3 "
        "projected mix (100+ clients), engineering maintenance drops below $100/client/month and "
        "blended gross margin reaches 85–92 percent — premium SaaS economics.",
    )

    # 9. V2 upside
    cursor = add_heading_after(cursor, "V2 upside (priced separately when shipped)", level=2)
    v2_headers = ["V2 capability", "Additional hours saved/client/mo", "Tentative pricing"]
    v2_rows = [
        ["Auto-remediation (gated; 7-day clean trace + per-op approval)", "5–15", "+$10/user or new top tier"],
        ["Jian Voice — (949) 502-0074 inbound deflection",                 "2–8",  "+$5/user/mo"],
        ["Jian Remote Agent — OS / device / network / client-MSSQL data", "3–10", "$99/agent or +$3/endpoint"],
    ]
    _add_table_at(doc, v2_headers, v2_rows, cursor)
    cursor = add_para_after(cursor, "")
    cursor = add_para_after(
        cursor,
        "These are pricing upside for the V2 conversation, not part of V1 justification. V2 SKU "
        "decisions are finalized at V2 production launch.",
    )

    # Save
    doc.save(str(DOCX_PATH))
    print(f"OK — replaced/inserted V3 Pricing & Packaging section in {DOCX_PATH.name}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
