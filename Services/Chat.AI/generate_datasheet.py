"""
Chat.AI Marketing Datasheet — ReportLab PDF generation.
Full-bleed, 4-page letter-size document with Technijian brand identity.
Also generates DOCX via python-docx for email attachment use.
"""

import os
import subprocess
import requests
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, Color, white, black
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
W, H = letter  # 612 x 792 pts

# ── Brand Colors ──
BLUE       = HexColor('#006DB6')
BLUE_DARK  = HexColor('#004D80')
ORANGE     = HexColor('#F67D4B')
ORANGE_DK  = HexColor('#E5663A')
TEAL       = HexColor('#1EAAC8')
CHARTREUSE = HexColor('#CBDB2D')
DARK       = HexColor('#1A1A2E')
DARK2      = HexColor('#0A1628')
DARK3      = HexColor('#0D2847')
NEAR_BLACK = HexColor('#2D2D2D')
GREY       = HexColor('#59595B')
MED_GREY   = HexColor('#ADB5BD')
LIGHT_GREY = HexColor('#E9ECEF')
OFF_WHITE  = HexColor('#F8F9FA')
WHITE_C    = HexColor('#FFFFFF')
LIGHT_BLUE = HexColor('#EBF5FB')
HIGHLIGHT_BG = Color(0, 0.43, 0.71, alpha=0.06)

# Content margins
MARGIN = 36

# ── Helpers ──

def draw_rect(c, x, y, w, h, color, radius=0):
    c.setFillColor(color)
    c.setStrokeColor(color)
    if radius:
        c.roundRect(x, y, w, h, radius, fill=1, stroke=0)
    else:
        c.rect(x, y, w, h, fill=1, stroke=0)

def draw_card(c, x, y, w, h, radius=6, border_color=None, shadow=True):
    if shadow:
        c.setFillColor(Color(0, 0, 0, alpha=0.06))
        c.roundRect(x+1.5, y-1.5, w, h, radius, fill=1, stroke=0)
    c.setFillColor(WHITE_C)
    c.setStrokeColor(LIGHT_GREY)
    c.setLineWidth(0.5)
    c.roundRect(x, y, w, h, radius, fill=1, stroke=1)
    if border_color:
        c.setFillColor(border_color)
        c.rect(x, y + 2, 3.5, h - 4, fill=1, stroke=0)

def text(c, x, y, txt, size=9, color=GREY, bold=False, font='Helvetica'):
    fname = f'{font}-Bold' if bold else font
    c.setFont(fname, size)
    c.setFillColor(color)
    c.drawString(x, y, str(txt))

def text_center(c, x, y, txt, size=9, color=GREY, bold=False):
    fname = 'Helvetica-Bold' if bold else 'Helvetica'
    c.setFont(fname, size)
    c.setFillColor(color)
    c.drawCentredString(x, y, str(txt))

def text_right(c, x, y, txt, size=9, color=GREY, bold=False):
    fname = 'Helvetica-Bold' if bold else 'Helvetica'
    c.setFont(fname, size)
    c.setFillColor(color)
    c.drawRightString(x, y, str(txt))

def draw_circle(c, x, y, r, color):
    c.setFillColor(color)
    c.circle(x, y, r, fill=1, stroke=0)

def draw_check(c, x, y, color=TEAL, size=9):
    c.setFont('Helvetica-Bold', size)
    c.setFillColor(color)
    c.drawString(x, y, '\u2713')

def draw_dot(c, x, y, color, r=3):
    c.setFillColor(color)
    c.circle(x, y, r, fill=1, stroke=0)

def wrap_text(c, x, y, txt, max_width, size=8, color=GREY, leading=10, bold=False):
    fname = 'Helvetica-Bold' if bold else 'Helvetica'
    c.setFont(fname, size)
    c.setFillColor(color)
    words = txt.split()
    lines = []
    current = ''
    for w in words:
        test = f'{current} {w}'.strip()
        if c.stringWidth(test, fname, size) <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = w
    if current:
        lines.append(current)
    for line in lines:
        c.drawString(x, y, line)
        y -= leading
    return y

def fetch_logo():
    try:
        resp = requests.get("https://technijian.com/wp-content/uploads/2023/08/Logo.jpg", timeout=5)
        if resp.status_code == 200:
            return BytesIO(resp.content)
    except:
        pass
    return None

def draw_page_bg(c, color=OFF_WHITE):
    """Full-bleed background for entire page."""
    draw_rect(c, 0, 0, W, H, color)

def draw_orange_top_line(c):
    """2pt orange accent at very top of page."""
    draw_rect(c, 0, H - 2, W, 2, ORANGE)

def draw_decorative_glows(c, positions=None):
    """Add subtle decorative glows on dark sections."""
    if positions:
        for (x, y, r, clr, alpha) in positions:
            c.setFillColor(Color(clr[0], clr[1], clr[2], alpha=alpha))
            c.circle(x, y, r, fill=1, stroke=0)


# ══════════════════════════════════════════════════
# PAGE 1: Hero + Problems + Features start
# ══════════════════════════════════════════════════

def draw_page1(c, logo_data):
    draw_page_bg(c, OFF_WHITE)

    # ── HERO STRIP ── full-bleed dark, 110pt
    hero_h = 110
    hero_y = H - hero_h
    draw_rect(c, 0, hero_y, W, hero_h, DARK)
    draw_rect(c, 0, hero_y, W * 0.4, hero_h, DARK2)
    c.setFillColor(Color(0.05, 0.1, 0.18, alpha=0.5))
    c.rect(W * 0.7, hero_y, W * 0.3, hero_h, fill=1, stroke=0)

    # Decorative glows
    draw_decorative_glows(c, [
        (W - 60, H - 20, 80, (0.12, 0.67, 0.78), 0.1),
        (180, hero_y + 10, 60, (0.96, 0.49, 0.29), 0.07),
    ])

    # Orange accent line at bottom of hero (2pt)
    draw_rect(c, 0, hero_y, W, 2, ORANGE)

    # Orange top line
    draw_orange_top_line(c)

    # Logo — 36x36 square (1:1 aspect ratio preserved)
    if logo_data:
        logo_data.seek(0)
        logo = ImageReader(logo_data)
        logo_s = 36
        c.drawImage(logo, MARGIN + 4, hero_y + 62, width=logo_s, height=logo_s, mask='auto')

    # Title
    title_x = MARGIN + 46
    text(c, title_x, hero_y + 72, 'Technijian', size=24, color=WHITE_C, bold=True)
    tw = c.stringWidth('Technijian ', 'Helvetica-Bold', 24)
    text(c, title_x + tw, hero_y + 72, 'Chat.AI', size=24, color=ORANGE, bold=True)

    # Teal badge pill
    badge_text = 'ENTERPRISE AI PLATFORM'
    badge_x = title_x
    badge_y = hero_y + 56
    c.setFillColor(Color(0.12, 0.67, 0.78, alpha=0.15))
    c.roundRect(badge_x, badge_y, 130, 14, 7, fill=1, stroke=0)
    c.setStrokeColor(Color(0.12, 0.67, 0.78, alpha=0.4))
    c.setLineWidth(0.5)
    c.roundRect(badge_x, badge_y, 130, 14, 7, fill=0, stroke=1)
    text(c, badge_x + 8, badge_y + 3, badge_text, size=6.5, color=TEAL, bold=True)

    # Subtitle
    text(c, title_x, hero_y + 40, 'Secure, multi-model AI with collective intelligence,', size=8, color=Color(0.7, 0.7, 0.73, alpha=1))
    text(c, title_x, hero_y + 30, 'workflow automation, and enterprise governance.', size=8, color=Color(0.7, 0.7, 0.73, alpha=1))

    # 4 stats — right side of hero
    stats = [('4+', 'AI Providers'), ('32+', 'Endpoints'), ('>60%', 'Win Rate'), ('99.9%', 'SLA')]
    sx = W - MARGIN - 8
    for num, lbl in reversed(stats):
        nw = c.stringWidth(num, 'Helvetica-Bold', 16)
        lw = c.stringWidth(lbl, 'Helvetica', 7)
        total_w = nw + 4 + lw
        sx -= total_w
        text(c, sx, hero_y + 22, num, size=16, color=ORANGE, bold=True)
        text(c, sx + nw + 4, hero_y + 25, lbl, size=7, color=Color(1, 1, 1, alpha=0.5))
        if lbl != 'AI Providers':
            sx -= 16
            draw_dot(c, sx + 6, hero_y + 28, Color(1, 1, 1, alpha=0.2), 1.5)

    # ── PROBLEMS SECTION ── off-white bg
    y = hero_y - 14

    # Section title
    text(c, MARGIN, y, 'The Challenge with ', size=14, color=DARK, bold=True)
    tw2 = c.stringWidth('The Challenge with ', 'Helvetica-Bold', 14)
    text(c, MARGIN + tw2, y, 'Enterprise AI Today', size=14, color=ORANGE, bold=True)
    y -= 14
    text(c, MARGIN, y, 'Organizations face real barriers adopting AI at scale. Chat.AI was built to solve each one.', size=8, color=GREY)
    y -= 16

    problems = [
        ('Single-Vendor Lock-In', BLUE,
         'Tying your organization to one AI provider means you inherit their outages, pricing changes, and limitations.',
         'Chat.AI orchestrates OpenAI, Anthropic, Google, and Azure simultaneously with automatic failover.'),
        ('No Cost Visibility', ORANGE,
         'ChatGPT Team and Copilot charge flat per-seat fees with no transparency into actual token consumption.',
         'Per-call cost attribution, prepaid budgets, and overage alerts give you complete financial control.'),
        ('Hallucinations & Unreliable Output', TEAL,
         'Single-model responses have no checks or balances. Users can\'t tell when the AI is confidently wrong.',
         'AI Council deliberation across 3+ models reduces hallucinations by over 30% with consensus validation.'),
        ('Security & Compliance Gaps', CHARTREUSE,
         'Consumer AI tools lack tenant isolation, audit trails, and the controls required for regulated industries.',
         'Row-level security, stored-procedure-only data access, full audit logging, and HIPAA/SOC 2 architecture.'),
    ]

    card_w = W - 2 * MARGIN
    card_h = 52
    for title, border_color, problem, solution in problems:
        # Card with colored left border
        draw_card(c, MARGIN, y - card_h, card_w, card_h, radius=5, border_color=border_color, shadow=True)
        # Problem title
        text(c, MARGIN + 10, y - 12, title, size=9, color=DARK, bold=True)
        # Problem text
        wrap_text(c, MARGIN + 10, y - 24, problem, card_w - 20, size=7, color=GREY, leading=9)
        # Arrow + solution
        text(c, MARGIN + 10, y - 42, '\u2192 ', size=8, color=ORANGE, bold=True)
        arrow_w = c.stringWidth('\u2192 ', 'Helvetica-Bold', 8)
        wrap_text(c, MARGIN + 10 + arrow_w, y - 42, solution, card_w - 20 - arrow_w, size=7, color=DARK, leading=9, bold=True)
        y -= card_h + 6

    # ── START FEATURES (if room) ──
    y -= 6
    text(c, MARGIN, y, 'Platform ', size=14, color=DARK, bold=True)
    tw3 = c.stringWidth('Platform ', 'Helvetica-Bold', 14)
    text(c, MARGIN + tw3, y, 'Capabilities', size=14, color=ORANGE, bold=True)
    y -= 12
    text(c, MARGIN, y, 'Everything your organization needs to adopt AI securely and effectively.', size=8, color=GREY)


# ══════════════════════════════════════════════════
# PAGE 2: Features (6 in 2-col) + AI Council
# ══════════════════════════════════════════════════

def draw_page2(c):
    draw_page_bg(c, OFF_WHITE)
    draw_orange_top_line(c)

    y = H - 20

    features = [
        ('Multi-Model Chat', BLUE, [
            'GPT-4o, Claude Sonnet, Gemini Pro \u2014 switch mid-chat',
            'Streaming responses, file attachments & analysis',
            'Thread sharing with permissions, export & history',
        ]),
        ('Custom AI Assistants', ORANGE, [
            'Purpose-built assistants with custom instructions',
            'Per-assistant model selection and connectors (REST, GraphQL, DB)',
            'Shared or private assistants, MyGPTs for productivity',
        ]),
        ('Projects & Workspaces', TEAL, [
            'Organize chats, files, instructions into project containers',
            'Project-level system prompts inherited across all chats',
            'Granular view/edit sharing permissions',
        ]),
        ('Workflow Automation', DARK, [
            'Trigger N8N workflows from chat with parameter extraction',
            'Multi-step business process execution from conversation',
            'Run tracking, retry mechanisms, cost attribution',
        ]),
        ('Knowledge & Memory (RAG)', CHARTREUSE, [
            'Upload documents, vectorize for persistent org memory',
            'Weaviate, Pinecone, or Chroma vector stores per tenant',
            'Semantic similarity search and retrieval-augmented generation',
        ]),
        ('Billing & Cost Controls', BLUE, [
            'Prepaid balance model with token-level metering per model',
            'Budget alerts at 75% and 100%, automatic overage handling',
            'Per-user, per-conversation, per-model cost attribution',
        ]),
    ]

    col_w = (W - 2 * MARGIN - 16) / 2
    left_x = MARGIN
    right_x = MARGIN + col_w + 16

    for idx, (title, color, items) in enumerate(features):
        col = idx % 2
        row = idx // 2
        fx = left_x if col == 0 else right_x
        fy = y - row * 155

        # Colored circle icon
        draw_circle(c, fx + 10, fy - 4, 7, color)
        text_center(c, fx + 10, fy - 7, '\u2726', size=7, color=WHITE_C, bold=True)

        # Title
        text(c, fx + 22, fy - 8, title, size=9.5, color=DARK, bold=True)

        # Bullet items with teal checks
        by = fy - 22
        for item in items:
            draw_check(c, fx + 22, by, TEAL, 7)
            by = wrap_text(c, fx + 32, by, item, col_w - 38, size=7, color=GREY, leading=9)
            by -= 3

    # ── AI COUNCIL SECTION ── dark background block
    council_y = y - 3 * 155 - 10
    council_h = H - 20 - (y - 3 * 155 - 10) + council_y
    # Recalc: council block from council_y down to bottom margin
    council_top = y - 3 * 155 + 5
    council_bottom = 20
    council_h_actual = council_top - council_bottom

    draw_rect(c, 0, council_bottom, W, council_h_actual, DARK)
    # Decorative glows
    draw_decorative_glows(c, [
        (W - 80, council_bottom + council_h_actual - 20, 60, (0.12, 0.67, 0.78), 0.08),
        (60, council_bottom + 40, 50, (0.96, 0.49, 0.29), 0.06),
    ])

    cy = council_top - 16

    # PATENT PENDING badge
    c.setFillColor(Color(0.96, 0.49, 0.29, alpha=0.15))
    c.roundRect(MARGIN, cy + 2, 100, 14, 7, fill=1, stroke=0)
    text(c, MARGIN + 8, cy + 5, 'PATENT PENDING', size=7, color=ORANGE, bold=True)
    cy -= 16

    # Title
    text(c, MARGIN, cy, 'AI Council \u2014 ', size=16, color=WHITE_C, bold=True)
    tw = c.stringWidth('AI Council \u2014 ', 'Helvetica-Bold', 16)
    text(c, MARGIN + tw, cy, 'Collective Intelligence', size=16, color=ORANGE, bold=True)
    cy -= 14
    wrap_text(c, MARGIN, cy, 'The industry\'s first multi-agent deliberation system. Chat.AI convenes a council of 3+ diverse models through a structured 3-stage pipeline.', W - 2 * MARGIN, size=8, color=Color(0.65, 0.65, 0.68, alpha=1), leading=10)
    cy -= 26

    # 3 pipeline stages with orange numbered circles
    stages = [
        ('Divergence', 'Query routed to 3+ AI models in parallel. Each responds independently.'),
        ('Convergence', 'Responses anonymized and peer-reviewed with mandatory dissent.'),
        ('Synthesis', 'Highest-performing model synthesizes consensus using Bayesian weighting.'),
    ]
    stage_w = (W - 2 * MARGIN - 20) / 3
    for i, (title, desc) in enumerate(stages):
        sx = MARGIN + i * (stage_w + 10)
        # Orange numbered circle
        draw_circle(c, sx + 10, cy + 4, 9, ORANGE)
        text_center(c, sx + 10, cy + 1, str(i + 1), size=9, color=WHITE_C, bold=True)
        text(c, sx + 24, cy + 1, title, size=8.5, color=WHITE_C, bold=True)
        wrap_text(c, sx + 24, cy - 10, desc, stage_w - 30, size=7, color=Color(0.7, 0.7, 0.73, alpha=1), leading=9)

    cy -= 40

    # 4 metrics in teal
    metrics = [
        ('>60%', 'Win Rate vs\nSingle Models'),
        ('>30%', 'Hallucination\nReduction'),
        ('3+', 'Models Per\nDeliberation'),
        ('100%', 'Anonymized\nPeer Review'),
    ]
    metric_w = (W - 2 * MARGIN - 30) / 4
    for i, (val, lbl) in enumerate(metrics):
        mx = MARGIN + i * (metric_w + 10)
        # Dark card
        draw_rect(c, mx, cy - 30, metric_w, 40, Color(0.16, 0.16, 0.3, alpha=1), radius=4)
        text_center(c, mx + metric_w / 2, cy - 2, val, size=16, color=TEAL, bold=True)
        lines = lbl.split('\n')
        for li, line in enumerate(lines):
            text_center(c, mx + metric_w / 2, cy - 16 - li * 9, line, size=6.5, color=Color(0.53, 0.53, 0.56, alpha=1))
    cy -= 44

    # Supported models
    text(c, MARGIN, cy, 'Supported: ', size=7.5, color=Color(0.53, 0.53, 0.56, alpha=1), bold=True)
    sw = c.stringWidth('Supported: ', 'Helvetica-Bold', 7.5)
    text(c, MARGIN + sw, cy, 'GPT-4o  \u2022  Claude Sonnet  \u2022  Gemini Pro  \u2022  GPT-4o mini  \u2022  Claude Haiku  \u2022  Gemini Flash  \u2022  Claude Opus  \u2022  Azure OpenAI',
         size=7, color=Color(0.63, 0.63, 0.67, alpha=1))


# ══════════════════════════════════════════════════
# PAGE 3: Pricing + Comparison
# ══════════════════════════════════════════════════

def draw_page3(c):
    draw_page_bg(c, WHITE_C)
    # Teal top accent line
    draw_rect(c, 0, H - 2, W, 2, TEAL)

    y = H - 24

    # ── PRICING TABLE ──
    text(c, MARGIN, y, 'Simple, ', size=14, color=DARK, bold=True)
    tw = c.stringWidth('Simple, ', 'Helvetica-Bold', 14)
    text(c, MARGIN + tw, y, 'Transparent Pricing', size=14, color=ORANGE, bold=True)
    y -= 12
    text(c, MARGIN, y, 'No hidden fees. No per-seat surprises. Annual plans save 17%.', size=8, color=GREY)
    y -= 16

    # 5 columns x 13 rows
    headers = ['', 'Starter', 'Growth', 'Professional \u2605', 'Enterprise']
    rows = [
        ['Monthly Price',      '$299',    '$899',     '$1,699',        'Custom'],
        ['Annual Price',       '$2,990',  '$8,990',   '$16,990',       'Custom'],
        ['Users',              '10',      '50',       '100',           '250+'],
        ['Tokens/month',       '10M',     '35M',      '75M',          '200M+'],
        ['Storage',            '10 GB',   '50 GB',    '150 GB',       '500 GB+'],
        ['Custom Assistants',  '5',       '25',       'Unlimited',    'Unlimited'],
        ['Workflows',          '10',      '50',       'Unlimited',    'Unlimited'],
        ['AI Council',         '\u2014',  '\u2014',   '\u2713 Included', '\u2713 Included'],
        ['Model Tier',         'Standard', 'Pro',     'Pro + Premium', 'All Premium'],
        ['SLA',                '99.0%',   '99.5%',    '99.5%',        '99.9%+'],
        ['Support',            '72 hours', '48 hours', '24 hours',    '4\u20131 hour'],
        ['Overage',            'Hard cap', 'Auto top-up', 'Auto top-up', 'Custom'],
    ]

    col_widths = [100, 90, 90, 110, 110]  # total ~500
    table_w = sum(col_widths)
    table_x = (W - table_w) / 2
    row_h = 16
    hl_col = 3  # Professional column

    # Header row
    hx = table_x
    for ci, hdr in enumerate(headers):
        bg = BLUE if ci == hl_col else DARK
        draw_rect(c, hx, y - row_h, col_widths[ci], row_h, bg)
        text_center(c, hx + col_widths[ci] / 2, y - row_h + 4, hdr, size=7, color=WHITE_C, bold=True)
        hx += col_widths[ci]
    y -= row_h

    # Data rows
    for ri, row in enumerate(rows):
        hx = table_x
        for ci, val in enumerate(row):
            is_hl = (ci == hl_col)
            # Alternating tint
            if ri % 2 == 0:
                bg = LIGHT_BLUE if is_hl else Color(0.97, 0.97, 0.98, alpha=1)
            else:
                bg = Color(0.88, 0.93, 0.97, alpha=1) if is_hl else WHITE_C
            draw_rect(c, hx, y - row_h, col_widths[ci], row_h, bg)
            is_bold = (ci == 0 or is_hl)
            txt_color = DARK if (is_hl or ci == 0) else GREY
            text_center(c, hx + col_widths[ci] / 2, y - row_h + 4, val, size=6.5, color=txt_color, bold=is_bold)
            hx += col_widths[ci]
        y -= row_h

    # Add-ons
    y -= 8
    text(c, MARGIN, y, 'Add-Ons: ', size=7, color=DARK, bold=True)
    aw = c.stringWidth('Add-Ons: ', 'Helvetica-Bold', 7)
    text(c, MARGIN + aw, y, 'Premium Models ($299\u2013499/mo) \u2022 Additional Users ($20/user/mo) \u2022 Additional Storage ($0.15/GB/mo) \u2022 White-Label ($1,500/mo)',
         size=7, color=GREY)

    # ── COMPARISON TABLE ──
    y -= 24
    text(c, MARGIN, y, 'How Chat.AI ', size=14, color=DARK, bold=True)
    tw2 = c.stringWidth('How Chat.AI ', 'Helvetica-Bold', 14)
    text(c, MARGIN + tw2, y, 'Compares', size=14, color=ORANGE, bold=True)
    y -= 12
    text(c, MARGIN, y, 'Feature-for-feature, Chat.AI delivers more value at a lower total cost.', size=8, color=GREY)
    y -= 16

    comp_headers = ['Capability', 'ChatGPT Team', 'MS Copilot', 'Gemini Business', 'Chat.AI']
    comp_rows = [
        ['Multi-model access',       'OpenAI only',  'OpenAI only',     'Google only',   '\u2713 4+ providers'],
        ['AI Council (consensus)',    '\u2014',       '\u2014',          '\u2014',        '\u2713 Patent pending'],
        ['Per-call cost attribution', '\u2014',       '\u2014',          '\u2014',        '\u2713 Token-level'],
        ['Workflow automation',       'GPTs only',    'Power Automate ($)','\u2014',      '\u2713 N8N built-in'],
        ['Multi-tenant isolation',    '\u2014',       'Basic',           'Basic',         '\u2713 Row-level'],
        ['Full audit logging',        '\u2014',       'Limited',         'Limited',       '\u2713 SP-enforced'],
        ['Custom vector stores',      'OpenAI only',  '\u2014',          '\u2014',        '\u2713 Weaviate+'],
        ['On-premise deployment',     '\u2014',       '\u2014',          '\u2014',        '\u2713 Term license'],
        ['Budget controls & alerts',  '\u2014',       '\u2014',          '\u2014',        '\u2713 75%/100%'],
        ['Price (25 users)',          '$625/mo',      '$750/mo',         '$500/mo',       '$899/mo (50 seats)'],
    ]

    comp_col_widths = [120, 85, 85, 85, 110]
    comp_table_w = sum(comp_col_widths)
    comp_table_x = (W - comp_table_w) / 2
    comp_hl = 4  # Chat.AI column

    # Header
    cx = comp_table_x
    for ci, hdr in enumerate(comp_headers):
        bg = BLUE if ci == comp_hl else DARK
        draw_rect(c, cx, y - row_h, comp_col_widths[ci], row_h, bg)
        text_center(c, cx + comp_col_widths[ci] / 2, y - row_h + 4, hdr, size=6.5, color=WHITE_C, bold=True)
        cx += comp_col_widths[ci]
    y -= row_h

    for ri, row in enumerate(comp_rows):
        cx = comp_table_x
        for ci, val in enumerate(row):
            is_hl = (ci == comp_hl)
            if ri % 2 == 0:
                bg = LIGHT_BLUE if is_hl else Color(0.97, 0.97, 0.98, alpha=1)
            else:
                bg = Color(0.88, 0.93, 0.97, alpha=1) if is_hl else WHITE_C
            draw_rect(c, cx, y - row_h, comp_col_widths[ci], row_h, bg)

            # Color: teal for check marks in Chat.AI col, grey dash for others
            if is_hl:
                txt_color = TEAL
                is_bold = True
            elif ci == 0:
                txt_color = DARK
                is_bold = True
            elif val == '\u2014':
                txt_color = MED_GREY
                is_bold = False
            else:
                txt_color = GREY
                is_bold = False

            text_center(c, cx + comp_col_widths[ci] / 2, y - row_h + 4, val, size=6.5, color=txt_color, bold=is_bold)
            cx += comp_col_widths[ci]
        y -= row_h


# ══════════════════════════════════════════════════
# PAGE 4: ROI + Security + Deployment + CTA
# ══════════════════════════════════════════════════

def draw_page4(c):
    draw_page_bg(c, WHITE_C)
    draw_orange_top_line(c)

    # ── ROI SECTION ── dark background
    roi_h = 140
    roi_y = H - roi_h
    draw_rect(c, 0, roi_y, W, roi_h, DARK)
    draw_decorative_glows(c, [
        (W - 60, H - 30, 70, (0.12, 0.67, 0.78), 0.08),
        (80, roi_y + 20, 50, (0.96, 0.49, 0.29), 0.06),
    ])

    y = H - 24
    text(c, MARGIN, y, 'What You ', size=16, color=WHITE_C, bold=True)
    tw = c.stringWidth('What You ', 'Helvetica-Bold', 16)
    text(c, MARGIN + tw, y, 'Save', size=16, color=ORANGE, bold=True)
    y -= 12
    text(c, MARGIN, y, 'Real cost comparisons for typical organizations switching to Chat.AI.', size=8, color=Color(0.53, 0.53, 0.56, alpha=1))
    y -= 20

    # 3-column ROI cards
    roi_data = [
        ('Small Team', '10 users', '$389/mo \u2192 $299/mo', '$1,080/yr'),
        ('Mid-Market', '50 users', '$2,500/mo \u2192 $899/mo', '$19,212/yr'),
        ('Enterprise', '250 users', '$18,500/mo \u2192 $3,999/mo', '$174,012/yr'),
    ]
    roi_card_w = (W - 2 * MARGIN - 24) / 3
    for i, (label, users, cost, savings) in enumerate(roi_data):
        rx = MARGIN + i * (roi_card_w + 12)
        # Card
        draw_rect(c, rx, y - 60, roi_card_w, 72, Color(0.16, 0.16, 0.3, alpha=1), radius=5)
        # Top accent
        colors = [BLUE, ORANGE, TEAL]
        draw_rect(c, rx, y + 10, roi_card_w, 2, colors[i])
        # Content
        text_center(c, rx + roi_card_w / 2, y, label, size=9, color=WHITE_C, bold=True)
        text_center(c, rx + roi_card_w / 2, y - 12, users, size=7, color=Color(0.53, 0.53, 0.56, alpha=1))
        text_center(c, rx + roi_card_w / 2, y - 28, cost, size=7, color=Color(0.7, 0.7, 0.73, alpha=1))
        text_center(c, rx + roi_card_w / 2, y - 46, savings, size=14, color=CHARTREUSE, bold=True)

    # ── SECURITY SECTION ── white bg
    y = roi_y - 16
    text(c, MARGIN, y, 'Enterprise-Grade ', size=13, color=DARK, bold=True)
    tw2 = c.stringWidth('Enterprise-Grade ', 'Helvetica-Bold', 13)
    text(c, MARGIN + tw2, y, 'Security & Compliance', size=13, color=ORANGE, bold=True)
    y -= 12

    sec_items = [
        ('Azure AD SSO', BLUE, 'OAuth 2.0/OIDC with JWT bearer tokens and MSAL integration'),
        ('Row-Level Security', ORANGE, 'Database-enforced tenant isolation on every query via stored procedures'),
        ('Full Audit Trail', TEAL, 'Every data mutation logged with user, timestamp, action \u2014 exportable'),
        ('Azure Key Vault', DARK, 'All secrets and credentials encrypted at rest with automatic rotation'),
        ('Prompt Injection Defense', ORANGE, 'Salted sequence tags and delimiter separation protect against adversarial inputs'),
        ('HIPAA / SOC 2 Ready', BLUE, 'Architecture designed from the ground up for regulated industry compliance'),
        ('RBAC Roles', TEAL, 'Platform Admin, Client Admin, User hierarchies with granular permissions'),
        ('Rate Limiting', DARK, 'Per-user (100/min), per-tenant (1,000/min), per-endpoint throttling'),
    ]

    # 2-column layout for security items
    sec_col_w = (W - 2 * MARGIN - 16) / 2
    for idx, (title, color, desc) in enumerate(sec_items):
        col = idx % 2
        row = idx // 2
        sx = MARGIN if col == 0 else MARGIN + sec_col_w + 16
        sy = y - row * 28

        draw_dot(c, sx + 4, sy + 3, color, 3.5)
        text(c, sx + 12, sy, title, size=8, color=DARK, bold=True)
        wrap_text(c, sx + 12, sy - 10, desc, sec_col_w - 16, size=6.5, color=GREY, leading=8)

    y -= 4 * 28 + 10

    # ── DEPLOYMENT SECTION ──
    text(c, MARGIN, y, 'Flexible ', size=13, color=DARK, bold=True)
    tw3 = c.stringWidth('Flexible ', 'Helvetica-Bold', 13)
    text(c, MARGIN + tw3, y, 'Deployment', size=13, color=ORANGE, bold=True)
    y -= 14

    # 2 deployment cards
    dep_card_w = (W - 2 * MARGIN - 16) / 2

    # Cloud SaaS card
    draw_card(c, MARGIN, y - 52, dep_card_w, 52, radius=5, border_color=BLUE, shadow=True)
    text(c, MARGIN + 10, y - 12, '\u2601 Cloud (SaaS)', size=10, color=BLUE, bold=True)
    wrap_text(c, MARGIN + 10, y - 24, 'Fully managed on Azure. Automatic updates, monitoring, 24/7 support. Zero infrastructure overhead.', dep_card_w - 20, size=7, color=GREY, leading=9)
    text(c, MARGIN + 10, y - 46, 'Starting at $299/mo', size=8, color=BLUE, bold=True)

    # On-Premise card
    op_x = MARGIN + dep_card_w + 16
    draw_card(c, op_x, y - 52, dep_card_w, 52, radius=5, border_color=DARK, shadow=True)
    text(c, op_x + 10, y - 12, '\U0001F3E0 On-Premise (Term License)', size=10, color=DARK, bold=True)
    wrap_text(c, op_x + 10, y - 24, 'Deploy inside your infrastructure. Maximum data sovereignty. Ideal for regulated industries.', dep_card_w - 20, size=7, color=GREY, leading=9)
    text(c, op_x + 10, y - 46, 'Starting at $45,000/yr (50 users)', size=8, color=DARK, bold=True)

    y -= 66

    # ── BLUE CTA BANNER ──
    cta_h = 52
    cta_y = y - cta_h
    draw_rect(c, 0, cta_y, W, cta_h, BLUE)
    c.setFillColor(Color(0, 0.2, 0.4, alpha=0.3))
    c.rect(W * 0.65, cta_y, W * 0.35, cta_h, fill=1, stroke=0)

    text_center(c, W / 2, cta_y + 32, 'Ready to Transform How Your Team Uses AI?', size=14, color=WHITE_C, bold=True)
    text_center(c, W / 2, cta_y + 18, 'Schedule a personalized demo and see how Chat.AI delivers better answers at a lower cost.', size=8, color=Color(0.8, 0.87, 0.93, alpha=1))
    text_center(c, W / 2, cta_y + 6, 'RJain@technijian.com  \u2022  949.379.8500', size=10, color=WHITE_C, bold=True)

    # ── DARK FOOTER ──
    footer_h = 18
    draw_rect(c, 0, 0, W, footer_h, NEAR_BLACK)
    text_center(c, W / 2, 5,
                '\u00A9 2026 Technijian  \u2022  18 Technology Dr. Ste 141, Irvine CA 92618  \u2022  technijian.com',
                size=6.5, color=Color(1, 1, 1, alpha=0.4))


# ══════════════════════════════════════════════════
# BUILD PDF
# ══════════════════════════════════════════════════

def build_pdf():
    pdf_path = os.path.join(OUT_DIR, 'Chat.AI Marketing Datasheet.pdf')
    c = canvas.Canvas(pdf_path, pagesize=letter)
    c.setTitle('Technijian Chat.AI \u2014 Marketing Datasheet')
    c.setAuthor('Technijian')

    logo_data = fetch_logo()

    # Page 1
    draw_page1(c, logo_data)
    c.showPage()

    # Page 2
    draw_page2(c)
    c.showPage()

    # Page 3
    draw_page3(c)
    c.showPage()

    # Page 4
    draw_page4(c)

    c.save()
    print(f"  PDF: {pdf_path}")
    return pdf_path


# ══════════════════════════════════════════════════
# BUILD DOCX (secondary output for email attachment)
# ══════════════════════════════════════════════════

def build_docx():
    """Generate DOCX version using python-docx for email attachment use."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    BLUE_RGB = RGBColor(0x00, 0x6D, 0xB6)
    ORANGE_RGB = RGBColor(0xF6, 0x7D, 0x4B)
    TEAL_RGB = RGBColor(0x1E, 0xAA, 0xC8)
    DARK_RGB = RGBColor(0x1A, 0x1A, 0x2E)
    GREY_RGB = RGBColor(0x59, 0x59, 0x5B)
    WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
    CHARTREUSE_RGB = RGBColor(0xCB, 0xDB, 0x2D)

    def _font(r, name='Open Sans'):
        r.font.name = name
        rPr = r._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:cs="{name}"/>')
            rPr.insert(0, rFonts)

    def shade(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def cell_margins(cell, top=50, bottom=50, left=90, right=90):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'  <w:top w:w="{top}" w:type="dxa"/>'
            f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
            f'  <w:start w:w="{left}" w:type="dxa"/>'
            f'  <w:end w:w="{right}" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        old = tcPr.find(qn('w:tcMar'))
        if old is not None:
            tcPr.remove(old)
        tcPr.append(tcMar)

    def remove_borders(table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            tbl.insert(0, tblPr)
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            + ''.join(f'<w:{s} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                      for s in ['top','left','bottom','right','insideH','insideV'])
            + '</w:tblBorders>'
        )
        old = tblPr.find(qn('w:tblBorders'))
        if old is not None:
            tblPr.remove(old)
        tblPr.append(borders)

    def add_run(p, txt, bold=False, color=None, size=None, font_name='Open Sans'):
        r = p.add_run(txt)
        _font(r, font_name)
        if bold: r.bold = True
        if color: r.font.color.rgb = color
        if size: r.font.size = Pt(size)
        return r

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(1.8)
        sec.right_margin = Cm(1.8)

    style = doc.styles['Normal']
    style.font.name = 'Open Sans'
    style.font.size = Pt(10)
    style.font.color.rgb = GREY_RGB

    # Hero
    hero = doc.add_table(rows=1, cols=1)
    hero.alignment = WD_TABLE_ALIGNMENT.CENTER
    hc = hero.rows[0].cells[0]
    hc.text = ''
    shade(hc, "1A1A2E")
    cell_margins(hc, 300, 300, 300, 300)
    remove_borders(hero)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    hero._tbl.tblPr.append(tblW)

    hp = hc.paragraphs[0]
    try:
        resp = requests.get("https://technijian.com/wp-content/uploads/2023/08/Logo.jpg", timeout=5)
        if resp.status_code == 200:
            r = hp.add_run()
            r.add_picture(BytesIO(resp.content), width=Inches(1.6))
            hp.add_run('\n')
    except:
        pass

    add_run(hp, '\nENTERPRISE AI PLATFORM\n', bold=True, color=TEAL_RGB, size=9)
    hp2 = hc.add_paragraph()
    add_run(hp2, 'Technijian ', bold=True, color=WHITE_RGB, size=30)
    add_run(hp2, 'Chat.AI', bold=True, color=ORANGE_RGB, size=30)

    hp3 = hc.add_paragraph()
    hp3.space_after = Pt(16)
    add_run(hp3, 'Secure, multi-model AI with collective intelligence, workflow automation, '
        'and enterprise governance.', color=RGBColor(0xB0, 0xB0, 0xB8), size=10.5)

    doc.add_paragraph().space_after = Pt(6)

    # Problems
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(2)
    add_run(p, 'The Challenge with ', bold=True, color=DARK_RGB, size=18)
    add_run(p, 'Enterprise AI Today', bold=True, color=ORANGE_RGB, size=18)

    for title, color, problem, solution in [
        ('Single-Vendor Lock-In', BLUE_RGB,
         'Tying your organization to one AI provider means you inherit their outages, pricing changes, and limitations.',
         'Chat.AI orchestrates OpenAI, Anthropic, Google, and Azure simultaneously with automatic failover.'),
        ('No Cost Visibility', ORANGE_RGB,
         'ChatGPT Team and Copilot charge flat per-seat fees with no transparency into actual token consumption.',
         'Per-call cost attribution, prepaid budgets, and overage alerts give you complete financial control.'),
        ('Hallucinations & Unreliable Output', TEAL_RGB,
         'Single-model responses have no checks. Users can\'t tell when the AI is confidently wrong.',
         'AI Council deliberation across 3+ models reduces hallucinations by over 30%.'),
        ('Security & Compliance Gaps', DARK_RGB,
         'Consumer AI tools lack tenant isolation, audit trails, and controls for regulated industries.',
         'Row-level security, stored-procedure-only data access, full audit logging, HIPAA/SOC 2 architecture.'),
    ]:
        bp = doc.add_paragraph()
        bp.space_after = Pt(1)
        bp.space_before = Pt(0)
        add_run(bp, title, bold=True, color=color, size=11.5)
        bp2 = doc.add_paragraph()
        bp2.space_after = Pt(1)
        add_run(bp2, problem, color=GREY_RGB, size=9.5)
        bp3 = doc.add_paragraph()
        bp3.space_after = Pt(10)
        add_run(bp3, '\u2192 ', bold=True, color=ORANGE_RGB, size=10)
        add_run(bp3, solution, color=DARK_RGB, size=9.5, bold=True)

    # Features
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(2)
    add_run(p, 'Platform ', bold=True, color=DARK_RGB, size=18)
    add_run(p, 'Capabilities', bold=True, color=ORANGE_RGB, size=18)

    for title, color, bullets in [
        ('Multi-Model Chat', BLUE_RGB, ['GPT-4o, Claude Sonnet, Gemini Pro \u2014 switch mid-chat', 'Streaming responses, file attachments & analysis', 'Thread sharing with permissions, export & history']),
        ('Custom AI Assistants', ORANGE_RGB, ['Purpose-built assistants with custom instructions', 'Per-assistant model selection and connectors', 'Shared or private assistants, MyGPTs']),
        ('Projects & Workspaces', TEAL_RGB, ['Organize chats, files, instructions into containers', 'Project-level system prompts inherited across chats', 'Granular view/edit sharing permissions']),
        ('Workflow Automation', DARK_RGB, ['Trigger N8N workflows from chat', 'Multi-step business process execution', 'Run tracking, retry, cost attribution']),
        ('Knowledge & Memory (RAG)', RGBColor(0x8B, 0xA8, 0x2E), ['Upload docs, vectorize for persistent org memory', 'Weaviate, Pinecone, or Chroma per tenant', 'Semantic search and RAG']),
        ('Billing & Cost Controls', BLUE_RGB, ['Prepaid balance with token-level metering', 'Budget alerts at 75%/100%, auto overage', 'Per-user/conversation/model cost attribution']),
    ]:
        fp = doc.add_paragraph()
        fp.space_after = Pt(2)
        add_run(fp, title, bold=True, color=color, size=11.5)
        for b in bullets:
            bp = doc.add_paragraph(style='List Bullet')
            bp.space_after = Pt(1)
            bp.space_before = Pt(0)
            bp.clear()
            add_run(bp, b, color=GREY_RGB, size=9.5)
        doc.add_paragraph().space_after = Pt(4)

    # CTA
    cta = doc.add_table(rows=1, cols=1)
    cta.alignment = WD_TABLE_ALIGNMENT.CENTER
    ctc = cta.rows[0].cells[0]
    ctc.text = ''
    shade(ctc, "006DB6")
    cell_margins(ctc, 300, 300, 200, 200)
    remove_borders(cta)
    tblW4 = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    cta._tbl.tblPr.append(tblW4)
    cp = ctc.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cp, 'Ready to Transform How Your Team Uses AI?', bold=True, color=WHITE_RGB, size=20)
    cp2 = ctc.add_paragraph()
    cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cp2, 'RJain@technijian.com', bold=True, color=WHITE_RGB, size=13)
    add_run(cp2, '  \u2022  ', color=RGBColor(0xCC, 0xDD, 0xEE), size=13)
    add_run(cp2, '949.379.8500', bold=True, color=WHITE_RGB, size=13)

    # Footer
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(fp, '\u00A9 2026 Technijian  \u2022  18 Technology Dr. Ste 141, Irvine CA 92618  \u2022  technijian.com',
            color=GREY_RGB, size=8)

    return doc


# ══════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════

def main():
    print("Building Chat.AI Marketing Datasheet...")

    # Primary output: ReportLab PDF
    build_pdf()

    # Secondary output: DOCX for email attachment
    try:
        doc = build_docx()
        docx_path = os.path.join(OUT_DIR, 'Chat.AI Marketing Datasheet.docx')
        doc.save(docx_path)
        print(f"  DOCX: {docx_path}")
    except ImportError:
        print("  DOCX: skipped (python-docx not installed)")
    except Exception as e:
        print(f"  DOCX: error - {e}")

    print("\nDone!")


if __name__ == '__main__':
    main()
