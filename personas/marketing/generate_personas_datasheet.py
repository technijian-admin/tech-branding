"""
Technijian Customer Personas Datasheet Generator
Produces a brand-compliant DOCX + PDF for the marketing team.

Covers 13 personas (10 vertical + 3 horizontal) with:
- Pain points, triggers, primary services, Year-1 ARR
- Objections, decision criteria, messaging headlines
- Killer verbatim quotes, watering holes, campaign playbook

Brand: Technijian - Blue #006DB6, Orange #F67D4B, Teal #1EAAC8, Dark #1A1A2E
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import requests
import os

# ==========================================================================
# Brand colors (hex strings for XML shading + RGBColor for fonts)
# ==========================================================================
C_BLUE       = '006DB6'
C_ORANGE     = 'F67D4B'
C_TEAL       = '1EAAC8'
C_CHARTREUSE = 'CBDB2D'
C_DARK       = '1A1A2E'
C_NEAR_BLACK = '2D2D2D'
C_GREY       = '59595B'
C_OFFWHITE   = 'F8F9FA'
C_LIGHTGREY  = 'E9ECEF'
C_WHITE      = 'FFFFFF'
C_DARKBLUE   = '0d2847'
C_MIDBLUE    = '004d80'

RGB_BLUE       = RGBColor(0x00, 0x6D, 0xB6)
RGB_ORANGE     = RGBColor(0xF6, 0x7D, 0x4B)
RGB_TEAL       = RGBColor(0x1E, 0xAA, 0xC8)
RGB_CHARTREUSE = RGBColor(0xCB, 0xDB, 0x2D)
RGB_DARK       = RGBColor(0x1A, 0x1A, 0x2E)
RGB_GREY       = RGBColor(0x59, 0x59, 0x5B)
RGB_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
RGB_BLACK      = RGBColor(0x1A, 0x1A, 0x2E)

LOGO_URL = "https://technijian.com/wp-content/uploads/2023/08/Logo.jpg"


# ==========================================================================
# XML Helpers
# ==========================================================================
def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def shade_paragraph(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def set_cell_borders(cell, top=None, left=None, bottom=None, right=None):
    """Each arg is dict like {'sz': '8', 'val': 'single', 'color': '006DB6'} or None to clear."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge, opts in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        tag = f'w:{edge}'
        el = tcBorders.find(qn(tag))
        if el is not None:
            tcBorders.remove(el)
        el = OxmlElement(tag)
        if opts is None:
            el.set(qn('w:val'), 'nil')
        else:
            el.set(qn('w:val'), opts.get('val', 'single'))
            el.set(qn('w:sz'), str(opts.get('sz', '4')))
            el.set(qn('w:color'), opts.get('color', 'auto'))
        tcBorders.append(el)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_cell_margins(cell, top=100, left=120, bottom=100, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val in [('top', top), ('left', left), ('bottom', bottom), ('right', right)]:
        node = OxmlElement(f'w:{edge}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_col_widths(table, widths_inches):
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_inches):
                cell.width = Inches(widths_inches[idx])


def set_row_height(row, height_pt, rule='atLeast'):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_pt * 20)))
    trHeight.set(qn('w:hRule'), rule)
    trPr.append(trHeight)


def set_row_cant_split(row):
    """Prevent row from splitting across page breaks."""
    trPr = row._tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    trPr.append(cantSplit)


def add_run(paragraph, text, *, bold=False, italic=False, size=10, color=None, font='Open Sans'):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for k in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
        rFonts.set(qn(f'w:{k}'), font)
    return run


def set_paragraph_spacing(paragraph, before=0, after=4, line=1.15):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_horizontal_bar(paragraph, hex_color='006DB6', height_pt=3):
    """Add a thin colored horizontal bar via paragraph border-bottom."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(height_pt * 8)))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ==========================================================================
# Persona Data
# ==========================================================================
PERSONAS = [
    {
        'num': '01', 'name': 'Mark Tanaka', 'tier': 'Tier 2',
        'title': 'SMB Owner / CEO / Founder',
        'industry': 'Orange County Professional Services SMB',
        'size': '20-100 emp | $5M-$30M rev',
        'arr': '$30K-$90K Yr-1',
        'pains': [
            'One-person MSP is single-point-of-failure; no backup when "IT guy" is unavailable.',
            'No asset inventory, no passwords in one place - institutional knowledge lives in one head.',
            'Cyber insurance questionnaires getting longer; faking answers because he doesn\'t know.',
        ],
        'triggers': [
            'Outage that stopped revenue for half a day',
            'Insurance carrier imposes cybersecurity requirements',
            'Acquisition interest / exit-prep diligence',
        ],
        'services': 'My IT, My Office, My Security, My Continuity',
        'objections': [
            '"It\'s too expensive - my current guy is $X"',
            '"I don\'t want to sign a long-term contract"',
        ],
        'criteria': [
            'Single accountable team (pod model)',
            'Local OC presence, same-time-zone techs',
            '60-day out-clause, transparent monthly invoice',
        ],
        'headline': 'Run your business. We\'ll run your IT.',
        'quote': 'I just want to call one number. I don\'t want to figure out if it\'s a printer problem or a network problem or a Microsoft problem.',
        'holes': 'Vistage / EO peer groups | OC Business Journal events | Chamber of Commerce (Irvine, Newport, Anaheim)',
    },
    {
        'num': '02', 'name': 'Angela Santos', 'tier': 'Tier 1',
        'title': 'Practice Administrator / COO',
        'industry': 'Healthcare (multi-site clinic / medical group)',
        'size': '25-300 emp | $5M-$60M rev',
        'arr': '$60K-$180K Yr-1',
        'pains': [
            'EHR downtime during clinic halts documentation and patient flow.',
            'OCR breach-notification nightmare: PHI in sent folder, unclear 60-day clock.',
            'HIPAA training and phishing awareness on autopilot with no real adherence.',
        ],
        'triggers': [
            'Cyber insurance renewal tightens HIPAA requirements',
            'PE acquisition due diligence reveals IT gaps',
            'Staff email compromise / peer ransomware event',
        ],
        'services': 'My IT + HIPAA, My Compliance HIPAA, My Continuity, Chat.AI Healthcare',
        'objections': [
            '"Do you have healthcare-specific references?"',
            '"Our EHR vendor won\'t let anyone else touch integrations"',
        ],
        'criteria': [
            'Signed BAA pre-engagement, 24/7 MDR with IR SLA',
            'Healthcare-specific customer references',
            'HIPAA-fluent vCISO / vCCO on retainer',
        ],
        'headline': 'HIPAA-fluent managed IT. No more translating between your MSP and your attorney.',
        'quote': 'I want the docs to not know anything about IT. I want them to see patients. That\'s my job.',
        'holes': 'MGMA Annual (Oct) | OCMA chapter | HIPAA Journal | Becker\'s ASC Review',
    },
    {
        'num': '03', 'name': "Brian O'Connell", 'tier': 'Tier 1',
        'title': 'Chief Compliance Officer / Ops Principal',
        'industry': 'Financial Services (Broker-Dealer, RIA, Wealth)',
        'size': '15-500 emp | $3M-$100M AUM',
        'arr': '$75K-$250K Yr-1',
        'pains': [
            'Advisor off-channel texting/iMessage compliant on paper, uncontrolled in reality.',
            'No AI governance - advisors paste client data into ChatGPT unmonitored.',
            'Vendor risk program is manual spreadsheets with overdue Tier-1 reviews.',
        ],
        'triggers': [
            'Upcoming FINRA / SEC exam or deficiency letter',
            'Peer enforcement action announced',
            'M&A integration or Reg S-P amendment deadline',
        ],
        'services': 'My Security (SOC+SIEM+EDR), My Compliance vCCO, My Continuity (WORM)',
        'objections': [
            '"Do you have BD/RIA references?"',
            '"How do you integrate with Smarsh / Global Relay?"',
        ],
        'criteria': [
            'SOC 2 Type II on vendor (non-negotiable)',
            '15-min IR SLA + 72-hour Reg S-P playbook',
            'WORM archiving integration (Smarsh/Global Relay)',
        ],
        'headline': 'Compliance-fluent IT for broker-dealers, RIAs, and wealth managers.',
        'quote': 'I\'ll spend 20% more for the firm that can handle a 72-hour Reg S-P clock. That calculus is not close.',
        'holes': 'NSCP Annual (Oct) | Schwab IMPACT (Nov) | IA Watch newsletter',
    },
    {
        'num': '04', 'name': 'Greg Mahoney', 'tier': 'Tier 1',
        'title': 'COO / VP Operations / GM',
        'industry': 'Defense / Precision Manufacturing (DIB)',
        'size': '40-400 emp | $10M-$150M rev',
        'arr': '$90K-$300K Yr-1',
        'pains': [
            'CMMC 2.0 Level 2 - 110 NIST 800-171 controls MSP can\'t read.',
            'CUI scoping undefined across ERP, CAD, email, shop floor.',
            'Legacy ERP on Server 2012 that can\'t patch but must segment.',
        ],
        'triggers': [
            'Prime contractor requires CMMC L2 on new RFP',
            'Contract loss risk without certification',
            'DCSA / DCMA inspection scheduled',
        ],
        'services': 'My Compliance CMMC, My Cloud GCC High, My Security (CMMC SOC)',
        'objections': [
            '"CMMC will kill us, we can\'t afford it"',
            '"Our current MSP says they can do CMMC"',
        ],
        'criteria': [
            'CMMC / NIST 800-171 depth with sample SSP',
            'DIB client references (existing prime flowdown)',
            'GCC High migration experience',
        ],
        'headline': 'CMMC Level 2 readiness on your timeline - without losing your prime.',
        'quote': 'My #1 customer just sent a flowdown saying I need CMMC Level 2. I have 14 months. My MSP doesn\'t know what an SSP is.',
        'holes': 'NDIA | AIA / AUSA events | CMMC Day / CMMC-AB events',
    },
    {
        'num': '05', 'name': 'Priya Sharma', 'tier': 'Tier 1',
        'title': 'CTO / Founding Engineer / VP Eng',
        'industry': 'B2B SaaS Startup (Seed - Series B)',
        'size': '15-150 emp | $500K-$20M ARR',
        'arr': '$40K-$150K Yr-1',
        'pains': [
            'SOC 2 Type II on critical path - enterprise deals blocked.',
            'Security questionnaires eat 2-3 weeks per deal, no repeatable answers.',
            'No dedicated security engineer - she IS the SOC plus her CTO job.',
        ],
        'triggers': [
            'Enterprise deal gated on SOC 2 Type II',
            'Series A / B funding close approaching',
            'First regulated customer (HIPAA/PCI/FedRAMP)',
        ],
        'services': 'My Compliance SOC 2, My Security 24/7 MDR, My Cloud (AWS/Azure)',
        'objections': [
            '"Vanta / Drata handles this, why do I need you?"',
            '"Show me your own SOC 2"',
        ],
        'criteria': [
            'Technijian\'s SOC 2 posture (TPX reliance + CUEC mapping) on file',
            'AWS / Azure Well-Architected depth (IAM, KMS)',
            'Startup-friendly commercial terms, evidence automation',
        ],
        'headline': 'SOC 2 Type II in 90-120 days. 24/7 MDR. Audit-friendly evidence automation.',
        'quote': 'I have three enterprise deals waiting on SOC 2 Type II. Sales wants it yesterday. I\'m one person on security.',
        'holes': 'AWS re:Invent (Dec) | Pragmatic Engineer newsletter | Hacker News / Rands Slack',
    },
    {
        'num': '06', 'name': 'Catherine Weiss', 'tier': 'Tier 1',
        'title': 'Managing Partner / Firm Administrator',
        'industry': 'Law Firm (litigation, corporate, IP)',
        'size': '40-400 total | $10M-$100M rev',
        'arr': '$60K-$200K Yr-1',
        'pains': [
            'Attorneys use personal email for transfers, leaking client-confidential data.',
            'DMS is a nightmare - 20 yrs of docs, broken search, attorneys save to desktop.',
            'Corporate-client security questionnaires consume firm admin + outside IT.',
        ],
        'triggers': [
            'Lateral partner group onboarding (data migration)',
            'Corporate client cyber addendum signed',
            'Wire-fraud attempt or peer ransomware',
        ],
        'services': 'My IT + Office, My Security, My Compliance Guardian + Vault',
        'objections': [
            '"Our DMS vendor says to use their IT partner only"',
            '"The partners will never agree to change"',
        ],
        'criteria': [
            'Legal-industry references and NDA pre-access',
            'DMS experience (iManage / NetDocs / Worldox)',
            'Duty-of-technology-competence narrative',
        ],
        'headline': 'IT built for the duty of technology competence.',
        'quote': 'I\'ll pay for quiet. I\'ll pay double for quiet. What I won\'t pay for is excuses.',
        'holes': 'ILTACON (Aug) | ABA Techshow (Mar) | ALA / OCBA chapters',
    },
    {
        'num': '07', 'name': 'Tom Anderson', 'tier': 'Tier 2',
        'title': 'Broker of Record / Owner',
        'industry': 'Residential Real Estate / Property Mgmt',
        'size': '15-250 agents | $5M-$60M GCI',
        'arr': '$25K-$100K Yr-1',
        'pains': [
            'Zillow / paid lead channels expensive and declining in conversion.',
            'Agents ignore the CRM - paid $60K/yr for KVCore, half don\'t log in.',
            'Tech stack sprawl: 8-15 disconnected tools, website on SEO page 3.',
        ],
        'triggers': [
            'Lost a top agent to a better-tech brokerage',
            'NAR settlement fallout changing commissions',
            'Competitor launches AI-powered service',
        ],
        'services': 'My AI Lead Gen, My SEO, My AI Real Estate',
        'objections': [
            '"Will my agents actually use this?"',
            '"I already have a CRM / marketing firm"',
        ],
        'criteria': [
            'Lead volume impact visible in 60 days',
            'Agent-friendly simplicity (adoption-proof UI)',
            'Month-to-month commercial terms',
        ],
        'headline': 'More listings. Cheaper per lead. On autopilot.',
        'quote': 'Zillow is killing me. I\'m paying $8K a month and conversion is dropping. I need something that works.',
        'holes': 'Inman Connect (Jan/Aug) | NAR Annual (Nov) | Lab Coat Agents FB group',
    },
    {
        'num': '08', 'name': 'Denise Washington', 'tier': 'Tier 3',
        'title': 'Executive Director / COO / CEO',
        'industry': 'Non-Profit (social services / faith-based)',
        'size': '10-100 emp | $2M-$25M budget',
        'arr': '$18K-$60K Yr-1',
        'pains': [
            'Old hardware everywhere - 6-10 yr donated PCs, refresh only on failure.',
            'Donor database is a mess: duplicates, broken imports, merged records.',
            'No one owns IT - volunteer board member, part-time coord, or MSP junior.',
        ],
        'triggers': [
            'Major grant with IT modernization component',
            'Capital campaign kickoff',
            'Audit finding on IT controls',
        ],
        'services': 'My Office (M365 Non-Profit), My IT (NP SKU), My Cloud',
        'objections': [
            '"We can\'t afford it"',
            '"My board will never approve this"',
        ],
        'criteria': [
            'Non-profit pricing / M365 NP discount awareness',
            'Sector references (other NPs same size)',
            'Data-privacy posture for donor / beneficiary data',
        ],
        'headline': 'Mission-grade IT at non-profit pricing.',
        'quote': 'We had a phishing email a year ago. One of our program managers almost wired $28K to a fake invoice. Keeps me up.',
        'holes': 'CalNonprofits Annual | NTEN Nonprofit Technology Conf | Chronicle of Philanthropy',
    },
    {
        'num': '09', 'name': 'Antonio DiMarco', 'tier': 'Tier 3',
        'title': 'Owner / Multi-Unit Operator',
        'industry': 'Hospitality (Restaurants / QSR / Boutique Hotels)',
        'size': '50-500 emp | 3-30 locations',
        'arr': '$20K-$70K Yr-1',
        'pains': [
            'Google Maps / SEO / reviews fragmented across locations.',
            'PCI is a mystery - POS vendor claims coverage, merchant controls undocumented.',
            'POS outages during Friday peak cost $12K+ in lost covers per incident.',
        ],
        'triggers': [
            'New unit opening / grand opening push',
            'POS outage at peak service',
            'Credit-card breach at peer restaurant',
        ],
        'services': 'My SEO, My Security (PCI-aligned), My IT (multi-unit)',
        'objections': [
            '"My POS vendor does all this"',
            '"SEO is a scam"',
        ],
        'criteria': [
            'Restaurant / hospitality customer references',
            'Off-peak deployment windows (night / Sun AM)',
            'Month-to-month pricing, no long lockup',
        ],
        'headline': 'More covers. Cleaner PCI. Zero nights lost.',
        'quote': 'If it puts asses in seats, I pay. If it doesn\'t, I don\'t care what it costs.',
        'holes': 'National Restaurant Show (Chicago, May) | CRA local chapter | Nation\'s Restaurant News',
    },
    {
        'num': '10', 'name': 'Javier Torres', 'tier': 'Tier 2',
        'title': 'President / Owner / VP Ops',
        'industry': 'Construction (Luxury Builder / Commercial GC)',
        'size': '20-200 emp | $15M-$250M rev',
        'arr': '$40K-$150K Yr-1',
        'pains': [
            'Luxury buyer lead-gen inconsistent - realtor referrals and stale site.',
            'Site / trailer Wi-Fi patchy; superintendent can\'t open RFI on site.',
            'Plan files scattered across SharePoint, email, Dropbox, site tablets.',
        ],
        'triggers': [
            'New community / development launch',
            'Stolen trailer laptop with plan sets',
            'Wire-fraud attempt at a closing',
        ],
        'services': 'My AI Lead Gen, My IT (office + site), My Cloud (Azure AD)',
        'objections': [
            '"Builders don\'t need fancy IT"',
            '"Procore is all the software I need"',
        ],
        'criteria': [
            'Construction / builder customer references',
            'Procore / Sage fluency in delivery team',
            'Willingness to do site visits (not just office)',
        ],
        'headline': 'Luxury buyer leads on autopilot. Site + office IT that actually works.',
        'quote': 'A laptop went missing from a trailer last month. I don\'t know what was on it. I don\'t know who to call.',
        'holes': 'International Builders\' Show (Feb) | HBA of OC | Builder Magazine',
    },
    {
        'num': '11', 'name': 'Sarah Okafor', 'tier': 'Tier 1',
        'title': 'IT Director / VP IT (Co-Managed)',
        'industry': 'Cross-Industry Mid-Market',
        'size': '100-800 emp | $30M-$500M rev',
        'arr': '$60K-$250K Yr-1',
        'pains': [
            '24/7 coverage on a 9-to-6 team - on-call burns engineers, drives turnover.',
            'Specialist skills she can\'t hire for (Azure, security, K8s) - FTE ROI doesn\'t pencil.',
            'Project burst consumes 50% of team for 6 months; BAU slips.',
        ],
        'triggers': [
            'Major project (cloud mig / ERP / SOC 2 / M&A)',
            'Key engineer resigns without warning',
            'CFO mandates 24/7 coverage',
        ],
        'services': 'My IT Co-Managed, My Security (SOC), My Cloud specialists',
        'objections': [
            '"I don\'t want a vendor replacing my team"',
            '"I\'ve been burned by MSPs before"',
        ],
        'criteria': [
            'Co-managed model - extend not replace',
            'RACI documented pre-contract',
            'Technijian\'s SOC 2 posture (TPX reliance + CUEC mapping)',
        ],
        'headline': 'Extend your team - don\'t replace it. Co-managed IT that respects your ops.',
        'quote': 'My senior engineer is going to quit if he\'s on-call one more weekend. I can\'t hire fast enough. I need a partner, not a replacement.',
        'holes': 'SIM local chapter | Gartner IT Symposium | ISACA / ISSA OC chapter',
    },
    {
        'num': '12', 'name': 'Jennifer Hayes', 'tier': 'Tier 3',
        'title': 'Director of Marketing / VP Growth',
        'industry': 'Cross-Industry (RE, Hospitality, Construction, B2B, SaaS)',
        'size': '20-300 emp | $5M-$100M rev',
        'arr': '$30K-$120K Yr-1',
        'pains': [
            'Paid-media inflation - CPMs rising, CPL following, budget flat.',
            'SEO agency disappointment - $2-5K/mo for years, little rank movement.',
            'Content bottleneck - one in-house writer, AI tools hallucinate.',
        ],
        'triggers': [
            'New product / location launch',
            'CEO declares "year of AI"',
            'Agency contract expiration',
        ],
        'services': 'My SEO, My AI Lead Gen, Chat.AI, My Dev',
        'objections': [
            '"I already have an SEO agency"',
            '"How are you different from every SEO firm?"',
        ],
        'criteria': [
            'Industry case studies with real numbers',
            'Transparent monthly results (rank / leads / CPL)',
            'SLA on deliverables and response',
        ],
        'headline': 'Unlimited SEO. AI-enriched leads. Done-for-you content - without another agency.',
        'quote': 'We\'ve been with the same SEO agency for 3 years. Our organic traffic is flat. The reports look beautiful. Nothing is moving.',
        'holes': 'Inbound HubSpot (Sep) | Content Marketing World (Sep) | MarketingProfs',
    },
    {
        'num': '13', 'name': 'Terry Williams', 'tier': 'Tier 2',
        'title': 'CFO / Controller / VP Finance',
        'industry': 'Cross-Industry Mid-Market Economic Buyer',
        'size': '50-800 emp | $10M-$500M rev',
        'arr': 'Varies (signs all >$50K)',
        'pains': [
            'Cyber insurance renewal pain - longer questionnaires, rising premiums.',
            'Audit evidence - auditors ask for IT-control evidence he has to chase.',
            'IT spend transparency - SaaS / vendor / license sprawl hard to optimize.',
        ],
        'triggers': [
            'Cyber insurance renewal window',
            'PE sponsor requiring SOC 2 / security program',
            'M&A prep (buy or sell-side)',
        ],
        'services': 'My Continuity, My Compliance suite, co-signatory on all major deals',
        'objections': [
            '"This wasn\'t in the budget"',
            '"Cheaper alternatives exist"',
        ],
        'criteria': [
            'SOC 2 Type II on vendor (audit-ready)',
            '$5M+ cyber COI + E&O insurance',
            'Written SLAs with credit / remedy structure',
        ],
        'headline': 'OpEx-predictable IT + compliance that pleases your auditor, your board, and your carrier.',
        'quote': 'My cyber renewal questionnaire this year had 60 questions. Our MSP filled out 35. I\'m signing an attestation on the other 25.',
        'holes': 'AFP Annual | FEI Summit / CFO Leadership Council | CFO.com / WSJ CFO Journal',
    },
]


# ==========================================================================
# Fetch logo
# ==========================================================================
def fetch_logo():
    try:
        r = requests.get(LOGO_URL, timeout=10)
        r.raise_for_status()
        return BytesIO(r.content)
    except Exception as e:
        print(f"  [warn] Logo fetch failed: {e}")
        return None


# ==========================================================================
# Document Builders
# ==========================================================================
def set_default_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Open Sans'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    for k in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
        rFonts.set(qn(f'w:{k}'), 'Open Sans')


def set_page_margins(section, top=0.5, bottom=0.5, left=0.6, right=0.6):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def build_hero(doc, logo_stream):
    """Dark cover page with Technijian logo, title, and stats."""
    # Hero is a 1x1 table with dark fill, spanning content width
    hero = doc.add_table(rows=1, cols=1)
    hero.alignment = WD_TABLE_ALIGNMENT.CENTER
    hero.autofit = False
    remove_table_borders(hero)
    set_col_widths(hero, [7.3])
    cell = hero.rows[0].cells[0]
    shade_cell(cell, C_DARK)
    set_cell_margins(cell, top=360, left=400, bottom=360, right=400)
    set_row_height(hero.rows[0], 480)

    # Wordmark (text treatment — logo JPG with white bg looks awkward on dark)
    logo_para = cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(logo_para, after=0, line=1.0)
    add_run(logo_para, 'TECHNIJIAN', bold=True, size=26, color=RGB_WHITE)
    add_run(logo_para, '.', bold=True, size=26, color=RGB_ORANGE)
    tag_para = cell.add_paragraph()
    set_paragraph_spacing(tag_para, before=0, after=12, line=1.0)
    add_run(tag_para, 'TECHNOLOGY AS A SOLUTION', bold=True, size=7, color=RGB_TEAL)

    # Badge
    badge_para = cell.add_paragraph()
    set_paragraph_spacing(badge_para, before=18, after=6)
    add_run(badge_para, 'CUSTOMER PERSONAS  |  SALES + MARKETING REFERENCE',
            bold=True, size=9, color=RGB_TEAL)

    # Title
    title_para = cell.add_paragraph()
    set_paragraph_spacing(title_para, before=2, after=4, line=1.1)
    add_run(title_para, 'Who We Sell To', bold=True, size=34, color=RGB_WHITE)

    # Subtitle
    sub_para = cell.add_paragraph()
    set_paragraph_spacing(sub_para, before=0, after=10, line=1.25)
    add_run(sub_para,
            'Thirteen buyer profiles built for sales plays, campaign targeting, messaging, and pricing. ',
            size=12, color=RGB_WHITE)
    add_run(sub_para,
            'Every persona is validated against a real Technijian win.',
            bold=True, size=12, color=RGB_ORANGE)

    # Stats bar (nested table)
    stats_container = cell.add_paragraph()
    set_paragraph_spacing(stats_container, before=8, after=0)
    add_horizontal_bar(stats_container, hex_color='F67D4B', height_pt=1)

    stats_tbl = cell.add_table(rows=1, cols=4)
    remove_table_borders(stats_tbl)
    stats_tbl.autofit = False
    stats_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    stats = [
        ('13', 'Total Personas'),
        ('10', 'Industry Verticals'),
        ('6', 'Tier-1 Primary ICPs'),
        ('3', 'Horizontal Buyer Roles'),
    ]
    col_widths = [1.7, 1.7, 1.7, 1.7]
    for i, (num, label) in enumerate(stats):
        c = stats_tbl.rows[0].cells[i]
        shade_cell(c, C_DARK)
        set_cell_margins(c, top=120, left=80, bottom=120, right=80)
        c.width = Inches(col_widths[i])
        p1 = c.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p1, after=2)
        add_run(p1, num, bold=True, size=26, color=RGB_ORANGE)
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p2, after=0)
        add_run(p2, label.upper(), bold=True, size=8, color=RGB_WHITE)


def build_section_title(doc, eyebrow, title, subtitle=None, color=C_BLUE):
    """Colored section divider + heading."""
    # Thin colored bar
    bar_para = doc.add_paragraph()
    set_paragraph_spacing(bar_para, before=8, after=0)
    add_horizontal_bar(bar_para, hex_color=color, height_pt=3)

    # Eyebrow
    eye_para = doc.add_paragraph()
    set_paragraph_spacing(eye_para, before=8, after=2)
    add_run(eye_para, eyebrow.upper(),
            bold=True, size=9, color=RGBColor.from_string(color))

    # Title
    t_para = doc.add_paragraph()
    set_paragraph_spacing(t_para, before=0, after=2, line=1.1)
    add_run(t_para, title, bold=True, size=20, color=RGB_DARK)

    # Subtitle
    if subtitle:
        s_para = doc.add_paragraph()
        set_paragraph_spacing(s_para, before=0, after=8, line=1.35)
        add_run(s_para, subtitle, size=11, color=RGB_GREY)


def build_exec_summary(doc):
    build_section_title(
        doc,
        'Executive Summary',
        'Why these personas, how to use them',
        'Technijian sells on two axes - industry vertical (how they buy) and buyer role (who signs). Every deal addresses at least one of each.',
        color=C_BLUE,
    )

    # Two-axes explainer in a 2-col table
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    remove_table_borders(tbl)
    set_col_widths(tbl, [3.6, 3.6])
    set_row_cant_split(tbl.rows[0])

    left = tbl.rows[0].cells[0]
    shade_cell(left, C_OFFWHITE)
    set_cell_margins(left, top=200, left=220, bottom=200, right=220)
    set_cell_borders(left, left={'sz': '32', 'val': 'single', 'color': C_BLUE})
    p = left.paragraphs[0]
    set_paragraph_spacing(p, after=4)
    add_run(p, 'VERTICAL AXIS', bold=True, size=9, color=RGB_BLUE)
    p = left.add_paragraph()
    set_paragraph_spacing(p, after=4)
    add_run(p, 'Industry-primary decision-maker', bold=True, size=13, color=RGB_DARK)
    p = left.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.35)
    add_run(p,
        'Regulation, operations, and language differ so radically by industry that a healthcare admin and a defense manufacturer buy cybersecurity differently. 10 vertical personas capture this. Lead the pitch with the vertical.',
        size=10, color=RGB_GREY)

    right = tbl.rows[0].cells[1]
    shade_cell(right, C_OFFWHITE)
    set_cell_margins(right, top=200, left=220, bottom=200, right=220)
    set_cell_borders(right, left={'sz': '32', 'val': 'single', 'color': C_ORANGE})
    p = right.paragraphs[0]
    set_paragraph_spacing(p, after=4)
    add_run(p, 'HORIZONTAL AXIS', bold=True, size=9, color=RGB_ORANGE)
    p = right.add_paragraph()
    set_paragraph_spacing(p, after=4)
    add_run(p, 'Cross-industry buyer role', bold=True, size=13, color=RGB_DARK)
    p = right.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.35)
    add_run(p,
        'Inside any account, multiple people shape the purchase: CFO (economic buyer), IT Director (co-managed technical counterpart), and Marketing Director (growth-side). 3 horizontal personas capture these - loop them in early.',
        size=10, color=RGB_GREY)

    # Usage grid
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10, after=6)
    add_run(p, 'HOW TO USE', bold=True, size=9, color=RGB_TEAL)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=8, line=1.15)
    add_run(p, 'Four plays, same personas', bold=True, size=15, color=RGB_DARK)

    use_tbl = doc.add_table(rows=1, cols=4)
    use_tbl.autofit = False
    remove_table_borders(use_tbl)
    set_col_widths(use_tbl, [1.8, 1.8, 1.8, 1.8])
    set_row_cant_split(use_tbl.rows[0])
    uses = [
        ('Sales', C_BLUE,
         'Match prospect to 1 vertical + 1 horizontal on call 1. Use Pain Points as interview script. Anti-Persona flags disqualify fast.'),
        ('Marketing', C_ORANGE,
         'LinkedIn title/industry filters come from Identity + Organization. Ad hooks come from Top Pains. CTAs come from Triggers.'),
        ('Product', C_TEAL,
         'Pricing tiers, SLA structure, reporting formats, and contract preferences in Buying Behavior drive packaging decisions.'),
        ('Account Plan', C_CHARTREUSE,
         'Success Metrics + Communication Prefs set QBR cadence. Key Stakeholders + Influencers identify the account-plan map.'),
    ]
    for i, (title, color, body) in enumerate(uses):
        c = use_tbl.rows[0].cells[i]
        shade_cell(c, C_WHITE)
        set_cell_margins(c, top=160, left=160, bottom=160, right=160)
        set_cell_borders(c,
                         top={'sz': '24', 'val': 'single', 'color': color},
                         left={'sz': '4', 'val': 'single', 'color': C_LIGHTGREY},
                         bottom={'sz': '4', 'val': 'single', 'color': C_LIGHTGREY},
                         right={'sz': '4', 'val': 'single', 'color': C_LIGHTGREY})
        p = c.paragraphs[0]
        set_paragraph_spacing(p, after=4)
        add_run(p, title.upper(), bold=True, size=11,
                color=RGBColor.from_string(color))
        p = c.add_paragraph()
        set_paragraph_spacing(p, after=0, line=1.3)
        add_run(p, body, size=9, color=RGB_GREY)


def build_tier_divider(doc, tier_label, tier_description, color_hex):
    """Full-width colored band announcing a tier."""
    # Force page break via paragraph format (avoids trailing-empty-page)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.page_break_before = True
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)

    band = doc.add_table(rows=1, cols=1)
    band.autofit = False
    remove_table_borders(band)
    set_col_widths(band, [7.3])
    cell = band.rows[0].cells[0]
    shade_cell(cell, C_DARK)
    set_cell_margins(cell, top=260, left=300, bottom=260, right=300)

    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=4)
    add_run(p, tier_label.upper(), bold=True, size=10,
            color=RGBColor.from_string(color_hex))
    p = cell.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.3)
    add_run(p, tier_description, bold=True, size=14, color=RGB_WHITE)


def build_persona_card(doc, p_data, is_first_in_tier=False):
    """One persona per page to prevent split across breaks."""
    # Small top gap - use page_break_before to avoid leaving blank trailing page
    gap = doc.add_paragraph()
    if not is_first_in_tier:
        gap.paragraph_format.page_break_before = True
    set_paragraph_spacing(gap, before=4, after=0)

    # Outer card = single-col table with colored top accent
    card = doc.add_table(rows=1, cols=1)
    card.autofit = False
    remove_table_borders(card)
    set_col_widths(card, [7.3])
    wrap = card.rows[0].cells[0]
    shade_cell(wrap, C_WHITE)
    set_cell_margins(wrap, top=0, left=0, bottom=0, right=0)

    # Determine tier color
    tier_color = C_BLUE
    if p_data['tier'] == 'Tier 1':
        tier_color = C_BLUE
    elif p_data['tier'] == 'Tier 2':
        tier_color = C_ORANGE
    elif p_data['tier'] == 'Tier 3':
        tier_color = C_TEAL

    # 1. Dark header bar with persona # + name + title + tier badge
    hdr_tbl = wrap.add_table(rows=1, cols=3)
    hdr_tbl.autofit = False
    remove_table_borders(hdr_tbl)
    for idx, w in enumerate([0.7, 5.1, 1.5]):
        hdr_tbl.rows[0].cells[idx].width = Inches(w)

    # Number cell - tier-colored
    num_c = hdr_tbl.rows[0].cells[0]
    shade_cell(num_c, tier_color)
    set_cell_margins(num_c, top=140, left=120, bottom=140, right=120)
    p = num_c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0)
    add_run(p, p_data['num'], bold=True, size=20, color=RGB_WHITE)

    # Name+title cell - dark
    name_c = hdr_tbl.rows[0].cells[1]
    shade_cell(name_c, C_DARK)
    set_cell_margins(name_c, top=120, left=180, bottom=120, right=120)
    p = name_c.paragraphs[0]
    set_paragraph_spacing(p, after=1)
    add_run(p, p_data['name'], bold=True, size=14, color=RGB_WHITE)
    p = name_c.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.15)
    add_run(p, p_data['title'], size=9, color=RGB_TEAL)

    # Tier pill cell - dark
    tier_c = hdr_tbl.rows[0].cells[2]
    shade_cell(tier_c, C_DARK)
    set_cell_margins(tier_c, top=120, left=120, bottom=120, right=180)
    p = tier_c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, after=1)
    add_run(p, p_data['tier'].upper(), bold=True, size=9,
            color=RGBColor.from_string(tier_color))
    p = tier_c.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, after=0)
    add_run(p, 'ICP', bold=False, size=8, color=RGB_WHITE)

    # 2. Meta strip: industry + size + ARR
    meta_tbl = wrap.add_table(rows=1, cols=3)
    meta_tbl.autofit = False
    remove_table_borders(meta_tbl)
    for idx, w in enumerate([3.3, 2.2, 1.8]):
        meta_tbl.rows[0].cells[idx].width = Inches(w)
    meta_vals = [
        ('INDUSTRY', p_data['industry'], C_BLUE),
        ('COMPANY SIZE', p_data['size'], C_TEAL),
        ('YEAR-1 ARR', p_data['arr'], C_ORANGE),
    ]
    for i, (lab, val, col) in enumerate(meta_vals):
        c = meta_tbl.rows[0].cells[i]
        shade_cell(c, C_OFFWHITE)
        set_cell_margins(c, top=110, left=160, bottom=110, right=160)
        set_cell_borders(c,
                         bottom={'sz': '4', 'val': 'single', 'color': C_LIGHTGREY},
                         right={'sz': '4', 'val': 'single', 'color': C_LIGHTGREY} if i < 2 else None)
        p = c.paragraphs[0]
        set_paragraph_spacing(p, after=1)
        add_run(p, lab, bold=True, size=7, color=RGBColor.from_string(col))
        p = c.add_paragraph()
        set_paragraph_spacing(p, after=0, line=1.2)
        add_run(p, val, bold=True, size=9, color=RGB_DARK)

    # 3. Two-column body: pains | triggers/services
    body_tbl = wrap.add_table(rows=1, cols=2)
    body_tbl.autofit = False
    remove_table_borders(body_tbl)
    set_col_widths(body_tbl, [3.65, 3.65])

    # LEFT: pains + triggers
    left = body_tbl.rows[0].cells[0]
    shade_cell(left, C_WHITE)
    set_cell_margins(left, top=160, left=160, bottom=100, right=120)
    set_cell_borders(left,
                     right={'sz': '4', 'val': 'single', 'color': C_LIGHTGREY},
                     bottom={'sz': '4', 'val': 'single', 'color': C_LIGHTGREY})
    p = left.paragraphs[0]
    set_paragraph_spacing(p, after=2)
    add_run(p, 'TOP PAIN POINTS', bold=True, size=8, color=RGB_BLUE)
    for pain in p_data['pains']:
        p = left.add_paragraph()
        set_paragraph_spacing(p, after=2, line=1.25)
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        add_run(p, 'v  ', bold=True, size=9, color=RGB_TEAL, font='Segoe UI Symbol')
        add_run(p, pain, size=9, color=RGB_DARK)

    p = left.add_paragraph()
    set_paragraph_spacing(p, before=6, after=2)
    add_run(p, 'TRIGGERS & CATALYSTS', bold=True, size=8, color=RGB_ORANGE)
    for trig in p_data['triggers']:
        p = left.add_paragraph()
        set_paragraph_spacing(p, after=1, line=1.25)
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        add_run(p, '>  ', bold=True, size=9, color=RGB_ORANGE)
        add_run(p, trig, size=9, color=RGB_DARK)

    # RIGHT: services + objections + criteria
    right = body_tbl.rows[0].cells[1]
    shade_cell(right, C_WHITE)
    set_cell_margins(right, top=160, left=160, bottom=100, right=160)
    set_cell_borders(right,
                     bottom={'sz': '4', 'val': 'single', 'color': C_LIGHTGREY})
    p = right.paragraphs[0]
    set_paragraph_spacing(p, after=2)
    add_run(p, 'PRIMARY TECHNIJIAN SERVICES', bold=True, size=8, color=RGB_TEAL)
    p = right.add_paragraph()
    set_paragraph_spacing(p, after=5, line=1.3)
    add_run(p, p_data['services'], bold=True, size=9, color=RGB_DARK)

    p = right.add_paragraph()
    set_paragraph_spacing(p, before=2, after=2)
    add_run(p, 'TOP OBJECTIONS', bold=True, size=8, color=RGBColor.from_string(C_CHARTREUSE))
    for obj in p_data['objections']:
        p = right.add_paragraph()
        set_paragraph_spacing(p, after=1, line=1.25)
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        add_run(p, '-  ', bold=True, size=9, color=RGB_GREY)
        add_run(p, obj, italic=True, size=9, color=RGB_DARK)

    p = right.add_paragraph()
    set_paragraph_spacing(p, before=5, after=2)
    add_run(p, 'KEY DECISION CRITERIA', bold=True, size=8, color=RGB_BLUE)
    for cri in p_data['criteria']:
        p = right.add_paragraph()
        set_paragraph_spacing(p, after=1, line=1.25)
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        add_run(p, 'v  ', bold=True, size=9, color=RGB_TEAL, font='Segoe UI Symbol')
        add_run(p, cri, size=9, color=RGB_DARK)

    # 4. Messaging headline bar (blue)
    msg_tbl = wrap.add_table(rows=1, cols=1)
    msg_tbl.autofit = False
    remove_table_borders(msg_tbl)
    set_col_widths(msg_tbl, [7.3])
    mc = msg_tbl.rows[0].cells[0]
    shade_cell(mc, C_BLUE)
    set_cell_margins(mc, top=130, left=180, bottom=130, right=180)
    p = mc.paragraphs[0]
    set_paragraph_spacing(p, after=1)
    add_run(p, 'MESSAGING HEADLINE', bold=True, size=7, color=RGBColor(0xCB, 0xDB, 0x2D))
    p = mc.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.25)
    add_run(p, '"' + p_data['headline'] + '"', bold=True, size=11, color=RGB_WHITE)

    # 5. Killer quote + watering holes strip
    bot_tbl = wrap.add_table(rows=1, cols=2)
    bot_tbl.autofit = False
    remove_table_borders(bot_tbl)
    set_col_widths(bot_tbl, [4.4, 2.9])

    # Quote cell
    q = bot_tbl.rows[0].cells[0]
    shade_cell(q, C_OFFWHITE)
    set_cell_margins(q, top=140, left=200, bottom=140, right=180)
    set_cell_borders(q,
                     left={'sz': '28', 'val': 'single', 'color': C_ORANGE},
                     right={'sz': '4', 'val': 'single', 'color': C_LIGHTGREY})
    p = q.paragraphs[0]
    set_paragraph_spacing(p, after=2)
    add_run(p, 'VERBATIM QUOTE', bold=True, size=7, color=RGB_ORANGE)
    p = q.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.3)
    add_run(p, '"' + p_data['quote'] + '"', italic=True, size=9, color=RGB_DARK)

    # Watering holes cell
    w = bot_tbl.rows[0].cells[1]
    shade_cell(w, C_OFFWHITE)
    set_cell_margins(w, top=140, left=160, bottom=140, right=180)
    p = w.paragraphs[0]
    set_paragraph_spacing(p, after=2)
    add_run(p, 'WATERING HOLES', bold=True, size=7, color=RGB_TEAL)
    p = w.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.3)
    add_run(p, p_data['holes'], size=8, color=RGB_GREY)


def build_messaging_matrix(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.page_break_before = True
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)

    build_section_title(
        doc,
        'Quick-Reference',
        'Messaging Matrix',
        'One headline per persona. Use as LinkedIn ad hooks, subject lines, and opening lines in outbound sequences.',
        color=C_TEAL,
    )

    tbl = doc.add_table(rows=len(PERSONAS) + 1, cols=4)
    tbl.autofit = False
    remove_table_borders(tbl)
    widths = [0.5, 1.9, 3.3, 1.6]
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i])

    # Header
    headers = ['#', 'PERSONA', 'MESSAGING HEADLINE', 'YR-1 ARR']
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        shade_cell(c, C_DARK)
        set_cell_margins(c, top=70, left=120, bottom=70, right=120)
        p = c.paragraphs[0]
        set_paragraph_spacing(p, after=0)
        add_run(p, h, bold=True, size=8, color=RGB_WHITE)
    set_row_cant_split(tbl.rows[0])

    for r, p_data in enumerate(PERSONAS, start=1):
        row = tbl.rows[r]
        set_row_cant_split(row)
        is_alt = (r % 2 == 0)
        row_color = C_OFFWHITE if is_alt else C_WHITE
        tier_color = {
            'Tier 1': C_BLUE, 'Tier 2': C_ORANGE, 'Tier 3': C_TEAL
        }.get(p_data['tier'], C_BLUE)

        # #
        c = row.cells[0]
        shade_cell(c, tier_color)
        set_cell_margins(c, top=60, left=80, bottom=60, right=80)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=0)
        add_run(p, p_data['num'], bold=True, size=10, color=RGB_WHITE)

        # Persona name + title
        c = row.cells[1]
        shade_cell(c, row_color)
        set_cell_margins(c, top=60, left=120, bottom=60, right=120)
        p = c.paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.1)
        add_run(p, p_data['name'], bold=True, size=9, color=RGB_DARK)
        p = c.add_paragraph()
        set_paragraph_spacing(p, after=0, line=1.1)
        add_run(p, p_data['title'], size=7, color=RGB_GREY)

        # Headline
        c = row.cells[2]
        shade_cell(c, row_color)
        set_cell_margins(c, top=60, left=120, bottom=60, right=120)
        p = c.paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.2)
        add_run(p, p_data['headline'], size=8, color=RGB_DARK)

        # ARR
        c = row.cells[3]
        shade_cell(c, row_color)
        set_cell_margins(c, top=60, left=120, bottom=60, right=120)
        p = c.paragraphs[0]
        set_paragraph_spacing(p, after=0)
        add_run(p, p_data['arr'], bold=True, size=8, color=RGB_ORANGE)


def build_playbook(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.page_break_before = True
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)

    build_section_title(
        doc,
        'Campaign Playbook',
        'Where these personas gather',
        'Use this to plan events, sponsorship, content placement, geo-fencing, and ad targeting.',
        color=C_ORANGE,
    )

    playbook = [
        ('Healthcare', '02', C_BLUE, 'MGMA Annual (Oct), OCMA chapter, HIPAA Journal, Becker\'s ASC Review, AAMA local.'),
        ('Financial Services', '03', C_BLUE, 'NSCP Annual (Oct), Schwab IMPACT (Nov), IA Watch, FINRA Annual, ALI Wealth Mgmt.'),
        ('Defense / DIB', '04', C_BLUE, 'NDIA, AIA / AUSA, CMMC Day, CMMC-AB events, local DIB supplier forums.'),
        ('SaaS / Startups', '05', C_BLUE, 'AWS re:Invent (Dec), Pragmatic Engineer, Hacker News, Rands Slack, SaaStr.'),
        ('Legal', '06', C_BLUE, 'ILTACON (Aug), ABA Techshow (Mar), ALA / OCBA chapters, LegalTech NY.'),
        ('Real Estate', '07', C_ORANGE, 'Inman Connect (Jan/Aug), NAR Annual (Nov), Lab Coat Agents FB, T3 Sixty Summit.'),
        ('Non-Profit', '08', C_TEAL, 'CalNonprofits Annual, NTEN Nonprofit Tech Conf, Chronicle of Philanthropy.'),
        ('Hospitality', '09', C_TEAL, 'National Restaurant Show (Chicago, May), CRA local, Nation\'s Restaurant News.'),
        ('Construction', '10', C_ORANGE, 'Intl. Builders\' Show (Feb), HBA of OC, Builder Magazine, AGC local.'),
        ('SMB OC Generalist', '01', C_ORANGE, 'Vistage / EO, OC Business Journal, Chamber of Commerce (Irvine, Newport, Anaheim).'),
        ('IT Director (Co-Mg)', '11', C_BLUE, 'SIM local chapter, Gartner IT Symposium, ISACA / ISSA OC chapter, HDI.'),
        ('Marketing Director', '12', C_TEAL, 'Inbound HubSpot (Sep), Content Marketing World (Sep), MarketingProfs, MozCon.'),
        ('CFO / Controller', '13', C_ORANGE, 'AFP Annual, FEI Summit, CFO Leadership Council, CFO.com, WSJ CFO Journal.'),
    ]

    tbl = doc.add_table(rows=len(playbook) + 1, cols=3)
    tbl.autofit = False
    remove_table_borders(tbl)
    widths = [2.0, 0.7, 4.6]
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i])

    # Header
    for i, h in enumerate(['SEGMENT', 'PERSONA', 'WATERING HOLES / EVENTS / PUBS']):
        c = tbl.rows[0].cells[i]
        shade_cell(c, C_DARK)
        set_cell_margins(c, top=100, left=140, bottom=100, right=140)
        p = c.paragraphs[0]
        set_paragraph_spacing(p, after=0)
        add_run(p, h, bold=True, size=9, color=RGB_WHITE)

    for r, (seg, num, color, holes) in enumerate(playbook, start=1):
        is_alt = (r % 2 == 0)
        row_bg = C_OFFWHITE if is_alt else C_WHITE
        set_row_cant_split(tbl.rows[r])

        c = tbl.rows[r].cells[0]
        shade_cell(c, row_bg)
        set_cell_margins(c, top=100, left=140, bottom=100, right=140)
        set_cell_borders(c, left={'sz': '24', 'val': 'single', 'color': color})
        p = c.paragraphs[0]
        set_paragraph_spacing(p, after=0)
        add_run(p, seg, bold=True, size=10, color=RGB_DARK)

        c = tbl.rows[r].cells[1]
        shade_cell(c, row_bg)
        set_cell_margins(c, top=100, left=140, bottom=100, right=140)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, after=0)
        add_run(p, '#' + num, bold=True, size=10,
                color=RGBColor.from_string(color))

        c = tbl.rows[r].cells[2]
        shade_cell(c, row_bg)
        set_cell_margins(c, top=100, left=140, bottom=100, right=140)
        p = c.paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.3)
        add_run(p, holes, size=9, color=RGB_GREY)

    # Seasonal callout — force to new page so header + row stay together cleanly
    spacer2 = doc.add_paragraph()
    spacer2.paragraph_format.page_break_before = True
    spacer2.paragraph_format.space_before = Pt(0)
    spacer2.paragraph_format.space_after = Pt(0)

    # Reuse section title treatment for seasonal
    bar_p = doc.add_paragraph()
    set_paragraph_spacing(bar_p, before=4, after=0)
    add_horizontal_bar(bar_p, hex_color=C_TEAL, height_pt=3)
    bar_p.paragraph_format.keep_with_next = True

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=2)
    add_run(p, 'SEASONAL TIMING', bold=True, size=9, color=RGB_TEAL)
    p.paragraph_format.keep_with_next = True

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=4, line=1.1)
    add_run(p, 'Buying windows by quarter', bold=True, size=20, color=RGB_DARK)
    p.paragraph_format.keep_with_next = True

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=10, line=1.35)
    add_run(p, 'Where each persona tier is most reachable - time your outbound, events, and email cadence accordingly.',
            size=11, color=RGB_GREY)
    p.paragraph_format.keep_with_next = True

    season_tbl = doc.add_table(rows=1, cols=4)
    season_tbl.autofit = False
    remove_table_borders(season_tbl)
    set_col_widths(season_tbl, [1.8, 1.8, 1.8, 1.8])
    set_row_cant_split(season_tbl.rows[0])
    seasons = [
        ('Q1 (Jan-Mar)', 'Budget reset, insurance renewals, tax-season MSP stress. SMB + CFO personas active.', C_BLUE),
        ('Q2 (Apr-Jun)', 'Post-tax CPA/attorney buying, compliance prep for summer audits. Healthcare + Fin Svcs peaks.', C_ORANGE),
        ('Q3 (Jul-Sep)', 'Back-to-school & fiscal reset. ILTACON, Inbound, NSCP prep. Legal + Marketing personas most receptive.', C_TEAL),
        ('Q4 (Oct-Dec)', 'Renewal crunch, board reviews, year-end security spend. Every persona active, Nov-Dec is execution window.', C_CHARTREUSE),
    ]
    for i, (q, desc, col) in enumerate(seasons):
        c = season_tbl.rows[0].cells[i]
        shade_cell(c, C_OFFWHITE)
        set_cell_margins(c, top=140, left=160, bottom=140, right=160)
        set_cell_borders(c, top={'sz': '24', 'val': 'single', 'color': col})
        p = c.paragraphs[0]
        set_paragraph_spacing(p, after=4)
        add_run(p, q, bold=True, size=10, color=RGBColor.from_string(col))
        p = c.add_paragraph()
        set_paragraph_spacing(p, after=0, line=1.3)
        add_run(p, desc, size=9, color=RGB_GREY)


def build_cta_and_footer(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.page_break_before = True
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)

    # CTA band
    cta = doc.add_table(rows=1, cols=1)
    cta.autofit = False
    remove_table_borders(cta)
    set_col_widths(cta, [7.3])
    cc = cta.rows[0].cells[0]
    shade_cell(cc, C_BLUE)
    set_cell_margins(cc, top=360, left=360, bottom=360, right=360)

    p = cc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=4)
    add_run(p, 'FOR THE MARKETING TEAM', bold=True, size=9,
            color=RGBColor(0xCB, 0xDB, 0x2D))

    p = cc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=6, line=1.2)
    add_run(p, 'Start campaigns from this.', bold=True, size=22, color=RGB_WHITE)

    p = cc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=10, line=1.35)
    add_run(p,
        'The next sales play, LinkedIn audience, nurture sequence, or ABM account list should pull directly from these 13 personas. If a campaign doesn\'t map to a persona, it\'s a guess.',
        size=11, color=RGB_WHITE)

    p = cc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=6, after=0)
    add_run(p, 'Full attribute files (~25 fields each): ', size=10, color=RGB_WHITE)
    add_run(p, 'tech-branding/personas/vertical/ + /horizontal/',
            bold=True, size=10, color=RGBColor(0xCB, 0xDB, 0x2D))

    # About Technijian
    doc.add_paragraph()
    about_tbl = doc.add_table(rows=1, cols=2)
    about_tbl.autofit = False
    remove_table_borders(about_tbl)
    set_col_widths(about_tbl, [3.6, 3.6])

    left = about_tbl.rows[0].cells[0]
    shade_cell(left, C_OFFWHITE)
    set_cell_margins(left, top=220, left=220, bottom=220, right=220)
    set_cell_borders(left, left={'sz': '28', 'val': 'single', 'color': C_BLUE})
    p = left.paragraphs[0]
    set_paragraph_spacing(p, after=4)
    add_run(p, 'ABOUT TECHNIJIAN', bold=True, size=9, color=RGB_BLUE)
    p = left.add_paragraph()
    set_paragraph_spacing(p, after=6, line=1.2)
    add_run(p, '25+ years. 1,000+ networks. Orange County.',
            bold=True, size=14, color=RGB_DARK)
    p = left.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.35)
    add_run(p,
        'Pod-model managed IT, security, cloud, compliance, and AI services for regulated and high-growth organizations. Local, accountable, and audit-ready from day one.',
        size=10, color=RGB_GREY)

    right = about_tbl.rows[0].cells[1]
    shade_cell(right, C_DARK)
    set_cell_margins(right, top=220, left=220, bottom=220, right=220)
    p = right.paragraphs[0]
    set_paragraph_spacing(p, after=4)
    add_run(p, 'MAINTAINER', bold=True, size=9, color=RGB_TEAL)
    p = right.add_paragraph()
    set_paragraph_spacing(p, after=2, line=1.2)
    add_run(p, 'Ravi Jain', bold=True, size=14, color=RGB_WHITE)
    p = right.add_paragraph()
    set_paragraph_spacing(p, after=6, line=1.2)
    add_run(p, 'rjain@technijian.com', size=10, color=RGB_TEAL)
    p = right.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.3)
    add_run(p,
        'File a new persona or flag a drift when a closed deal doesn\'t match any existing profile. Review quarterly.',
        size=9, color=RGB_WHITE)

    # Footer
    doc.add_paragraph()
    footer = doc.add_table(rows=1, cols=1)
    footer.autofit = False
    remove_table_borders(footer)
    set_col_widths(footer, [7.3])
    fc = footer.rows[0].cells[0]
    shade_cell(fc, C_NEAR_BLACK)
    set_cell_margins(fc, top=180, left=260, bottom=180, right=260)
    p = fc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=2)
    add_run(p, 'TECHNIJIAN', bold=True, size=10, color=RGB_WHITE)
    p = fc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0, line=1.3)
    add_run(p,
            '18 Technology Dr. Ste 141, Irvine CA 92618  |  949.379.8500  |  technijian.com  |  Customer Personas v1.0 - April 2026',
            size=8, color=RGB_WHITE)


# ==========================================================================
# Build document
# ==========================================================================
def build_document(out_path):
    doc = Document()
    set_default_style(doc)
    section = doc.sections[0]
    set_page_margins(section, top=0.5, bottom=0.5, left=0.6, right=0.6)

    # Page width 8.5" - 1.2" margins = 7.3" content width

    print("  - Fetching logo...")
    logo = fetch_logo()

    print("  - Building hero cover...")
    build_hero(doc, logo)

    # Small gap after hero
    gap = doc.add_paragraph()
    set_paragraph_spacing(gap, before=4, after=0)

    print("  - Building executive summary...")
    build_exec_summary(doc)

    # Tier 1 Section
    print("  - Tier 1 personas...")
    build_tier_divider(
        doc,
        'TIER 1 - PRIMARY ICPS',
        'Six personas. Highest Year-1 ARR ($40K-$300K). Compliance-heavy sales motions, technical depth required.',
        C_BLUE,
    )
    tier1_nums = ['02', '03', '04', '05', '06', '11']
    tier1 = [p for p in PERSONAS if p['num'] in tier1_nums]
    for i, pd in enumerate(tier1):
        build_persona_card(doc, pd, is_first_in_tier=(i == 0))

    # Tier 2 Section
    print("  - Tier 2 personas...")
    build_tier_divider(
        doc,
        'TIER 2 - PROVEN SECONDARY',
        'Four personas. ROI-anchored, faster cycles. CFO co-signatory on deals >$50K.',
        C_ORANGE,
    )
    tier2_nums = ['01', '07', '10', '13']
    tier2 = [p for p in PERSONAS if p['num'] in tier2_nums]
    for i, pd in enumerate(tier2):
        build_persona_card(doc, pd, is_first_in_tier=(i == 0))

    # Tier 3 Section
    print("  - Tier 3 personas...")
    build_tier_divider(
        doc,
        'TIER 3 - OPPORTUNISTIC / REFERRAL',
        'Three personas. Shorter cycles, smaller ARR ($18K-$120K). Often land inside another persona\'s account.',
        C_TEAL,
    )
    tier3_nums = ['08', '09', '12']
    tier3 = [p for p in PERSONAS if p['num'] in tier3_nums]
    for i, pd in enumerate(tier3):
        build_persona_card(doc, pd, is_first_in_tier=(i == 0))

    print("  - Messaging matrix...")
    build_messaging_matrix(doc)

    print("  - Campaign playbook...")
    build_playbook(doc)

    print("  - CTA + footer...")
    build_cta_and_footer(doc)

    doc.save(out_path)
    print(f"  = Saved: {out_path}")


if __name__ == '__main__':
    import sys
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, 'Technijian Customer Personas Datasheet.docx')
    build_document(out_path)
    print("\nDone.")
